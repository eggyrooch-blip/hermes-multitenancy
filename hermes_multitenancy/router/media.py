"""Media/artifact helpers split out of router god-node (pure move).

Shim helpers/state routed through ``_m`` for monkeypatch fidelity.
"""
from __future__ import annotations

import asyncio
import base64
import hashlib
import importlib
import ipaddress
import inspect
import json
import os
import re
import shutil
import socket
import zlib
import zipfile
from contextlib import asynccontextmanager
from itertools import zip_longest
from pathlib import Path
from xml.etree import ElementTree as ET
import urllib.parse
import urllib.request
from typing import Any, Optional

from .. import router as _m


def _materialize_response_artifacts(response: str, profile_home: Path) -> str:
    """Write model-emitted artifact JSON blocks into profile workspace files.

    The model has no direct filesystem tool in some Feishu profiles. This
    controlled bridge lets it emit file content declaratively while preserving
    the existing outbound MEDIA safety boundary.
    """
    text = str(response or "")
    if "```hermes-artifact-json" not in text.lower():
        return text
    root = profile_home.expanduser().resolve(strict=False)
    workspace = (root / "workspace").resolve(strict=False)
    downloads = (workspace / "Downloads").resolve(strict=False)
    media_additions: list[str] = []

    for match in _m._ARTIFACT_JSON_RE.finditer(text):
        try:
            spec = json.loads(match.group("body"))
        except Exception as exc:
            _m.logger.debug("multitenancy: invalid artifact json skipped: %s", exc)
            continue
        if not isinstance(spec, dict):
            continue
        raw_path = str(spec.get("path") or "").strip()
        filename = str(spec.get("filename") or spec.get("name") or "").strip()
        if raw_path == "/workspace" or raw_path.startswith("/workspace/"):
            candidate = workspace / raw_path.removeprefix("/workspace").lstrip("/")
        else:
            if raw_path:
                candidate = Path(raw_path).expanduser()
            elif filename:
                candidate = downloads / filename
            else:
                continue
            if not candidate.is_absolute():
                candidate = workspace / candidate
        target = candidate.resolve(strict=False)
        if not (target == downloads or downloads in target.parents):
            _m.logger.warning("multitenancy: blocked response artifact outside workspace downloads path=%s", target)
            continue
        try:
            target.parent.mkdir(parents=True, exist_ok=True)
            existing_source = _m._existing_profile_source_for_artifact_spec(spec, target, root)
            if existing_source is not None:
                shutil.copy2(existing_source, target)
                _m.logger.info(
                    "multitenancy: reused existing profile artifact source=%s target=%s",
                    existing_source,
                    target,
                )
            else:
                _m._write_response_artifact(target, spec)
            try:
                media_path = "/workspace/" + str(target.relative_to(workspace).as_posix())
            except ValueError:
                media_path = "/workspace/Downloads/" + target.name
            if f"MEDIA:{media_path}" not in text:
                if (
                    bool(spec.get("as_document"))
                    and target.suffix.lower() not in _m._MARKDOWN_DOCUMENT_EXTENSIONS
                    and "[[as_document]]" not in text
                ):
                    media_additions.append("[[as_document]]")
                media_additions.append(f"MEDIA:{media_path}")
        except Exception as exc:
            _m.logger.warning("multitenancy: failed to materialize response artifact path=%s error=%s", target, exc)
    if not media_additions:
        return text
    return f"{text.rstrip()}\n" + "\n".join(media_additions)


def _append_remote_image_media_directives(response: str, profile_home: Path) -> str:
    """Attach public markdown image URLs as profile-local media artifacts."""
    text = str(response or "")
    if "![" not in text or "http" not in text.lower():
        return text
    additions: list[str] = []
    seen_urls: set[str] = set()
    existing_media_paths = {match.group("path").strip() for match in _m._MEDIA_DIRECTIVE_RE.finditer(text)}
    for match in _m._MARKDOWN_REMOTE_IMAGE_RE.finditer(text):
        url = match.group("url").strip()
        if url in seen_urls:
            continue
        seen_urls.add(url)
        image_path = _m._materialize_remote_image_url(url, profile_home)
        if image_path is None:
            continue
        path_text = str(image_path)
        if path_text in existing_media_paths:
            continue
        existing_media_paths.add(path_text)
        additions.append(f"MEDIA:{path_text}")
    if not additions:
        return text
    return f"{text.rstrip()}\n" + "\n".join(additions)


async def _append_remote_image_media_directives_async(response: str, profile_home: Path) -> str:
    """Async wrapper for remote image materialization.

    DNS and HTTP reads are blocking in the sync materializer. Keep those calls
    off the gateway event loop so Feishu/WebUI streaming stays responsive.
    """
    return await asyncio.to_thread(_m._append_remote_image_media_directives, response, profile_home)


def _materialize_remote_image_url(url: str, profile_home: Path) -> Optional[Path]:
    parsed = urllib.parse.urlparse(url)
    if not _m._is_public_remote_image_url(parsed):
        _m.logger.info("multitenancy: skipped remote image delivery for untrusted url host=%s", parsed.hostname or "")
        return None

    root = profile_home.expanduser().resolve(strict=False)
    target_dir = (root / "workspace" / "Downloads" / "remote-images").resolve(strict=False)
    url_ext = _m._remote_image_extension_from_url(parsed)
    digest = hashlib.sha256(url.encode("utf-8")).hexdigest()[:20]
    tentative_ext = url_ext or ".img"
    target = (target_dir / f"remote-image-{digest}{tentative_ext}").resolve(strict=False)
    if not (target == root or root in target.parents):
        return None
    if target.exists() and target.is_file():
        return target

    try:
        request = urllib.request.Request(url, headers={"User-Agent": "Hermes-Multitenancy/remote-image-delivery"})
        with urllib.request.urlopen(request, timeout=_m._REMOTE_IMAGE_DOWNLOAD_TIMEOUT_S) as response:
            final_url = str(getattr(response, "geturl", lambda: url)() or url)
            final_parsed = urllib.parse.urlparse(final_url)
            if final_url != url and not _m._is_public_remote_image_url(final_parsed):
                _m.logger.warning(
                    "multitenancy: blocked remote image redirect host=%s final_host=%s",
                    parsed.hostname or "",
                    final_parsed.hostname or "",
                )
                return None
            peer_ip = _m._remote_image_response_peer_ip(response)
            if not peer_ip or not _m._is_global_ip_address(peer_ip):
                _m.logger.warning(
                    "multitenancy: blocked remote image peer host=%s peer_ip=%s",
                    final_parsed.hostname or parsed.hostname or "",
                    peer_ip or "",
                )
                return None
            headers = getattr(response, "headers", {}) or {}
            content_type = _m._header_value(headers, "Content-Type").split(";", 1)[0].strip().lower()
            content_length = _m._header_value(headers, "Content-Length").strip()
            if content_length:
                try:
                    if int(content_length) > _m._REMOTE_IMAGE_DELIVERY_MAX_BYTES:
                        _m.logger.info(
                            "multitenancy: skipped oversized remote image host=%s size=%s",
                            parsed.hostname or "",
                            content_length,
                        )
                        return None
                except ValueError:
                    pass
            ext = _m._remote_image_extension_from_content_type(content_type) or url_ext
            if ext not in _m._MEDIA_IMAGE_EXTENSIONS:
                _m.logger.info(
                    "multitenancy: skipped remote image with unsupported content type host=%s content_type=%s",
                    parsed.hostname or "",
                    content_type,
                )
                return None
            data = response.read(_m._REMOTE_IMAGE_DELIVERY_MAX_BYTES + 1)
            if not data or len(data) > _m._REMOTE_IMAGE_DELIVERY_MAX_BYTES:
                _m.logger.info(
                    "multitenancy: skipped oversized or empty remote image host=%s size=%s",
                    parsed.hostname or "",
                    len(data),
                )
                return None
    except Exception as exc:
        _m.logger.warning(
            "multitenancy: failed to materialize remote image host=%s error=%s",
            parsed.hostname or "",
            exc,
        )
        return None

    if ext != tentative_ext:
        target = (target_dir / f"remote-image-{digest}{ext}").resolve(strict=False)
        if not (target == root or root in target.parents):
            return None
        if target.exists() and target.is_file():
            return target
    try:
        target_dir.mkdir(parents=True, exist_ok=True)
        tmp_target = target.with_name(target.name + ".tmp")
        tmp_target.write_bytes(data)
        tmp_target.replace(target)
        _m.logger.info("multitenancy: materialized remote image url host=%s path=%s", parsed.hostname or "", target)
        return target
    except Exception as exc:
        _m.logger.warning(
            "multitenancy: failed to write remote image artifact host=%s path=%s error=%s",
            parsed.hostname or "",
            target,
            exc,
        )
        return None


def _is_public_remote_image_url(parsed: urllib.parse.ParseResult) -> bool:
    if parsed.scheme.lower() not in {"http", "https"}:
        return False
    hostname = (parsed.hostname or "").strip().lower().rstrip(".")
    if not hostname:
        return False
    if hostname in {"localhost"} or hostname.endswith((".localhost", ".local", ".internal", ".lan")):
        return False
    try:
        addresses = socket.getaddrinfo(hostname, parsed.port or (443 if parsed.scheme == "https" else 80))
    except OSError as exc:
        _m.logger.info("multitenancy: remote image host resolution failed host=%s error=%s", hostname, exc)
        return False
    if not addresses:
        return False
    for address in addresses:
        sockaddr = address[4]
        if not sockaddr:
            return False
        if not _m._is_global_ip_address(str(sockaddr[0])):
            return False
    return True


def _is_global_ip_address(value: str) -> bool:
    try:
        ip = ipaddress.ip_address(value)
    except ValueError:
        return False
    return ip.is_global


def _remote_image_response_peer_ip(response: Any) -> Optional[str]:
    explicit = str(getattr(response, "_hermes_peer_ip", "") or "").strip()
    if explicit:
        return explicit
    seen: set[int] = set()
    pending: list[Any] = [response]
    for _ in range(32):
        if not pending:
            break
        obj = pending.pop(0)
        if obj is None:
            continue
        obj_id = id(obj)
        if obj_id in seen:
            continue
        seen.add(obj_id)
        getpeername = getattr(obj, "getpeername", None)
        if callable(getpeername):
            try:
                peer = getpeername()
            except OSError:
                peer = None
            if peer:
                return str(peer[0])
        for attr in ("fp", "raw", "_fp", "_sock", "sock"):
            try:
                child = getattr(obj, attr)
            except Exception:
                continue
            if child is not None and id(child) not in seen:
                pending.append(child)
    return None


def _remote_image_extension_from_url(parsed: urllib.parse.ParseResult) -> Optional[str]:
    path = urllib.parse.unquote(parsed.path or "")
    suffix = Path(path).suffix.lower()
    return suffix if suffix in _m._MEDIA_IMAGE_EXTENSIONS else None


def _remote_image_extension_from_content_type(content_type: str) -> Optional[str]:
    return _m._REMOTE_IMAGE_CONTENT_TYPE_EXTENSIONS.get(content_type.lower())


def _header_value(headers: Any, name: str) -> str:
    getter = getattr(headers, "get", None)
    if callable(getter):
        return str(getter(name) or getter(name.lower()) or "")
    try:
        return str(headers[name])
    except Exception:
        return ""


def _existing_profile_source_for_artifact_spec(spec: dict[str, Any], target: Path, profile_home: Path) -> Optional[Path]:
    """Find a real profile-local source when an artifact JSON block is only a placeholder.

    Some skills already write the complete markdown to `.ai-docs/...` and then
    emit a marker-only artifact JSON block with the same filename. In that case
    the marker is a delivery hint, not file content.
    """
    fmt = str(spec.get("format") or target.suffix.lstrip(".")).lower()
    if fmt not in {"md", "markdown"} and target.suffix.lower() not in _m._MARKDOWN_DOCUMENT_EXTENSIONS:
        return None
    if any(spec.get(key) not in (None, "") for key in ("content", "data", "rows")):
        return None
    filename = Path(str(spec.get("filename") or spec.get("name") or target.name)).name
    if not filename:
        return None
    root = profile_home.expanduser().resolve(strict=False)
    resolved_target = target.expanduser().resolve(strict=False)
    search_roots = (
        root / ".ai-docs",
        root / "home" / "Downloads",
        root / "data",
        root / "tmp",
    )
    for search_root in search_roots:
        if not search_root.exists():
            continue
        try:
            matches = search_root.rglob(filename) if search_root.is_dir() else [search_root]
            for candidate in matches:
                resolved = candidate.expanduser().resolve(strict=False)
                if resolved == resolved_target:
                    continue
                if _m._is_deliverable_profile_file(resolved, root):
                    return resolved
        except OSError:
            continue
    return None


def _write_response_artifact(path: Path, spec: dict[str, Any]) -> None:
    fmt = str(spec.get("format") or path.suffix.lstrip(".")).lower()
    marker = str(spec.get("marker") or "")
    if fmt in {"md", "markdown", "txt", "csv"}:
        content = str(spec.get("content") or marker or "")
        path.write_text(content, encoding="utf-8")
    elif fmt == "json":
        data = spec.get("data")
        if data is None:
            data = {"marker": marker, "content": spec.get("content") or marker}
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    elif fmt == "xlsx":
        _m._write_artifact_xlsx(path, spec)
    elif fmt == "docx":
        _m._write_artifact_docx(path, spec)
    elif fmt == "pdf":
        _m._write_artifact_pdf(path, spec)
    elif fmt in {"png", "jpg", "jpeg"}:
        _m._write_artifact_image(path, spec)
    else:
        content = str(spec.get("content") or marker or "")
        path.write_text(content, encoding="utf-8")


def _artifact_rows(spec: dict[str, Any]) -> list[list[Any]]:
    rows = spec.get("rows")
    if isinstance(rows, list) and all(isinstance(row, list) for row in rows):
        return rows
    marker = str(spec.get("marker") or "")
    return [["marker", "value", "note"], [marker, 42, "generated by Hermes artifact bridge"]]


def _write_artifact_xlsx(path: Path, spec: dict[str, Any]) -> None:
    from openpyxl import Workbook  # type: ignore

    wb = Workbook()
    ws = wb.active
    ws.title = str(spec.get("sheet") or "matrix")[:31]
    for row in _m._artifact_rows(spec):
        ws.append(row)
    wb.save(path)


def _write_artifact_docx(path: Path, spec: dict[str, Any]) -> None:
    from docx import Document  # type: ignore

    doc = Document()
    title = str(spec.get("title") or "Hermes generated document")
    doc.add_heading(title, level=1)
    content = str(spec.get("content") or spec.get("marker") or "")
    for paragraph in content.splitlines() or [content]:
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    doc.save(path)


def _write_artifact_pdf(path: Path, spec: dict[str, Any]) -> None:
    from reportlab.lib.pagesizes import letter  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore

    c = canvas.Canvas(str(path), pagesize=letter)
    y = 720
    for line in str(spec.get("content") or spec.get("marker") or "").splitlines() or [str(spec.get("marker") or "")]:
        c.drawString(72, y, line[:120])
        y -= 24
        if y < 72:
            c.showPage()
            y = 720
    c.save()


def _write_artifact_image(path: Path, spec: dict[str, Any]) -> None:
    from PIL import Image, ImageDraw  # type: ignore

    marker = str(spec.get("marker") or spec.get("content") or "")
    image = Image.new("RGB", (1000, 420), color=(245, 250, 255))
    draw = ImageDraw.Draw(image)
    draw.rectangle((40, 40, 960, 380), outline=(20, 120, 80), width=5)
    draw.text((70, 120), marker[:100], fill=(0, 0, 0))
    draw.text((70, 190), "Hermes generated image artifact", fill=(20, 120, 80))
    image.save(path)


def _append_profile_file_media_directives(response: str, profile_home: Path) -> str:
    """Attach profile-local files that the model mentioned as plain paths."""
    text = str(response or "")
    if not text:
        return text
    root = profile_home.expanduser().resolve(strict=False)
    media_paths = {match.group("path").strip() for match in _m._MEDIA_DIRECTIVE_RE.finditer(text)}
    additions: list[str] = []
    seen: set[Path] = set()
    for match in _m._PROFILE_FILE_PATH_RE.finditer(text):
        raw_path = match.group("path").strip().rstrip(".,;:)]}")
        if not raw_path or raw_path in media_paths:
            continue
        published = _m._publish_mentioned_profile_file(raw_path, root)
        if published is None or published in seen:
            continue
        seen.add(published)
        additions.append(f"MEDIA:{published}")
    if not additions:
        return text
    return f"{text.rstrip()}\n" + "\n".join(additions)


def _strip_plain_profile_file_paths_for_display(text: str, profile_home: Path) -> str:
    """Avoid exposing host absolute paths when the file will be attached."""
    raw = str(text or "")
    if not raw:
        return raw
    root = profile_home.expanduser().resolve(strict=False)

    def repl(match: re.Match[str]) -> str:
        raw_path = match.group("path").strip().rstrip(".,;:)]}")
        if not raw_path:
            return match.group(0)
        published = _m._publish_mentioned_profile_file(raw_path, root)
        if published is not None:
            if _m._should_deliver_as_feishu_document(published):
                return "[Markdown 源文件已自动发送]"
            return "[文件已作为附件发送]"
        candidate = Path(raw_path).expanduser()
        if raw_path == "/workspace" or raw_path.startswith("/workspace/"):
            candidate = root / "workspace" / raw_path.removeprefix("/workspace").lstrip("/")
        if not candidate.is_absolute():
            candidate = root / candidate
        resolved = candidate.resolve(strict=False)
        if resolved.exists() and (resolved == root or root in resolved.parents):
            return "[受保护文件路径已隐藏]"
        return match.group(0)

    return _m._PROFILE_FILE_PATH_RE.sub(repl, raw)


def _should_deliver_as_feishu_document(path: Path) -> bool:
    return path.suffix.lower() in _m._MARKDOWN_DOCUMENT_EXTENSIONS


def _append_media_denied_security_event(*, reason: str, path: Path | str) -> None:
    try:
        try:
            open_id = _m._current_sender_open_id()
        except Exception:
            open_id = None
        _m.append_security_event(
            event_type="media.denied",
            reason=reason,
            path=str(path),
            open_id=open_id,
        )
    except Exception:
        _m.logger.exception("multitenancy: failed to append media denied audit event")


def _profile_scoped_media_response(response: str, profile_home: Path) -> str:
    """Drop MEDIA directives outside profile scope and publish artifacts to workspace."""
    root = profile_home.expanduser().resolve(strict=False)

    def repl(match: re.Match[str]) -> str:
        raw_path = match.group("path").strip()
        if raw_path == "/workspace" or raw_path.startswith("/workspace/"):
            workspace_relative = raw_path.removeprefix("/workspace").lstrip("/")
            candidate = root / "workspace" / workspace_relative
        else:
            candidate = Path(raw_path).expanduser()
        if not candidate.is_absolute():
            candidate = root / candidate
        resolved = candidate.resolve(strict=False)
        if resolved == root or root in resolved.parents:
            if not _m._is_deliverable_profile_file(resolved, root):
                _m.logger.warning(
                    "multitenancy: blocked outbound MEDIA for non-deliverable profile file path=%s profile_home=%s",
                    resolved,
                    root,
                )
                _m._append_media_denied_security_event(
                    reason="non_deliverable_profile_file",
                    path=resolved,
                )
                return ""
            workspace_artifact = _m._publish_profile_media_artifact(resolved, root)
            deliver_path = workspace_artifact or resolved
            return f"{match.group('prefix')}{deliver_path}{match.group('suffix')}"
        profile_artifact = _m._resolve_profile_media_artifact(raw_path, root)
        if profile_artifact is not None:
            _m.logger.info(
                "multitenancy: rewrote outbound MEDIA to profile workspace artifact path=%s resolved=%s",
                raw_path,
                profile_artifact,
            )
            return f"{match.group('prefix')}{profile_artifact}{match.group('suffix')}"
        _m.logger.warning(
            "multitenancy: blocked outbound MEDIA outside profile home path=%s profile_home=%s",
            raw_path,
            root,
        )
        _m._append_media_denied_security_event(
            reason="outside_profile_home",
            path=raw_path,
        )
        return ""

    return _m._MEDIA_DIRECTIVE_RE.sub(repl, str(response or ""))


def _webui_profile_scoped_media_response(
    response: str,
    profile_home: Path,
    *,
    materialize_remote_images: bool = True,
) -> str:
    """Scope outbound MEDIA and expose workspace files through browser-safe aliases."""
    root = profile_home.expanduser().resolve(strict=False)
    response_with_remote_images = (
        _m._append_remote_image_media_directives(response, root)
        if materialize_remote_images
        else str(response or "")
    )
    scoped = _m._profile_scoped_media_response(response_with_remote_images, root)

    def repl(match: re.Match[str]) -> str:
        raw_path = match.group("path").strip()
        candidate = Path(raw_path).expanduser()
        if raw_path == "/workspace" or raw_path.startswith("/workspace/"):
            return match.group(0)
        if not candidate.is_absolute():
            candidate = root / candidate
        resolved = candidate.resolve(strict=False)
        workspace_root = (root / "workspace").resolve(strict=False)
        if resolved == workspace_root or workspace_root in resolved.parents:
            return f"{match.group('prefix')}{_m._workspace_alias_for_profile_file(resolved, root)}{match.group('suffix')}"
        return match.group(0)

    return _m._MEDIA_DIRECTIVE_RE.sub(repl, scoped)


def _publish_mentioned_profile_file(raw_path: str, profile_home: Path) -> Optional[Path]:
    if raw_path == "/workspace" or raw_path.startswith("/workspace/"):
        workspace_relative = raw_path.removeprefix("/workspace").lstrip("/")
        candidate = profile_home / "workspace" / workspace_relative
    else:
        candidate = Path(raw_path).expanduser()
    if not candidate.is_absolute():
        candidate = profile_home / candidate
    source = candidate.resolve(strict=False)
    if not _m._is_deliverable_profile_file(source, profile_home):
        return None
    workspace_root = (profile_home / "workspace").resolve(strict=False)
    if source == workspace_root or workspace_root in source.parents:
        return source
    target_dir = (workspace_root / "Downloads").resolve(strict=False)
    target = (target_dir / source.name).resolve(strict=False)
    if not (target == profile_home or profile_home in target.parents):
        return None
    try:
        target_dir.mkdir(parents=True, exist_ok=True)
        if source != target:
            shutil.copy2(source, target)
        _m.logger.info(
            "multitenancy: auto-attached mentioned profile file source=%s target=%s",
            source,
            target,
        )
        return target
    except Exception as exc:
        _m.logger.warning(
            "multitenancy: failed to auto-attach mentioned profile file source=%s target=%s error=%s",
            source,
            target,
            exc,
        )
        return None


def _is_deliverable_profile_file(source: Path, profile_home: Path) -> bool:
    root = profile_home.resolve(strict=False)
    if not (source.exists() and source.is_file() and root in source.parents):
        return False
    try:
        if source.stat().st_size > _m._AUTO_FILE_DELIVERY_MAX_BYTES:
            _m.logger.info(
                "multitenancy: skipped auto file delivery for oversized file path=%s size=%s",
                source,
                source.stat().st_size,
            )
            return False
    except OSError:
        return False
    relative_parts = source.relative_to(root).parts
    lowered = [part.lower() for part in relative_parts]
    if any(part in _m._SENSITIVE_PROFILE_DIR_NAMES for part in lowered[:-1]):
        _m.logger.warning("multitenancy: blocked auto file delivery for sensitive directory path=%s", source)
        return False
    name = lowered[-1] if lowered else ""
    if name in _m._SENSITIVE_PROFILE_FILE_NAMES:
        _m.logger.warning("multitenancy: blocked auto file delivery for sensitive file path=%s", source)
        return False
    return True


def _remember_recent_profile_file(profile_name: str, chat_id: str, path: Path, profile_home: Path) -> None:
    if not profile_name or not chat_id:
        return
    resolved = path.expanduser().resolve(strict=False)
    if not _m._is_deliverable_profile_file(resolved, profile_home):
        return
    key = (profile_name, chat_id)
    existing = [item for item in _m._recent_profile_files_by_chat.get(key, []) if item != str(resolved)]
    existing.append(str(resolved))
    _m._recent_profile_files_by_chat[key] = existing[-_m._RECENT_PROFILE_FILE_CONTEXT_MAX:]


def _should_append_recent_profile_file_context(text: str) -> bool:
    return bool(_m._RECENT_FILE_CONTEXT_TRIGGER_RE.search(str(text or "")))


def _workspace_alias_for_profile_file(path: Path, profile_home: Path) -> str:
    workspace_root = (profile_home / "workspace").resolve(strict=False)
    resolved = path.expanduser().resolve(strict=False)
    try:
        relative = resolved.relative_to(workspace_root)
    except ValueError:
        return str(resolved)
    return "/workspace/" + relative.as_posix()


def _recent_profile_files_from_history(prior_messages: list[dict], profile_home: Path) -> list[tuple[Path, Path]]:
    candidates: list[tuple[Path, Path]] = []
    seen: set[str] = set()
    for message in reversed(prior_messages[-_m._SESSION_HISTORY_MAX:]):
        if message.get("role") not in {"assistant", "tool"}:
            continue
        content = str(message.get("content") or "")
        for match in _m._PROFILE_FILE_PATH_RE.finditer(content):
            raw_path = match.group("path").strip().strip("`\"'")
            source = Path(raw_path).expanduser()
            if not source.is_absolute():
                source = profile_home / source
            source = source.resolve(strict=False)
            published = _m._publish_mentioned_profile_file(raw_path, profile_home)
            if published is None:
                continue
            resolved = published.resolve(strict=False)
            if str(resolved) in seen:
                continue
            seen.add(str(resolved))
            candidates.append((source, resolved))
            if len(candidates) >= _m._RECENT_PROFILE_FILE_CONTEXT_MAX:
                return candidates
    return candidates


def _append_recent_profile_file_context(
    text: str,
    *,
    profile_name: str,
    chat_id: str,
    profile_home: Path,
    prior_messages: list[dict],
) -> str:
    raw = str(text or "")
    if not _m._should_append_recent_profile_file_context(raw):
        return raw

    candidates: list[tuple[Path, Path]] = []
    seen: set[str] = set()
    for stored in _m._recent_profile_files_by_chat.get((profile_name, chat_id), []):
        path = Path(stored).expanduser().resolve(strict=False)
        if _m._is_deliverable_profile_file(path, profile_home) and str(path) not in seen:
            candidates.append((path, path))
            seen.add(str(path))

    for source, path in _m._recent_profile_files_from_history(prior_messages, profile_home):
        if str(path) not in seen:
            candidates.append((source, path))
            seen.add(str(path))
        if len(candidates) >= _m._RECENT_PROFILE_FILE_CONTEXT_MAX:
            break

    if not candidates:
        return raw

    lines = [
        "",
        "",
        "[Hermes context: 最近 Hermes 已投递给当前会话的文件]",
    ]
    for source, path in candidates[:_m._RECENT_PROFILE_FILE_CONTEXT_MAX]:
        lines.extend(
            [
                f"- file_name: {path.name}",
                f"  workspace_path: {_m._workspace_alias_for_profile_file(path, profile_home)}",
                f"  profile_path: {path}",
            ]
        )
        if source != path:
            lines.append(f"  source_path: {source}")
    lines.append("[/Hermes context]")
    return raw + "\n".join(lines)


def _resolve_profile_media_artifact(raw_path: str, profile_home: Path) -> Optional[Path]:
    """Map tool-reported temp media paths to same-name artifacts in the workspace."""
    name = Path(raw_path).name
    if not name:
        return None
    search_dirs = (
        profile_home / "home",
        profile_home / "home" / "Downloads",
        profile_home / "cache" / "images",
        profile_home / "tmp",
        profile_home / "data",
    )
    for directory in search_dirs:
        candidate = (directory / name).resolve(strict=False)
        if candidate.exists() and candidate.is_file() and profile_home in candidate.parents:
            return _m._publish_profile_media_artifact(candidate, profile_home)
    return None


def _publish_profile_media_artifact(source: Path, profile_home: Path) -> Optional[Path]:
    """Copy a profile-local generated artifact into the WebUI-visible workspace."""
    source = source.resolve(strict=False)
    root = profile_home.resolve(strict=False)
    if not (source.exists() and source.is_file() and root in source.parents):
        return None
    workspace_root = (root / "workspace").resolve(strict=False)
    if source == workspace_root or workspace_root in source.parents:
        return source
    artifact_dirs = (
        root / "home" / "Downloads",
        root / "cache" / "images",
        root / "tmp",
        root / "data",
    )
    direct_home_file = source.parent == (root / "home").resolve(strict=False)
    source_in_artifacts = any(
        source == directory.resolve(strict=False)
        or directory.resolve(strict=False) in source.parents
        for directory in artifact_dirs
    ) or direct_home_file
    if not source_in_artifacts:
        return None
    target_dir = (workspace_root / "Downloads").resolve(strict=False)
    target = (target_dir / source.name).resolve(strict=False)
    if not (target == root or root in target.parents):
        return None
    try:
        target_dir.mkdir(parents=True, exist_ok=True)
        if source != target:
            shutil.copy2(source, target)
        return target
    except Exception as exc:
        _m.logger.warning(
            "multitenancy: failed to publish media artifact to workspace source=%s target=%s error=%s",
            source,
            target,
            exc,
        )
        return None


def _event_has_image_media(event: Any) -> bool:
    media_urls = getattr(event, "media_urls", None) or []
    media_types = getattr(event, "media_types", None) or []
    if not media_urls:
        return False
    message_type_obj = getattr(event, "message_type", None)
    message_type_parts = [
        str(getattr(message_type_obj, "value", "") or ""),
        str(getattr(message_type_obj, "name", "") or ""),
        str(message_type_obj or ""),
    ]
    message_type = " ".join(message_type_parts).lower()
    if "photo" in message_type or "image" in message_type:
        return True
    for raw_path, raw_mtype in zip_longest(media_urls, media_types, fillvalue=""):
        raw = str(raw_path or "")
        suffix = Path(raw).suffix.lower()
        media_type = str(raw_mtype or "").lower()
        normalized = raw.replace("\\", "/").lower()
        if media_type.startswith("image") or suffix in _m._IMAGE_FILE_EXTENSIONS or "/cache/images/" in normalized:
            return True
    return False


def _image_vision_unavailable_response(event: Any, enriched_text: Optional[str]) -> Optional[str]:
    """Return a direct blocked reply when upstream image analysis is unavailable."""
    if not _m._event_has_image_media(event):
        return None
    text = str(enriched_text or "")
    lowered = text.lower()
    timeout_markers = (
        "vision auto-analysis timed out",
        "image preprocessing timed out",
    )
    failure_markers = (
        "something went wrong when i tried to look at it",
        "couldn't quite see it",
        "vision auto-analysis error",
    )
    is_timeout = any(marker in lowered for marker in timeout_markers)
    if not (is_timeout or any(marker in lowered for marker in failure_markers)):
        return None

    names: list[str] = []
    for raw_path in getattr(event, "media_urls", None) or []:
        name = Path(str(raw_path or "")).name
        if name:
            names.append(name)
    suffix = f"\n已收到图片附件：{', '.join(names[:3])}。" if names else ""
    message_id = _m._event_message_id(event)
    message_note = f"\nFeishu message_id: {message_id}" if message_id else ""
    if is_timeout:
        return (
            "无法读取图片内容：已收到图片附件，但自动视觉分析超时。"
            f"{suffix}{message_note}\n可以稍后重试，或提高图片预分析超时时间后再试。"
        )
    return (
        "无法读取图片内容：当前图片视觉分析不可用，vision_analyze provider rejected the request。"
        f"{suffix}{message_note}\n请修复 profile 的 vision provider/key 后重试。"
    )


def _image_prep_unavailable_note(event: Any, *, reason: str = "provider") -> str:
    paths = ", ".join(str(path) for path in (getattr(event, "media_urls", None) or [])[:3])
    suffix = f" using image_url: {paths}" if paths else ""
    message_id = _m._event_message_id(event)
    message_note = f" Feishu message_id: {message_id}." if message_id else ""
    if reason == "timeout":
        return (
            "[The user sent an image but vision auto-analysis timed out before vision_analyze completed. "
            "You can try again later or examine it yourself with vision_analyze"
            f"{suffix}.{message_note}]"
        )
    return (
        "[The user sent an image but something went wrong when I tried to look at it~ "
        "You can try examining it yourself with vision_analyze"
        f"{suffix}.{message_note}]"
    )


def _matching_custom_provider_entry(config: dict[str, Any], provider: str) -> Optional[dict[str, Any]]:
    """Return the custom provider entry that backs ``provider`` if present."""
    provider_l = str(provider or "").strip().lower()
    if not (provider_l == "custom" or provider_l.startswith("custom:")):
        return None
    custom_providers = config.get("custom_providers")
    if not isinstance(custom_providers, list):
        return None
    model_cfg = config.get("model") if isinstance(config.get("model"), dict) else {}
    model_base_url = str(model_cfg.get("base_url") or "").strip().rstrip("/")
    want_name = provider_l.split(":", 1)[1].strip() if ":" in provider_l else ""
    for entry in custom_providers:
        if not isinstance(entry, dict):
            continue
        name = str(entry.get("name") or "").strip().lower().replace(" ", "-")
        base_url = str(entry.get("base_url") or entry.get("url") or entry.get("api") or "").strip().rstrip("/")
        if want_name:
            matched = name == want_name
        else:
            matched = bool(model_base_url) and base_url == model_base_url
        if matched:
            return entry
    return None


def _profile_main_runtime_for_image_prep(profile_home: Path) -> Optional[dict[str, str]]:
    """Build a normalized main-model runtime for Hermes auxiliary image prep."""
    try:
        from ..agent_real import (
            _load_profile_config,
            _resolve_custom_provider_api_key,
            _split_model_spec,
        )
    except Exception as exc:
        _m.logger.debug("multitenancy: cannot import profile model helpers for image prep (%s)", exc)
        return None

    try:
        config = _load_profile_config(Path(profile_home))
    except Exception as exc:
        _m.logger.debug("multitenancy: cannot load profile config for image prep profile_home=%s (%s)", profile_home, exc)
        return None

    model_cfg_raw = config.get("model")
    if isinstance(model_cfg_raw, dict):
        model_cfg = model_cfg_raw
        default_spec = str(model_cfg.get("default") or model_cfg.get("model") or "").strip()
        provider = str(model_cfg.get("provider") or "").strip().lower()
        base_url = str(model_cfg.get("base_url") or "").strip()
    elif isinstance(model_cfg_raw, str):
        model_cfg = {}
        default_spec = model_cfg_raw.strip()
        provider = ""
        base_url = ""
    else:
        model_cfg = {}
        default_spec = ""
        provider = ""
        base_url = ""

    parsed_provider = ""
    model_name = default_spec
    if default_spec and "/" in default_spec:
        try:
            parsed_provider, parsed_model = _split_model_spec(default_spec)
            if not provider:
                provider = parsed_provider
            if not provider or provider == parsed_provider or parsed_provider.startswith("custom:"):
                model_name = parsed_model
        except Exception as exc:
            _m.logger.debug("multitenancy: cannot split profile model spec %r (%s)", default_spec, exc)

    entry = _m._matching_custom_provider_entry(config, provider)
    if entry is not None:
        if not base_url:
            base_url = str(entry.get("base_url") or entry.get("url") or entry.get("api") or "").strip()
        if not model_name:
            model_name = str(entry.get("model") or "").strip()
        api_mode = str(entry.get("api_mode") or "").strip()
    else:
        api_mode = str(model_cfg.get("api_mode") or "").strip()

    api_key = ""
    try:
        api_key = str(_resolve_custom_provider_api_key(config, provider) or "").strip()
    except Exception as exc:
        _m.logger.debug("multitenancy: cannot resolve custom provider api key for image prep (%s)", exc)

    provider_l = provider.lower()
    if not (provider_l == "custom" or provider_l.startswith("custom:")):
        return None
    if not (provider and model_name):
        return None
    return {
        "provider": provider,
        "model": model_name,
        "base_url": base_url,
        "api_key": api_key,
        "api_mode": api_mode,
    }


def _install_auxiliary_main_runtime_patch(runtime: Optional[dict[str, str]]) -> tuple[Optional[Any], dict[str, tuple[bool, Any]]]:
    """Patch hermes-agent auxiliary runtime readers under the env lock."""
    if not runtime:
        return None, {}
    try:
        auxiliary_client = importlib.import_module("agent.auxiliary_client")
    except Exception as exc:
        _m.logger.debug("multitenancy: agent.auxiliary_client unavailable for image prep runtime (%s)", exc)
        return None, {}

    attrs = tuple(_m._AUX_MAIN_RUNTIME_FIELDS) + ("_read_main_provider", "_read_main_model")
    saved = {name: (hasattr(auxiliary_client, name), getattr(auxiliary_client, name, None)) for name in attrs}
    provider = runtime.get("provider", "")
    model = runtime.get("model", "")
    base_url = runtime.get("base_url", "")
    api_key = runtime.get("api_key", "")
    api_mode = runtime.get("api_mode", "")

    set_runtime_main = getattr(auxiliary_client, "set_runtime_main", None)
    if callable(set_runtime_main) and api_key:
        try:
            set_runtime_main(provider, model, base_url=base_url, api_key=api_key, api_mode=api_mode)
        except Exception as exc:
            _m.logger.debug("multitenancy: agent.auxiliary_client.set_runtime_main failed (%s)", exc)

    setattr(auxiliary_client, "_read_main_provider", lambda: provider)
    setattr(auxiliary_client, "_read_main_model", lambda: model)
    return auxiliary_client, saved


def _install_vision_model_override(provider: str) -> tuple[Optional[Any], dict[str, tuple[bool, Any]]]:
    """Override the auto-detected vision model for ``provider`` when the core's
    hardcoded choice is unavailable on our account.

    The vision auto-detect resolves the multimodal model via
    ``auxiliary_client._PROVIDER_VISION_MODELS[provider]``; for zai it picks
    ``glm-5v-turbo``, which our z.ai plan does NOT include (HTTP 429 code 1311),
    so image vision fails even though TEXT works. ``glm-4.6v`` is on the plan and
    is multimodal. This runs independently of the main-runtime patch (which is a
    no-op when no complete custom runtime can be built, e.g. plain zai profiles),
    so the override applies for those profiles too. Restored after the prep call.
    """
    override_model = _m._VISION_MODEL_OVERRIDE.get(str(provider or "").strip().lower())
    if not override_model:
        return None, {}
    try:
        auxiliary_client = importlib.import_module("agent.auxiliary_client")
    except Exception as exc:
        _m.logger.debug("multitenancy: agent.auxiliary_client unavailable for vision override (%s)", exc)
        return None, {}
    current_map = getattr(auxiliary_client, "_PROVIDER_VISION_MODELS", None)
    key = str(provider or "").strip().lower()
    if not isinstance(current_map, dict) or current_map.get(key) == override_model:
        return None, {}
    saved = {"_PROVIDER_VISION_MODELS": (True, current_map)}
    patched_map = dict(current_map)
    patched_map[key] = override_model
    setattr(auxiliary_client, "_PROVIDER_VISION_MODELS", patched_map)
    return auxiliary_client, saved


def _restore_auxiliary_main_runtime_patch(auxiliary_client: Optional[Any], saved: dict[str, tuple[bool, Any]]) -> None:
    if auxiliary_client is None:
        return
    for name, (had_attr, value) in saved.items():
        try:
            if had_attr:
                setattr(auxiliary_client, name, value)
            else:
                delattr(auxiliary_client, name)
        except Exception as exc:
            _m.logger.debug("multitenancy: failed to restore auxiliary runtime attr %s (%s)", name, exc)


@asynccontextmanager
async def _profile_image_prep_runtime(profile_home: Optional[Path]):
    """Temporarily scope Hermes' private inbound media preprocessing to a profile."""
    if profile_home is None:
        yield
        return

    from ..runtime import HERMES_HOME_ENV, _PROFILE_HOME_VAR, _get_env_lock

    profile_home = Path(profile_home)
    runtime = _m._profile_main_runtime_for_image_prep(profile_home)
    token = _PROFILE_HOME_VAR.set(profile_home)
    try:
        async with _get_env_lock():
            had_home = HERMES_HOME_ENV in os.environ
            old_home = os.environ.get(HERMES_HOME_ENV)
            auxiliary_client = None
            saved_aux: dict[str, tuple[bool, Any]] = {}
            vis_client = None
            saved_vis: dict[str, tuple[bool, Any]] = {}
            try:
                os.environ[HERMES_HOME_ENV] = str(profile_home)
                auxiliary_client, saved_aux = _m._install_auxiliary_main_runtime_patch(runtime)
                # Apply the vision-model override independently of the main-runtime
                # patch (which no-ops for plain provider profiles like zai). Read
                # the profile's provider from its merged config.
                try:
                    from ..agent_real import _load_profile_config
                    _prof_cfg = _load_profile_config(profile_home)
                    _model_cfg = _prof_cfg.get("model")
                    _prof_provider = ""
                    if isinstance(_model_cfg, dict):
                        _prof_provider = str(_model_cfg.get("provider") or "").strip()
                        if not _prof_provider:
                            # Common shape: model.default == "zai/glm-5.1" with no
                            # separate provider field. _load_profile_config already
                            # ran _normalize_model_spec_inplace, so the prefix is
                            # the provider — derive it so default-only profiles
                            # aren't missed.
                            _default = str(_model_cfg.get("default") or "").strip()
                            if "/" in _default:
                                _prof_provider = _default.split("/", 1)[0].strip()
                    elif isinstance(_model_cfg, str) and "/" in _model_cfg:
                        # Legacy shape: model is a bare "provider/model" string.
                        _prof_provider = _model_cfg.split("/", 1)[0].strip()
                except Exception as exc:
                    _prof_provider = ""
                    _m.logger.debug("multitenancy: vision-override provider read failed (%s)", exc)
                if _prof_provider:
                    vis_client, saved_vis = _m._install_vision_model_override(_prof_provider)
                yield
            finally:
                _m._restore_auxiliary_main_runtime_patch(vis_client, saved_vis)
                _m._restore_auxiliary_main_runtime_patch(auxiliary_client, saved_aux)
                if had_home:
                    os.environ[HERMES_HOME_ENV] = old_home or ""
                else:
                    os.environ.pop(HERMES_HOME_ENV, None)
    finally:
        _PROFILE_HOME_VAR.reset(token)


def _materialize_inbound_media_for_profile(event: Any, profile_home: Path) -> None:
    """Copy gateway-cached inbound media into the routed profile boundary."""
    media_urls = getattr(event, "media_urls", None) or []
    if not media_urls:
        return
    media_types = getattr(event, "media_types", None) or []
    root = profile_home.resolve(strict=False)
    rewritten: list[str] = []
    replacements: dict[str, str] = {}
    changed = False

    for raw_path, raw_mtype in zip_longest(media_urls, media_types, fillvalue=""):
        raw = str(raw_path or "")
        if not raw or re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*://", raw):
            rewritten.append(raw)
            continue
        source = Path(raw).expanduser().resolve(strict=False)
        if not (source.exists() and source.is_file()):
            rewritten.append(raw)
            continue
        if source == root or root in source.parents:
            rewritten.append(str(source))
            continue

        suffix = source.suffix.lower()
        media_type = str(raw_mtype or "").lower()
        target_dir = root / "cache" / "images" if (
            media_type.startswith("image") or suffix in _m._IMAGE_FILE_EXTENSIONS
        ) else root / "uploads"
        target = (target_dir / source.name).resolve(strict=False)
        if target.exists() and target.resolve(strict=False) != source:
            digest = hashlib.sha1(str(source).encode("utf-8")).hexdigest()[:10]
            target = (target_dir / f"{source.stem}-{digest}{source.suffix}").resolve(strict=False)
        if not (target == root or root in target.parents):
            rewritten.append(raw)
            continue
        try:
            target_dir.mkdir(parents=True, exist_ok=True)
            if target.resolve(strict=False) != source:
                shutil.copy2(source, target)
            rewritten_path = str(target)
            rewritten.append(rewritten_path)
            replacements[raw] = rewritten_path
            changed = True
            _m.logger.warning(
                "multitenancy: materialized inbound media for profile source=%s target=%s",
                source,
                target,
            )
        except Exception as exc:
            _m.logger.warning(
                "multitenancy: failed to materialize inbound media source=%s profile_home=%s error=%s",
                source,
                root,
                exc,
            )
            rewritten.append(raw)

    if changed:
        setattr(event, "media_urls", rewritten)
        text = getattr(event, "text", None)
        if isinstance(text, str):
            for old, new in replacements.items():
                text = text.replace(old, new)
            setattr(event, "text", text)


async def _enrich_via_hermes_pipeline(
    event: Any,
    gateway: Any,
    *,
    profile_home: Optional[Path] = None,
) -> Optional[str]:
    """Delegate inbound preprocessing to hermes' ``_prepare_inbound_message_text``.

    This is the single call that mainstream uses to:
      - run vision_analyze_tool on attached images
      - run transcribe_audio on voice messages
      - inject text-file content (.txt / .md / .csv / etc.)
      - prepend reply-quoted context
      - attribute multi-user shared sessions

    By calling the same gateway method, the plugin behaves *identically* to
    mainstream for every multimodal input — no re-implementation, no drift.

    Caveat: this depends on a private GatewayRunner method. If hermes-agent
    refactors ``_prepare_inbound_message_text``, swap to local fallbacks
    (``_local_enrich_with_vision_only`` below as a minimal safety net).

    Returns:
        Enriched text string, or None on failure (caller falls back to event.text).
    """
    media_profile_home = profile_home if (getattr(event, "media_urls", None) or []) else None
    async with _m._profile_image_prep_runtime(media_profile_home):
        native_text: Optional[str] = None
        if _m._event_has_image_media(event):
            strategy = os.getenv("HERMES_MULTITENANCY_IMAGE_PREP_STRATEGY", "gateway").strip().lower()
            if strategy in {"blocked", "block", "skip", "disabled", "off"}:
                _m.logger.warning(
                    "multitenancy: image preprocessing blocked by strategy message_id=%s",
                    _m._event_message_id(event) or "",
                )
                return _m._image_prep_unavailable_note(event)

        prep = getattr(gateway, "_prepare_inbound_message_text", None)
        if gateway is None or prep is None or not callable(prep):
            _m.logger.debug("multitenancy: gateway._prepare_inbound_message_text unavailable")
        else:
            source = getattr(event, "source", None)
            if source is not None:
                try:
                    prep_call = prep(event=event, source=source, history=[])
                    if _m._event_has_image_media(event):
                        timeout_s = float(os.getenv("HERMES_MULTITENANCY_IMAGE_PREP_TIMEOUT_S", "30"))
                        native_text = await asyncio.wait_for(prep_call, timeout=max(0.1, timeout_s))
                    else:
                        native_text = await prep_call
                except asyncio.TimeoutError:
                    _m.logger.warning("multitenancy: image preprocessing timed out")
                    native_text = _m._image_prep_unavailable_note(event, reason="timeout")
                except Exception as exc:
                    _m.logger.debug("multitenancy: gateway._prepare_inbound_message_text failed (%s)", exc)

        local_file_text = _m._local_enrich_with_file_content(event, existing_text=native_text or "")
        if local_file_text:
            return _m._append_enrichment(native_text or getattr(event, "text", "") or "", local_file_text)

        if native_text:
            return native_text
        return await _m._local_enrich_with_vision_only(event)


async def _call_enrich_via_hermes_pipeline(
    event: Any,
    gateway: Any,
    *,
    profile_home: Optional[Path],
) -> Optional[str]:
    """Call enrichment with profile context while preserving monkeypatch compatibility."""
    try:
        params = inspect.signature(_m._enrich_via_hermes_pipeline).parameters
    except (TypeError, ValueError):
        params = {}
    accepts_profile_home = "profile_home" in params or any(
        param.kind == inspect.Parameter.VAR_KEYWORD
        for param in params.values()
    )
    if accepts_profile_home:
        return await _m._enrich_via_hermes_pipeline(event, gateway, profile_home=profile_home)
    return await _m._enrich_via_hermes_pipeline(event, gateway)


def _append_enrichment(base: str, enrichment: str) -> str:
    base = str(base or "").strip()
    enrichment = str(enrichment or "").strip()
    if not enrichment:
        return base
    if not base:
        return enrichment
    return f"{enrichment}\n{base}"


def _local_enrich_with_file_content(event: Any, *, existing_text: str = "") -> Optional[str]:
    """Plugin-owned fallback for Feishu document attachments Hermes does not inline.

    Hermes core remains the first choice for multimodal preprocessing. This
    fallback only covers plain/tabular files already cached as local event paths,
    keeping compatibility in multitenancy instead of patching Hermes-agent.
    """
    media_urls = getattr(event, "media_urls", None) or []
    media_types = getattr(event, "media_types", None) or []
    if not media_urls:
        return None
    existing = str(existing_text or "")
    parts: list[str] = []
    for raw_path, raw_mtype in zip_longest(media_urls, media_types, fillvalue=""):
        path = Path(str(raw_path))
        if not path.is_file():
            continue
        name = path.name
        suffix = path.suffix.lower()
        if suffix not in {".xlsx", ".docx", ".pdf"} and f"[Content of {name}]" in existing:
            continue
        try:
            if path.stat().st_size > _m._MAX_LOCAL_ENRICH_FILE_BYTES:
                _m.logger.debug("multitenancy: local file enrichment skipped oversized file %s", path)
                continue
        except OSError:
            continue
        media_type = str(raw_mtype or "").lower()
        try:
            if suffix == ".xlsx":
                content = _m._extract_xlsx_text(path)
            elif suffix == ".docx":
                content = _m._extract_docx_text(path)
            elif suffix == ".pdf":
                content = _m._extract_pdf_text(path)
            elif media_type.startswith("text/") or suffix in _m._TEXT_FILE_EXTENSIONS:
                with path.open("rb") as handle:
                    content = handle.read(_m._MAX_LOCAL_TEXT_PREVIEW_BYTES).decode("utf-8", errors="replace")
            else:
                continue
        except Exception as exc:
            _m.logger.debug("multitenancy: local file enrichment failed for %s: %s", path, exc)
            continue
        content = content.strip()
        if not content or content in existing:
            continue
        header = f"[Content of {name}]"
        if suffix in {".xlsx", ".docx", ".pdf"} and header in existing:
            header = f"[Content of {name} - multitenancy {suffix.lstrip('.')} fallback]"
        parts.append(f"{header}:\n{content}")
    if not parts:
        return None
    return "\n\n".join(parts)


def _extract_xlsx_text(path: Path, *, max_sheets: int = 3, max_rows: int = 50, max_cells: int = 20) -> str:
    """Extract a small text preview from an XLSX file using only stdlib."""
    ns = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    rel_ns = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
    with zipfile.ZipFile(path) as zf:
        shared_strings: list[str] = []
        if "xl/sharedStrings.xml" in zf.namelist():
            root = ET.fromstring(_m._read_zip_member_limited(zf, "xl/sharedStrings.xml"))
            for si in root.findall("main:si", ns):
                texts = [node.text or "" for node in si.findall(".//main:t", ns)]
                shared_strings.append("".join(texts))

        sheet_names: dict[str, str] = {}
        if "xl/workbook.xml" in zf.namelist():
            workbook = ET.fromstring(_m._read_zip_member_limited(zf, "xl/workbook.xml"))
            for idx, sheet in enumerate(workbook.findall(".//main:sheet", ns), start=1):
                rel_id = sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                sheet_names[rel_id or f"rId{idx}"] = sheet.attrib.get("name") or f"sheet{idx}"

        rel_targets: list[tuple[str, str]] = []
        if "xl/_rels/workbook.xml.rels" in zf.namelist():
            rels = ET.fromstring(_m._read_zip_member_limited(zf, "xl/_rels/workbook.xml.rels"))
            for rel in rels.findall("rel:Relationship", rel_ns):
                target = rel.attrib.get("Target") or ""
                if target.startswith("worksheets/"):
                    rel_targets.append((rel.attrib.get("Id") or "", "xl/" + target))
        if not rel_targets:
            rel_targets = [
                (f"rId{idx}", name)
                for idx, name in enumerate(sorted(
                    n for n in zf.namelist() if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")
                ), start=1)
            ]

        sections: list[str] = []
        for rel_id, sheet_path in rel_targets[:max_sheets]:
            if sheet_path not in zf.namelist():
                continue
            root = ET.fromstring(_m._read_zip_member_limited(zf, sheet_path))
            rows: list[str] = []
            for row in root.findall(".//main:sheetData/main:row", ns)[:max_rows]:
                cells: list[str] = []
                for cell in row.findall("main:c", ns)[:max_cells]:
                    value = cell.find("main:v", ns)
                    raw = value.text if value is not None else ""
                    cell_type = cell.attrib.get("t")
                    if cell_type == "s" and raw.isdigit():
                        idx = int(raw)
                        raw = shared_strings[idx] if idx < len(shared_strings) else raw
                    elif cell_type == "inlineStr":
                        texts = [node.text or "" for node in cell.findall(".//main:t", ns)]
                        raw = "".join(texts)
                    cells.append(raw or "")
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                sections.append(f"[{sheet_names.get(rel_id, rel_id or sheet_path)}]\n" + "\n".join(rows))
        return "\n\n".join(sections)


def _read_zip_member_limited(zf: zipfile.ZipFile, name: str) -> bytes:
    info = zf.getinfo(name)
    if info.file_size > _m._MAX_XLSX_XML_BYTES:
        raise ValueError(f"xlsx member too large: {name}")
    return zf.read(name)


def _extract_docx_text(path: Path, *, max_paragraphs: int = 80) -> str:
    """Extract a small text preview from a DOCX file using only stdlib."""
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as zf:
        if "word/document.xml" not in zf.namelist():
            return ""
        root = ET.fromstring(_m._read_zip_member_limited(zf, "word/document.xml"))
        paragraphs: list[str] = []
        for para in root.findall(".//w:p", ns)[:max_paragraphs]:
            texts = [node.text or "" for node in para.findall(".//w:t", ns)]
            text = "".join(texts).strip()
            if text:
                paragraphs.append(text)
        return "\n".join(paragraphs)


def _extract_pdf_text(path: Path, *, max_bytes: int = _m._MAX_LOCAL_TEXT_PREVIEW_BYTES) -> str:
    """Best-effort PDF text preview for small generated/text PDFs.

    This intentionally stays lightweight. It covers the real UAT fixtures
    generated by ReportLab (ASCII85Decode + FlateDecode text streams) and
    simple uncompressed text streams; scanned PDFs still require OCR/vision.
    """
    raw = path.read_bytes()[: _m._MAX_LOCAL_ENRICH_FILE_BYTES]
    chunks: list[str] = []
    for match in _m._PDF_STREAM_RE.finditer(raw):
        stream = match.group("body").strip()
        decoded = _m._decode_pdf_stream(stream)
        if not decoded:
            continue
        chunks.extend(_m._extract_pdf_literal_strings(decoded))
        if sum(len(item) for item in chunks) >= max_bytes:
            break
    if not chunks:
        chunks.extend(_m._extract_pdf_literal_strings(raw))
    text = "\n".join(item for item in chunks if item.strip())
    return text[:max_bytes]


def _decode_pdf_stream(stream: bytes) -> bytes:
    candidates = [stream]
    try:
        candidates.append(base64.a85decode(stream, adobe=True))
    except Exception:
        pass
    for candidate in list(candidates):
        try:
            decoded = zlib.decompress(candidate)
        except Exception:
            continue
        candidates.append(decoded)
    return candidates[-1] if candidates else b""


def _extract_pdf_literal_strings(raw: bytes) -> list[str]:
    values: list[str] = []
    for match in _m._PDF_TEXT_STRING_RE.finditer(raw):
        token = match.group(0)[1:-1]
        token = (
            token.replace(rb"\(", b"(")
            .replace(rb"\)", b")")
            .replace(rb"\\", b"\\")
            .replace(rb"\n", b"\n")
            .replace(rb"\r", b"\r")
            .replace(rb"\t", b"\t")
        )
        text = token.decode("utf-8", errors="replace").strip()
        if text:
            values.append(text)
    return values


async def _local_enrich_with_vision_only(event: Any) -> Optional[str]:
    """Local fallback if hermes' ``_prepare_inbound_message_text`` is unavailable.

    Only handles images (the most common multimodal input). Audio / files /
    reply context degrade gracefully — the model will see ``event.text`` only.
    """
    media_urls = getattr(event, "media_urls", None) or []
    media_types = getattr(event, "media_types", None) or []
    if not media_urls:
        return None
    try:
        from tools.vision_tools import vision_analyze_tool  # type: ignore
    except ImportError:
        return None
    import json as _json
    descriptions: list[str] = []
    for path, mtype in zip(media_urls, media_types or [""] * len(media_urls)):
        if mtype and not mtype.startswith("image"):
            continue
        try:
            result_json = await vision_analyze_tool(
                image_url=path,
                user_prompt="Describe this image in thorough detail.",
            )
            result = _json.loads(result_json) if isinstance(result_json, str) else result_json
            if isinstance(result, dict) and result.get("success"):
                descriptions.append(f"[Image: {result.get('analysis', '')}]")
            else:
                error = ""
                if isinstance(result, dict):
                    error = str(result.get("error") or result.get("analysis") or "").strip()
                descriptions.append(_m._image_analysis_unavailable_note(path, error))
        except Exception as exc:
            _m.logger.debug("multitenancy: local vision fallback error on %s: %s", path, exc)
            descriptions.append(_m._image_analysis_unavailable_note(path, str(exc)))
    if not descriptions:
        return None
    base = getattr(event, "text", "") or ""
    return "\n".join(descriptions) + ("\n" + base if base else "")


def _image_analysis_unavailable_note(path: Any, error: str = "") -> str:
    """Return recoverable context when an image is present but vision fails."""
    reason = re.sub(r"\s+", " ", str(error or "")).strip()
    if len(reason) > 300:
        reason = reason[:297] + "..."
    suffix = f" Reason: {reason}" if reason else ""
    return (
        "[Image analysis unavailable: the image is attached at "
        f"{path}, but automatic vision analysis failed.{suffix} "
        "If the user asks about the screenshot, explain that a vision-capable "
        "model or permission is required, or ask the user to describe the image.]"
    )
