"""Pre-gateway-dispatch hook callback (sync) + async dispatch entry point.

Wires together: SQLite RoutingTable (production) + in-memory _SPIKE_ROUTING
(fallback) + LRU RuntimePool (cached profile runtimes) + Hermes-derived slash
command dispatch.
"""
from __future__ import annotations

import asyncio
import base64
import copy
import hashlib
import inspect
import json
import logging
import os
import re
import shutil
import time
import zlib
import zipfile
from contextlib import asynccontextmanager
from itertools import zip_longest
from pathlib import Path
from typing import Any, Optional
from xml.etree import ElementTree as ET

logger = logging.getLogger(__name__)

_SKILL_SLASH_ALIASES = {
    "hades": "kep-hades-cli",
}

try:  # Shared Hermes Feishu card/typewriter transport.
    from gateway.stream_consumer import GatewayStreamConsumer, StreamConsumerConfig  # type: ignore
except Exception:  # pragma: no cover - allows plugin unit tests without gateway on path
    GatewayStreamConsumer = None  # type: ignore
    StreamConsumerConfig = None  # type: ignore

# Per-context in-flight dispatch tasks — used by /stop and replacement turns to
# cancel only the matching profile/chat task while preserving enough history to
# resume.
_user_inflight_tasks: dict[tuple[str, str, str], asyncio.Task] = {}
_user_inflight_history_keys: dict[tuple[str, str, str], tuple[str, str]] = {}
_suppress_interruption_marker_tasks: set[asyncio.Task] = set()
_synthetic_session_guards: dict[str, Any] = {}

# Group-chat profile state. Populated by the bot_added hook (Layer 4, runs
# on the Lark SDK thread) and consumed by the auto-provision path (asyncio
# loop thread). Persistent state lives in the SQLite routing table; this
# only covers the window between bot_added and the first @mention.
#
# Bounded + TTL'd + lock-guarded: an attacker spamming bot add/remove across
# throwaway chats would otherwise grow this dict without bound, and the
# SDK-thread writer can race the loop-thread reader/popper. OrderedDict +
# lock keeps eviction and the register/pop pair consistent across threads.
import threading as _threading
from collections import OrderedDict as _OrderedDict

_CHAT_INVITER_CACHE_MAX = 512
_CHAT_INVITER_CACHE_TTL_S = 3600  # bot_added → first @; generous for restarts
_chat_inviter_cache: "_OrderedDict[str, dict[str, Any]]" = _OrderedDict()
_chat_inviter_cache_lock = _threading.Lock()
_GROUP_PROFILE_PREFIX = "feishu_group_"
_GROUP_CHAT_TYPES: frozenset[str] = frozenset({"group", "topic"})
_LARK_CLI_PROFILE_TOOLSETS = [
    "lark-cli",
]
_LARK_CLI_SOUL_GUIDANCE = "\n".join(
    [
        "Feishu/Lark capability rules:",
        "- 飞书/Lark 的读写、导出和长尾 OpenAPI 能力必须通过 `lark_cli` 工具完成。",
        "- 当用户明确要求“调用 lark_cli / 使用 lark-cli / 必须真实调用工具”时，即使你从历史上下文知道答案，也必须重新调用 `lark_cli`，不得只凭记忆回答。",
        "- 不要通过 `terminal`、`code_execution`、`npx`、`which lark-cli` 或 shell 直接运行 lark-cli/lark-mcp 来绕开 profile runtime。",
        "- 如果 `lark_cli` 工具不可见或不可用，直接说明当前 profile 未暴露 lark-cli 能力；不要自行安装、探测或模拟结果。",
        "- 如果需要调用飞书能力，必须等待真实工具结果；不要编造 message_id、文档链接或调用结果。",
        "- 如果 `lark_cli` 返回 credential unavailable、validation、unsupported 或权限错误，立即停止重试并如实说明需要授权/权限/命令支持。",
        "- `lark_cli` 的 identity 使用 auto，profile runtime 会按个人/群聊边界选择 user 或 bot。",
        "- `lark_cli` 工具结果里的 `identity` 字段是身份事实；当 identity=user 时，不得说资源由 bot/应用身份创建。",
        "Feishu file output rules:",
        "- 当用户要求生成文件、图片、报表、PDF、docx、xlsx、csv、json、markdown 并发回时，不要要求用户提供本机路径。",
        "- 直接在回复中输出一个 ```hermes-artifact-json fenced block，字段使用 filename、format、marker/content/data/rows/title；不要使用宿主机绝对路径。",
        "- filename 只写普通文件名，例如 report.md、summary.pdf、chart.png；Hermes 会自动保存到当前 profile 的 Downloads 并通过飞书发送；markdown 源文件必须自动交付给用户。",
        "- 图片如果要作为可下载原文件发送，在 artifact JSON 中设置 as_document=true。",
        "- artifact JSON 是内部交付协议；除必要的测试标记和简短说明外，不要把本机路径、/workspace 路径或 MEDIA 指令解释给用户。",
    ]
)

_GROUP_EXTERNAL_TOOL_SOUL_GUIDANCE = "\n".join(
    [
        "External tool safety rules:",
        "- 当用户要求天气、网页查询或其它外部网络查询时，优先使用只读工具或只读命令。",
        "- 如果必须用 terminal 访问网络，URL 必须写完整 scheme（例如 `https://example.com`），并使用 short timeout，避免把危险命令推给用户审批。",
        "- 不要使用 schemeless URL（例如 `example.com/path`）或下载后执行的命令。",
    ]
)

# Multi-turn session memory: maps (profile_name, user_key) → list of messages.
# user_key is the canonical sender chosen by routing. Alternate IDs are lookup
# helpers only; using them as the memory key can merge distinct users that share
# a stale or tenant-global alias.
_SESSION_HISTORY_MAX = 20  # keep last N messages (user + assistant alternating)
_DEDUPE_MESSAGE_TTL_SECONDS = 24 * 60 * 60
_DEDUPE_CONTENT_TTL_SECONDS = 2 * 60 * 60
_DEDUPE_CONTENT_MIN_CHARS = 40
_session_history: dict[tuple[str, str], list[dict]] = {}
# Tracks which (profile, user_key) pairs we've lazy-loaded from SessionStore
# so we only hit SQLite once per pair per process lifetime.
_session_loaded: set[tuple[str, str]] = set()
_pending_approval_requests: dict[str, list[dict]] = {}
_RECENT_PROFILE_FILE_CONTEXT_MAX = 5
_recent_profile_files_by_chat: dict[tuple[str, str], list[str]] = {}
_RECENT_FILE_CONTEXT_TRIGGER_RE = re.compile(
    r"(这个文件|该文件|这个文档|该文档|刚才.*文件|上面.*文件|源文件|markdown|Markdown|\.md\b|转成飞书云文档|转云文档)"
)


def _history_key(profile_name: str, sender: str, sender_alt: Optional[str]) -> tuple[str, str]:
    """Return the per-(profile, user) key used to look up conversation history."""
    return (profile_name, _tenant_user_key(sender, sender_alt))


def _inflight_key(
    profile_name: Optional[str],
    sender: str,
    sender_alt: Optional[str],
    chat_id: str,
) -> tuple[str, str, str]:
    """Return the profile/chat/user scoped key for replace, /stop, and /status."""
    profile_key = str(profile_name or "").strip() or "_unrouted"
    chat_key = str(chat_id or "").strip() or "unknown"
    return (profile_key, chat_key, _tenant_user_key(sender, sender_alt))


def _tenant_user_key(sender: str, sender_alt: Optional[str]) -> str:
    """Return the canonical per-user key for memory/session scoping."""
    sender_key = str(sender or "").strip()
    if sender_key and sender_key != "unknown":
        return sender_key
    alt_key = str(sender_alt or "").strip()
    return alt_key or sender_key or "unknown"


def _trim_history(history: list[dict]) -> list[dict]:
    """Keep at most ``_SESSION_HISTORY_MAX`` most recent messages."""
    if len(history) <= _SESSION_HISTORY_MAX:
        return history
    return history[-_SESSION_HISTORY_MAX:]


_WRONG_LARK_CLI_BOT_IDENTITY_NOTE_RE = re.compile(
    r"(?im)^[^\n]*(?:(?:由|以|当前)\s*(?:bot|应用)\s*身份(?:创建|下)?)[^\n]*(?:\n|$)"
)


def _strip_wrong_lark_cli_bot_identity_note(text: str) -> str:
    """Remove stale hallucinated bot-identity notes from Feishu UAT history."""
    cleaned = _WRONG_LARK_CLI_BOT_IDENTITY_NOTE_RE.sub("", str(text or ""))
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _sanitize_history_messages(history: list[dict]) -> list[dict]:
    sanitized: list[dict] = []
    for item in history:
        if not isinstance(item, dict):
            continue
        copied = dict(item)
        if copied.get("role") == "assistant":
            copied["content"] = _strip_wrong_lark_cli_bot_identity_note(str(copied.get("content") or ""))
        sanitized.append(copied)
    return sanitized


def _load_history(key: tuple[str, str]) -> list[dict]:
    """Get history for ``key`` — first call hydrates from SessionStore, subsequent calls hit cache."""
    if key in _session_loaded:
        return _sanitize_history_messages(list(_session_history.get(key, [])))
    store = _get_session_store()
    if store is not None:
        try:
            persisted = store.load_recent(key[0], key[1], _SESSION_HISTORY_MAX)
        except Exception as exc:
            logger.debug("multitenancy: SessionStore.load_recent failed (%s)", exc)
            persisted = []
        if persisted:
            _session_history[key] = _sanitize_history_messages(persisted)
    _session_loaded.add(key)
    return _sanitize_history_messages(list(_session_history.get(key, [])))


def _persist_history_message(key: tuple[str, str], message: dict) -> None:
    """Append one user/assistant message to cache and SessionStore."""
    role = str(message.get("role") or "").strip()
    content = str(message.get("content") or "")
    if role not in {"user", "assistant"} or not content:
        return
    if role == "assistant":
        content = _strip_wrong_lark_cli_bot_identity_note(content)
    new_history = _session_history.get(key, []) + [{"role": role, "content": content}]
    _session_history[key] = _trim_history(new_history)
    store = _get_session_store()
    if store is None:
        return
    try:
        store.append(key[0], key[1], role, content)
    except Exception as exc:
        logger.debug("multitenancy: SessionStore.append(%s) failed (%s)", role, exc)


def _persist_turn(key: tuple[str, str], user_msg: dict, assistant_text: str) -> None:
    """Append a (user, assistant) turn to both in-memory cache and SessionStore."""
    _persist_history_message(key, user_msg)
    _persist_history_message(key, {"role": "assistant", "content": assistant_text})


def _persist_user_message(key: tuple[str, str], user_msg: dict) -> None:
    """Persist the user request before execution so interruption can resume."""
    _persist_history_message(key, user_msg)


def _persist_assistant_message(key: tuple[str, str], assistant_text: str) -> None:
    """Persist an assistant response or interruption marker."""
    _persist_history_message(key, {"role": "assistant", "content": assistant_text})


_INTERRUPTED_TASK_MESSAGE = (
    "上一个任务在完成前被中断或取消；如果用户要求继续，请根据上一条用户请求继续推进，不要丢失上下文。"
)
_FAILED_TASK_MESSAGE = (
    "上一个任务在完成前执行失败或中断；如果用户要求继续，请根据上一条用户请求继续推进，不要丢失上下文。"
)


def _persist_interruption_marker(key: tuple[str, str]) -> None:
    """Persist the canonical interruption marker, avoiding duplicate adjacent markers."""
    history = _session_history.get(key, [])
    if history and history[-1].get("role") == "assistant" and history[-1].get("content") == _INTERRUPTED_TASK_MESSAGE:
        return
    _persist_assistant_message(key, _INTERRUPTED_TASK_MESSAGE)


def _persist_failure_marker(key: tuple[str, str]) -> None:
    """Persist a resumable failure marker without exposing exception details."""
    history = _session_history.get(key, [])
    if history and history[-1].get("role") == "assistant" and history[-1].get("content") == _FAILED_TASK_MESSAGE:
        return
    _persist_assistant_message(key, _FAILED_TASK_MESSAGE)


def _cancel_inflight_task(key: tuple[str, str, str], *, preserve_resume_marker: bool) -> Optional[asyncio.Task]:
    """Cancel one active dispatch and optionally write a resumable marker first."""
    task = _user_inflight_tasks.pop(key, None)
    hist_key = _user_inflight_history_keys.pop(key, None)
    if task is None or task.done():
        return task
    if preserve_resume_marker and hist_key is not None:
        _persist_interruption_marker(hist_key)
    _suppress_interruption_marker_tasks.add(task)
    task.cancel()
    return task


def _clear_history(key: tuple[str, str]) -> bool:
    """Drop a user's history from cache + store. Returns True if anything was cleared."""
    had_cache = _session_history.pop(key, None) is not None
    _session_loaded.discard(key)
    store = _get_session_store()
    if store is None:
        return had_cache
    try:
        rows = store.clear(key[0], key[1])
    except Exception as exc:
        logger.debug("multitenancy: SessionStore.clear failed (%s)", exc)
        rows = 0
    return had_cache or rows > 0


def _dedupe_env_int(name: str, default: int) -> int:
    try:
        return max(0, int(os.environ.get(name, str(default)).strip()))
    except Exception:
        return default


def _event_message_id(event: Any) -> Optional[str]:
    source = getattr(event, "source", None)
    candidates = [
        getattr(event, "message_id", None),
        getattr(source, "message_id", None) if source is not None else None,
    ]
    raw_event = getattr(event, "raw_event", None)
    if isinstance(raw_event, dict):
        message = (raw_event.get("event") or {}).get("message") or {}
        candidates.append(message.get("message_id"))
    for value in candidates:
        text = str(value or "").strip()
        if text:
            return text
    return None


_REACTION_SYNTHETIC_PREFIXES = ("reaction:added:", "reaction:removed:")


def _is_reaction_synthetic_event(event: Any, text: str) -> bool:
    """Return True for Feishu reaction events routed as synthetic text."""
    if not isinstance(text, str):
        return False
    if not text.startswith(_REACTION_SYNTHETIC_PREFIXES):
        return False
    message_type = getattr(event, "message_type", None)
    if message_type is None:
        return True
    return str(getattr(message_type, "name", message_type)).upper() == "TEXT"


def _event_reply_to_message_id(event: Any) -> Optional[str]:
    """Only group/topic chats should use Feishu reply_to.

    In p2p/dm chats Feishu renders bot replies with ``reply_to`` as visible
    topics, which makes a private chat look like it was posted into a group.
    """
    if not _is_group_chat_type(_extract_chat_type(event)):
        return None
    return _event_message_id(event)


def _normalize_dedupe_text(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _run_request_dedupe_record(request: Any) -> Optional[tuple[str, Optional[str], Optional[str], int]]:
    idempotency_key = str(getattr(request, "idempotency_key", "") or "").strip()
    channel = str(getattr(request, "channel", "") or "").strip()
    message_id = str(getattr(request, "message_id", "") or "").strip()
    profile_name = str(getattr(request, "profile_name", "") or "").strip()
    user_key = str(getattr(request, "user_key", "") or "").strip()
    content = str(getattr(request, "content", "") or "")
    if idempotency_key:
        digest = hashlib.sha256(idempotency_key.encode("utf-8")).hexdigest()
        ttl = _dedupe_env_int(
            "HERMES_MULTITENANCY_EVENT_DEDUPE_TTL_SECONDS",
            _DEDUPE_MESSAGE_TTL_SECONDS,
        )
        return (f"idem:{channel}:{profile_name}:{user_key}:{digest}", None, None, ttl)
    if message_id:
        ttl = _dedupe_env_int(
            "HERMES_MULTITENANCY_EVENT_DEDUPE_TTL_SECONDS",
            _DEDUPE_MESSAGE_TTL_SECONDS,
        )
        return (f"msg:{profile_name}:{user_key}:{message_id}", message_id, None, ttl)
    if channel == "webui":
        return None

    normalized = _normalize_dedupe_text(content)
    min_chars = _dedupe_env_int(
        "HERMES_MULTITENANCY_CONTENT_DEDUPE_MIN_CHARS",
        _DEDUPE_CONTENT_MIN_CHARS,
    )
    if len(normalized) < min_chars:
        return None
    digest = hashlib.sha256(normalized.encode("utf-8")).hexdigest()
    ttl = _dedupe_env_int(
        "HERMES_MULTITENANCY_CONTENT_DEDUPE_TTL_SECONDS",
        _DEDUPE_CONTENT_TTL_SECONDS,
    )
    return (f"content:{profile_name}:{user_key}:{digest}", None, digest, ttl)


def _mark_run_request_seen(request: Any) -> bool:
    """Return False when this broker RunRequest was processed recently."""
    store = _get_session_store()
    if store is None:
        return True
    record = _run_request_dedupe_record(request)
    if record is None:
        return True
    event_key, message_id, content_hash, ttl = record
    try:
        return bool(store.mark_event_processed(
            event_key,
            profile_name=str(getattr(request, "profile_name", "") or ""),
            user_key=str(getattr(request, "user_key", "") or ""),
            message_id=message_id,
            content_hash=content_hash,
            ttl_seconds=ttl,
        ))
    except Exception as exc:
        logger.debug("multitenancy: run request dedupe check failed (%s)", exc)
        return True


def _host_tools_require_sandbox() -> bool:
    value = os.environ.get("HERMES_REQUIRE_SANDBOX_FOR_HOST_TOOLS", "").strip().lower()
    return value in {"1", "true", "yes", "on"}


def _router_sandbox_available() -> bool:
    value = os.environ.get("HERMES_USE_SANDBOX", "").strip().lower()
    return value in {"1", "true", "yes", "on"}


def _run_request_for_routed_event(
    *,
    event: Any,
    profile_name: str,
    sender: str,
    sender_alt: Optional[str],
    chat_id: str,
    text: str,
):
    from .run_models import RunRequest

    user_key = _tenant_user_key(sender, sender_alt)
    return RunRequest(
        channel="feishu",
        profile_name=profile_name,
        user_key=user_key,
        content=text,
        chat_id=chat_id,
        message_id=_event_message_id(event),
        credential_subject=user_key,
        requires_host_tools=_host_tools_require_sandbox(),
        metadata={"sender_alt": sender_alt} if sender_alt else {},
    )


def _make_routed_run_broker(*, dispatch_agent: Any = None):
    from .run_broker import RunBroker

    return RunBroker(
        dispatch_agent=dispatch_agent or (lambda _request: ""),
        mark_seen=_mark_run_request_seen,
        sandbox_available=_router_sandbox_available,
    )


def _build_user_message(event: Any, *, text_override: Optional[str] = None) -> dict:
    """Construct the OpenAI-format user message, splicing in reply context if any.

    Reply context: hermes mainstream sets ``event.reply_to_text`` when the
    user is quoting an earlier message. We surface it inline so the model
    knows what's being replied to.

    ``text_override`` lets ``_enrich_with_vision`` rewrite the text before
    the message is built (so reply context still wraps the enriched text).
    """
    text = text_override if text_override is not None else (getattr(event, "text", "") or "")
    reply_to_text = getattr(event, "reply_to_text", None)
    if reply_to_text:
        text = f"(replying to: {reply_to_text})\n{text}"
    return {"role": "user", "content": text}


def _normalize_feishu_open_id(value: Any) -> Optional[str]:
    raw = str(value or "").strip()
    if raw.startswith("user:"):
        raw = raw.split(":", 1)[1]
    return raw if raw.startswith("ou_") else None


def _is_feishu_open_id(value: Any) -> bool:
    return _normalize_feishu_open_id(value) is not None


def _nested_get(value: Any, path: tuple[str, ...]) -> Any:
    current = value
    for key in path:
        if current is None:
            return None
        if isinstance(current, dict):
            current = current.get(key)
        else:
            current = getattr(current, key, None)
    return current


def _current_sender_open_id() -> Optional[str]:
    """Return the adapter-provided Feishu open_id context, if available."""
    try:
        from tools import feishu_oapi_client as feishu_oapi  # type: ignore

        candidate = feishu_oapi.current_sender_open_id.get()
    except Exception:
        return None
    return _normalize_feishu_open_id(candidate)


def _resolve_sender_for_routing(event: Any, *, fallback: str = "unknown") -> str:
    """Pick the stable Feishu user key used by the multitenancy router.

    Feishu SDK events can expose ``source.user_id`` as a short tenant-local ID
    such as ``g41a5b5g``. The UAT files and explicit routes are keyed by the
    real app-scoped open_id (``ou_*``), so prefer that when the adapter or raw
    event makes it available.
    """
    source = getattr(event, "source", None)
    direct_candidates = (
        _current_sender_open_id(),
        getattr(event, "sender_open_id", None),
        getattr(source, "open_id", None) if source is not None else None,
        getattr(source, "user_id", None) if source is not None else None,
        getattr(source, "user_id_alt", None) if source is not None else None,
    )
    for candidate in direct_candidates:
        normalized = _normalize_feishu_open_id(candidate)
        if normalized:
            return normalized

    raw_candidates = (
        getattr(event, "raw", None),
        getattr(event, "raw_event", None),
        getattr(event, "event", None),
    )
    paths = (
        ("sender", "sender_id", "open_id"),
        ("event", "sender", "sender_id", "open_id"),
        ("event", "message", "sender", "sender_id", "open_id"),
        ("message", "sender", "sender_id", "open_id"),
        ("sender_id", "open_id"),
    )
    for raw in raw_candidates:
        for path in paths:
            candidate = _nested_get(raw, path)
            normalized = _normalize_feishu_open_id(candidate)
            if normalized:
                return normalized

    return fallback


def _event_with_text(event: Any, text: str) -> Any:
    """Return an event-shaped object whose ``text`` matches the agent prompt."""
    if text == (getattr(event, "text", "") or ""):
        return event
    cloned = copy.copy(event)
    setattr(cloned, "text", text)
    return cloned


def _clean_stream_display_text(text: str, profile_home: Optional[Path] = None) -> str:
    """Hide native media-delivery directives from visible streaming text."""
    try:
        from gateway.stream_consumer import GatewayStreamConsumer  # type: ignore

        cleaned = GatewayStreamConsumer._clean_for_display(text)
        cleaned = _ARTIFACT_JSON_RE.sub("", cleaned)
    except Exception:
        cleaned = str(text or "").replace("[[audio_as_voice]]", "")
        cleaned = _ARTIFACT_JSON_RE.sub("", cleaned)
        cleaned = re.sub(r'''[`"']?MEDIA:\s*\S+[`"']?''', "", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        cleaned = cleaned.rstrip()
    cleaned = cleaned.replace("[[as_document]]", "")
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).rstrip()
    cleaned = _linkify_feishu_document_ids(cleaned)
    if profile_home is not None:
        cleaned = _strip_plain_profile_file_paths_for_display(cleaned, profile_home)
        if not Path(profile_home).name.startswith(_GROUP_PROFILE_PREFIX):
            cleaned = _strip_wrong_lark_cli_bot_identity_note(cleaned)
    return cleaned


def _clean_stream_delta_text(text: str, profile_home: Optional[Path] = None) -> str:
    """Clean one streamed content delta without dropping token boundary spaces."""
    raw = str(text or "")
    cleaned = _clean_stream_display_text(raw, profile_home)
    if cleaned and raw[:1].isspace() and not cleaned[:1].isspace():
        leading = re.match(r"^\s+", raw)
        if leading:
            cleaned = leading.group(0) + cleaned
    return cleaned


def _sanitize_tool_event_payload(payload: Any, profile_home: Optional[Path]) -> dict[str, Any]:
    """Prepare tool progress metadata for user-visible cards."""
    data = payload if isinstance(payload, dict) else {"name": str(payload or "tool")}
    result: dict[str, Any] = {}
    for key, value in data.items():
        if key in {"preview", "args", "name", "tool_name", "duration", "is_error"}:
            result[key] = _sanitize_tool_event_value(value, profile_home, key=str(key))
    return result


def _sanitize_tool_event_value(value: Any, profile_home: Optional[Path], *, key: str = "") -> Any:
    if re.search(r"(token|secret|password|passwd|credential|authorization)", key, re.IGNORECASE):
        return "[已隐藏]"
    if isinstance(value, str):
        return _sanitize_tool_event_string(value, profile_home)
    if isinstance(value, dict):
        return {
            str(item_key): _sanitize_tool_event_value(item_value, profile_home, key=str(item_key))
            for item_key, item_value in value.items()
        }
    if isinstance(value, list):
        return [_sanitize_tool_event_value(item, profile_home, key=key) for item in value]
    if isinstance(value, tuple):
        return tuple(_sanitize_tool_event_value(item, profile_home, key=key) for item in value)
    return value


def _sanitize_tool_event_string(text: str, profile_home: Optional[Path]) -> str:
    raw = str(text or "")
    if not raw:
        return raw
    cleaned = raw
    root: Optional[Path] = None
    if profile_home is not None:
        root = Path(profile_home).expanduser().resolve(strict=False)
        root_text = str(root)
        cleaned = cleaned.replace(root_text + "/", "")
        if cleaned == root_text:
            cleaned = "."
        workspace_root = str(root / "workspace")
        cleaned = cleaned.replace(workspace_root + "/", "workspace/")
    cleaned = re.sub(r"(?<!\w)/workspace/", "workspace/", cleaned)

    def repl(match: re.Match[str]) -> str:
        candidate = match.group(0).rstrip(".,;:)]}")
        suffix = match.group(0)[len(candidate) :]
        if root is not None:
            try:
                resolved = Path(candidate).expanduser().resolve(strict=False)
                if resolved == root:
                    return "." + suffix
                if root in resolved.parents:
                    return resolved.relative_to(root).as_posix() + suffix
            except Exception:
                pass
        return "[宿主路径已隐藏]" + suffix

    cleaned = re.sub(r"/(?:Users|home)/[^\s`\"'<>]+", repl, cleaned)
    return cleaned


_MEDIA_DIRECTIVE_RE = re.compile(r'''(?P<prefix>[`"']?MEDIA:\s*)(?P<path>\S+)(?P<suffix>[`"']?)''')
_ARTIFACT_JSON_RE = re.compile(r"```hermes-artifact-json\s*(?P<body>.*?)\s*```", re.DOTALL | re.IGNORECASE)
_PROFILE_FILE_PATH_RE = re.compile(
    r'''(?P<path>(?:/workspace|/[^`"'<>\n\r]+?)'''
    r'''\.(?:png|jpe?g|gif|webp|mp4|mov|avi|mkv|webm|ogg|opus|mp3|wav|m4a|'''
    r'''epub|pdf|zip|rar|7z|docx?|xlsx?|pptx?|txt|csv|json|md|markdown))'''
)
_FEISHU_DOCUMENT_ID_LINE_RE = re.compile(
    r"(?m)^(?P<prefix>\s*(?:🆔\s*)?(?:(?:飞书)?(?:云)?文档|wiki|Wiki|WIKI|知识库)\s*"
    r"(?:ID|Id|id|Token|token|令牌)?\s*[:：]\s*)"
    r"[`\"']?(?P<token>[A-Za-z0-9]{20,})[`\"']?(?P<suffix>\s*)$"
)
_AUTO_FILE_DELIVERY_MAX_BYTES = int(os.getenv("HERMES_MULTITENANCY_AUTO_FILE_DELIVERY_MAX_BYTES", "52428800"))
_MARKDOWN_DOCUMENT_EXTENSIONS = {".md", ".markdown"}
_SENSITIVE_PROFILE_FILE_NAMES = {
    ".env",
    "auth.json",
    "config.yaml",
    "credential-materialization.yaml",
    "credential-materialization.yml",
}
_SENSITIVE_PROFILE_DIR_NAMES = {
    ".aws",
    ".config",
    ".gnupg",
    ".ssh",
    "credentials",
    "feishu_uat",
    "tokens",
}


async def _deliver_media_from_stream_response(
    gateway: Any,
    response: str,
    event: Any,
    adapter: Any,
    profile_home: Path,
) -> None:
    """Delegate post-stream media attachment delivery to Hermes' native gateway path."""
    response = _materialize_response_artifacts(response, profile_home)
    response_with_files = _append_profile_file_media_directives(response, profile_home)
    scoped_response = _profile_scoped_media_response(response_with_files, profile_home)
    if "MEDIA:" not in scoped_response:
        return
    delivered = await _deliver_profile_scoped_media_directives(
        adapter,
        event,
        gateway,
        scoped_response,
        profile_home=profile_home,
    )
    if delivered:
        return
    deliver = getattr(gateway, "_deliver_media_from_response", None)
    if not callable(deliver):
        logger.warning("multitenancy: no post-stream media delivery surface available")
        return
    await deliver(scoped_response, event, adapter)


def _feishu_document_base_url() -> Optional[str]:
    configured = os.getenv("HERMES_FEISHU_DOCUMENT_BASE_URL", "").strip()
    return configured.rstrip("/") if configured else "https://feishu.cn"


def _linkify_feishu_document_ids(text: str) -> str:
    """Turn visible Feishu document tokens into clickable document URLs."""
    raw = str(text or "")
    if not raw:
        return raw
    base = _feishu_document_base_url()

    def repl(match: re.Match[str]) -> str:
        token = match.group("token")
        prefix = match.group("prefix").lower()
        resource = "wiki" if "wiki" in prefix or "知识库" in prefix else "docx"
        label = "飞书 Wiki 链接" if resource == "wiki" else "飞书文档链接"
        return f"🔗 {label}：{base}/{resource}/{token}"

    return _FEISHU_DOCUMENT_ID_LINE_RE.sub(repl, raw)


_MEDIA_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif"}
_MEDIA_VIDEO_EXTENSIONS = {".mp4", ".mov", ".avi", ".mkv", ".webm", ".3gp"}
_MEDIA_AUDIO_EXTENSIONS = {".mp3", ".wav", ".ogg", ".opus", ".m4a", ".flac"}


async def _deliver_profile_scoped_media_directives(
    adapter: Any,
    event: Any,
    gateway: Any,
    scoped_response: str,
    *,
    profile_home: Optional[Path] = None,
) -> int:
    """Deliver already-scoped MEDIA directives without upstream path regex loss."""
    if adapter is None:
        return 0
    source = getattr(event, "source", None)
    chat_id = str(getattr(source, "chat_id", "") or "")
    if not chat_id:
        return 0
    force_document = "[[as_document]]" in str(scoped_response or "")
    audio_as_voice = "[[audio_as_voice]]" in str(scoped_response or "")
    metadata = _thread_metadata_for_media_delivery(gateway, event)
    reply_to = _event_reply_to_message_id(event)
    delivered = 0
    seen: set[Path] = set()
    for match in _MEDIA_DIRECTIVE_RE.finditer(str(scoped_response or "")):
        raw_path = match.group("path").strip().strip("`\"'")
        path = Path(raw_path).expanduser().resolve(strict=False)
        if path in seen or not path.exists() or not path.is_file():
            continue
        seen.add(path)
        ext = path.suffix.lower()
        try:
            if ext in _MEDIA_IMAGE_EXTENSIONS and not force_document and hasattr(adapter, "send_image_file"):
                result = await adapter.send_image_file(
                    chat_id=chat_id,
                    image_path=str(path),
                    reply_to=reply_to,
                    metadata=metadata,
                )
            elif ext in _MEDIA_VIDEO_EXTENSIONS and hasattr(adapter, "send_video"):
                result = await adapter.send_video(
                    chat_id=chat_id,
                    video_path=str(path),
                    reply_to=reply_to,
                    metadata=metadata,
                )
            elif ext in _MEDIA_AUDIO_EXTENSIONS and audio_as_voice and hasattr(adapter, "send_voice"):
                result = await adapter.send_voice(
                    chat_id=chat_id,
                    audio_path=str(path),
                    reply_to=reply_to,
                    metadata=metadata,
                )
            elif hasattr(adapter, "send_document"):
                result = await adapter.send_document(
                    chat_id=chat_id,
                    file_path=str(path),
                    file_name=path.name,
                    reply_to=reply_to,
                    metadata=metadata,
                )
            else:
                continue
            if getattr(result, "success", True):
                delivered += 1
                if profile_home is not None:
                    _remember_recent_profile_file(profile_home.name, chat_id, path, profile_home)
                logger.info("multitenancy: delivered post-stream media attachment path=%s", path)
            else:
                logger.warning(
                    "multitenancy: post-stream media delivery failed path=%s error=%s",
                    path,
                    getattr(result, "error", None),
                )
        except Exception as exc:
            logger.warning("multitenancy: post-stream media delivery failed path=%s error=%s", path, exc)
    return delivered


def _thread_metadata_for_media_delivery(gateway: Any, event: Any) -> Optional[dict[str, Any]]:
    try:
        if not _is_group_chat_type(_extract_chat_type(event)):
            return None
        source = getattr(event, "source", None)
        reply_anchor = None
        anchor_fn = getattr(gateway, "_reply_anchor_for_event", None)
        if callable(anchor_fn):
            reply_anchor = anchor_fn(event)
        meta_fn = getattr(gateway, "_thread_metadata_for_source", None)
        if callable(meta_fn):
            metadata = meta_fn(source, reply_anchor)
            return dict(metadata) if isinstance(metadata, dict) else metadata
    except Exception as exc:
        logger.debug("multitenancy: thread metadata lookup for media delivery failed: %s", exc)
    return None


def _materialize_response_artifacts(response: str, profile_home: Path) -> str:
    """Write model-emitted artifact JSON blocks into profile workspace files.

    The model has no direct filesystem tool in some Feishu profiles. This
    controlled bridge lets it emit file content declaratively while preserving
    the existing outbound MEDIA safety boundary.
    """
    text = str(response or "")
    if "```hermes-artifact-json" not in text.lower():
        return text
    root = profile_home.expanduser().resolve(strict=False)
    workspace = (root / "workspace").resolve(strict=False)
    downloads = (workspace / "Downloads").resolve(strict=False)
    media_additions: list[str] = []

    for match in _ARTIFACT_JSON_RE.finditer(text):
        try:
            spec = json.loads(match.group("body"))
        except Exception as exc:
            logger.debug("multitenancy: invalid artifact json skipped: %s", exc)
            continue
        if not isinstance(spec, dict):
            continue
        raw_path = str(spec.get("path") or "").strip()
        filename = str(spec.get("filename") or spec.get("name") or "").strip()
        if raw_path == "/workspace" or raw_path.startswith("/workspace/"):
            candidate = workspace / raw_path.removeprefix("/workspace").lstrip("/")
        else:
            if raw_path:
                candidate = Path(raw_path).expanduser()
            elif filename:
                candidate = downloads / filename
            else:
                continue
            if not candidate.is_absolute():
                candidate = workspace / candidate
        target = candidate.resolve(strict=False)
        if not (target == downloads or downloads in target.parents):
            logger.warning("multitenancy: blocked response artifact outside workspace downloads path=%s", target)
            continue
        try:
            target.parent.mkdir(parents=True, exist_ok=True)
            existing_source = _existing_profile_source_for_artifact_spec(spec, target, root)
            if existing_source is not None:
                shutil.copy2(existing_source, target)
                logger.info(
                    "multitenancy: reused existing profile artifact source=%s target=%s",
                    existing_source,
                    target,
                )
            else:
                _write_response_artifact(target, spec)
            try:
                media_path = "/workspace/" + str(target.relative_to(workspace).as_posix())
            except ValueError:
                media_path = "/workspace/Downloads/" + target.name
            if f"MEDIA:{media_path}" not in text:
                if (
                    bool(spec.get("as_document"))
                    and target.suffix.lower() not in _MARKDOWN_DOCUMENT_EXTENSIONS
                    and "[[as_document]]" not in text
                ):
                    media_additions.append("[[as_document]]")
                media_additions.append(f"MEDIA:{media_path}")
        except Exception as exc:
            logger.warning("multitenancy: failed to materialize response artifact path=%s error=%s", target, exc)
    if not media_additions:
        return text
    return f"{text.rstrip()}\n" + "\n".join(media_additions)


def _existing_profile_source_for_artifact_spec(spec: dict[str, Any], target: Path, profile_home: Path) -> Optional[Path]:
    """Find a real profile-local source when an artifact JSON block is only a placeholder.

    Some skills already write the complete markdown to `.ai-docs/...` and then
    emit a marker-only artifact JSON block with the same filename. In that case
    the marker is a delivery hint, not file content.
    """
    fmt = str(spec.get("format") or target.suffix.lstrip(".")).lower()
    if fmt not in {"md", "markdown"} and target.suffix.lower() not in _MARKDOWN_DOCUMENT_EXTENSIONS:
        return None
    if any(spec.get(key) not in (None, "") for key in ("content", "data", "rows")):
        return None
    filename = Path(str(spec.get("filename") or spec.get("name") or target.name)).name
    if not filename:
        return None
    root = profile_home.expanduser().resolve(strict=False)
    resolved_target = target.expanduser().resolve(strict=False)
    search_roots = (
        root / ".ai-docs",
        root / "home" / "Downloads",
        root / "data",
        root / "tmp",
    )
    for search_root in search_roots:
        if not search_root.exists():
            continue
        try:
            matches = search_root.rglob(filename) if search_root.is_dir() else [search_root]
            for candidate in matches:
                resolved = candidate.expanduser().resolve(strict=False)
                if resolved == resolved_target:
                    continue
                if _is_deliverable_profile_file(resolved, root):
                    return resolved
        except OSError:
            continue
    return None


def _write_response_artifact(path: Path, spec: dict[str, Any]) -> None:
    fmt = str(spec.get("format") or path.suffix.lstrip(".")).lower()
    marker = str(spec.get("marker") or "")
    if fmt in {"md", "markdown", "txt", "csv"}:
        content = str(spec.get("content") or marker or "")
        path.write_text(content, encoding="utf-8")
    elif fmt == "json":
        data = spec.get("data")
        if data is None:
            data = {"marker": marker, "content": spec.get("content") or marker}
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    elif fmt == "xlsx":
        _write_artifact_xlsx(path, spec)
    elif fmt == "docx":
        _write_artifact_docx(path, spec)
    elif fmt == "pdf":
        _write_artifact_pdf(path, spec)
    elif fmt in {"png", "jpg", "jpeg"}:
        _write_artifact_image(path, spec)
    else:
        content = str(spec.get("content") or marker or "")
        path.write_text(content, encoding="utf-8")


def _artifact_rows(spec: dict[str, Any]) -> list[list[Any]]:
    rows = spec.get("rows")
    if isinstance(rows, list) and all(isinstance(row, list) for row in rows):
        return rows
    marker = str(spec.get("marker") or "")
    return [["marker", "value", "note"], [marker, 42, "generated by Hermes artifact bridge"]]


def _write_artifact_xlsx(path: Path, spec: dict[str, Any]) -> None:
    from openpyxl import Workbook  # type: ignore

    wb = Workbook()
    ws = wb.active
    ws.title = str(spec.get("sheet") or "matrix")[:31]
    for row in _artifact_rows(spec):
        ws.append(row)
    wb.save(path)


def _write_artifact_docx(path: Path, spec: dict[str, Any]) -> None:
    from docx import Document  # type: ignore

    doc = Document()
    title = str(spec.get("title") or "Hermes generated document")
    doc.add_heading(title, level=1)
    content = str(spec.get("content") or spec.get("marker") or "")
    for paragraph in content.splitlines() or [content]:
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    doc.save(path)


def _write_artifact_pdf(path: Path, spec: dict[str, Any]) -> None:
    from reportlab.lib.pagesizes import letter  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore

    c = canvas.Canvas(str(path), pagesize=letter)
    y = 720
    for line in str(spec.get("content") or spec.get("marker") or "").splitlines() or [str(spec.get("marker") or "")]:
        c.drawString(72, y, line[:120])
        y -= 24
        if y < 72:
            c.showPage()
            y = 720
    c.save()


def _write_artifact_image(path: Path, spec: dict[str, Any]) -> None:
    from PIL import Image, ImageDraw  # type: ignore

    marker = str(spec.get("marker") or spec.get("content") or "")
    image = Image.new("RGB", (1000, 420), color=(245, 250, 255))
    draw = ImageDraw.Draw(image)
    draw.rectangle((40, 40, 960, 380), outline=(20, 120, 80), width=5)
    draw.text((70, 120), marker[:100], fill=(0, 0, 0))
    draw.text((70, 190), "Hermes generated image artifact", fill=(20, 120, 80))
    image.save(path)


def _append_profile_file_media_directives(response: str, profile_home: Path) -> str:
    """Attach profile-local files that the model mentioned as plain paths."""
    text = str(response or "")
    if not text:
        return text
    root = profile_home.expanduser().resolve(strict=False)
    media_paths = {match.group("path").strip() for match in _MEDIA_DIRECTIVE_RE.finditer(text)}
    additions: list[str] = []
    seen: set[Path] = set()
    for match in _PROFILE_FILE_PATH_RE.finditer(text):
        raw_path = match.group("path").strip().rstrip(".,;:)]}")
        if not raw_path or raw_path in media_paths:
            continue
        published = _publish_mentioned_profile_file(raw_path, root)
        if published is None or published in seen:
            continue
        seen.add(published)
        additions.append(f"MEDIA:{published}")
    if not additions:
        return text
    return f"{text.rstrip()}\n" + "\n".join(additions)


def _strip_plain_profile_file_paths_for_display(text: str, profile_home: Path) -> str:
    """Avoid exposing host absolute paths when the file will be attached."""
    raw = str(text or "")
    if not raw:
        return raw
    root = profile_home.expanduser().resolve(strict=False)

    def repl(match: re.Match[str]) -> str:
        raw_path = match.group("path").strip().rstrip(".,;:)]}")
        if not raw_path:
            return match.group(0)
        published = _publish_mentioned_profile_file(raw_path, root)
        if published is not None:
            if _should_deliver_as_feishu_document(published):
                return "[Markdown 源文件已自动发送]"
            return "[文件已作为附件发送]"
        candidate = Path(raw_path).expanduser()
        if raw_path == "/workspace" or raw_path.startswith("/workspace/"):
            candidate = root / "workspace" / raw_path.removeprefix("/workspace").lstrip("/")
        if not candidate.is_absolute():
            candidate = root / candidate
        resolved = candidate.resolve(strict=False)
        if resolved.exists() and (resolved == root or root in resolved.parents):
            return "[受保护文件路径已隐藏]"
        return match.group(0)

    return _PROFILE_FILE_PATH_RE.sub(repl, raw)


def _should_deliver_as_feishu_document(path: Path) -> bool:
    return path.suffix.lower() in _MARKDOWN_DOCUMENT_EXTENSIONS


def _profile_scoped_media_response(response: str, profile_home: Path) -> str:
    """Drop MEDIA directives outside profile scope and publish artifacts to workspace."""
    root = profile_home.expanduser().resolve(strict=False)

    def repl(match: re.Match[str]) -> str:
        raw_path = match.group("path").strip()
        if raw_path == "/workspace" or raw_path.startswith("/workspace/"):
            workspace_relative = raw_path.removeprefix("/workspace").lstrip("/")
            candidate = root / "workspace" / workspace_relative
        else:
            candidate = Path(raw_path).expanduser()
        if not candidate.is_absolute():
            candidate = root / candidate
        resolved = candidate.resolve(strict=False)
        if resolved == root or root in resolved.parents:
            if not _is_deliverable_profile_file(resolved, root):
                logger.warning(
                    "multitenancy: blocked outbound MEDIA for non-deliverable profile file path=%s profile_home=%s",
                    resolved,
                    root,
                )
                return ""
            workspace_artifact = _publish_profile_media_artifact(resolved, root)
            deliver_path = workspace_artifact or resolved
            return f"{match.group('prefix')}{deliver_path}{match.group('suffix')}"
        profile_artifact = _resolve_profile_media_artifact(raw_path, root)
        if profile_artifact is not None:
            logger.info(
                "multitenancy: rewrote outbound MEDIA to profile workspace artifact path=%s resolved=%s",
                raw_path,
                profile_artifact,
            )
            return f"{match.group('prefix')}{profile_artifact}{match.group('suffix')}"
        logger.warning(
            "multitenancy: blocked outbound MEDIA outside profile home path=%s profile_home=%s",
            raw_path,
            root,
        )
        return ""

    return _MEDIA_DIRECTIVE_RE.sub(repl, str(response or ""))


def _webui_profile_scoped_media_response(response: str, profile_home: Path) -> str:
    """Scope outbound MEDIA and expose workspace files through browser-safe aliases."""
    root = profile_home.expanduser().resolve(strict=False)
    scoped = _profile_scoped_media_response(response, root)

    def repl(match: re.Match[str]) -> str:
        raw_path = match.group("path").strip()
        candidate = Path(raw_path).expanduser()
        if raw_path == "/workspace" or raw_path.startswith("/workspace/"):
            return match.group(0)
        if not candidate.is_absolute():
            candidate = root / candidate
        resolved = candidate.resolve(strict=False)
        workspace_root = (root / "workspace").resolve(strict=False)
        if resolved == workspace_root or workspace_root in resolved.parents:
            return f"{match.group('prefix')}{_workspace_alias_for_profile_file(resolved, root)}{match.group('suffix')}"
        return match.group(0)

    return _MEDIA_DIRECTIVE_RE.sub(repl, scoped)


def _publish_mentioned_profile_file(raw_path: str, profile_home: Path) -> Optional[Path]:
    if raw_path == "/workspace" or raw_path.startswith("/workspace/"):
        workspace_relative = raw_path.removeprefix("/workspace").lstrip("/")
        candidate = profile_home / "workspace" / workspace_relative
    else:
        candidate = Path(raw_path).expanduser()
    if not candidate.is_absolute():
        candidate = profile_home / candidate
    source = candidate.resolve(strict=False)
    if not _is_deliverable_profile_file(source, profile_home):
        return None
    workspace_root = (profile_home / "workspace").resolve(strict=False)
    if source == workspace_root or workspace_root in source.parents:
        return source
    target_dir = (workspace_root / "Downloads").resolve(strict=False)
    target = (target_dir / source.name).resolve(strict=False)
    if not (target == profile_home or profile_home in target.parents):
        return None
    try:
        target_dir.mkdir(parents=True, exist_ok=True)
        if source != target:
            shutil.copy2(source, target)
        logger.info(
            "multitenancy: auto-attached mentioned profile file source=%s target=%s",
            source,
            target,
        )
        return target
    except Exception as exc:
        logger.warning(
            "multitenancy: failed to auto-attach mentioned profile file source=%s target=%s error=%s",
            source,
            target,
            exc,
        )
        return None


def _is_deliverable_profile_file(source: Path, profile_home: Path) -> bool:
    root = profile_home.resolve(strict=False)
    if not (source.exists() and source.is_file() and root in source.parents):
        return False
    try:
        if source.stat().st_size > _AUTO_FILE_DELIVERY_MAX_BYTES:
            logger.info(
                "multitenancy: skipped auto file delivery for oversized file path=%s size=%s",
                source,
                source.stat().st_size,
            )
            return False
    except OSError:
        return False
    relative_parts = source.relative_to(root).parts
    lowered = [part.lower() for part in relative_parts]
    if any(part in _SENSITIVE_PROFILE_DIR_NAMES for part in lowered[:-1]):
        logger.warning("multitenancy: blocked auto file delivery for sensitive directory path=%s", source)
        return False
    name = lowered[-1] if lowered else ""
    if name in _SENSITIVE_PROFILE_FILE_NAMES:
        logger.warning("multitenancy: blocked auto file delivery for sensitive file path=%s", source)
        return False
    return True


def _remember_recent_profile_file(profile_name: str, chat_id: str, path: Path, profile_home: Path) -> None:
    if not profile_name or not chat_id:
        return
    resolved = path.expanduser().resolve(strict=False)
    if not _is_deliverable_profile_file(resolved, profile_home):
        return
    key = (profile_name, chat_id)
    existing = [item for item in _recent_profile_files_by_chat.get(key, []) if item != str(resolved)]
    existing.append(str(resolved))
    _recent_profile_files_by_chat[key] = existing[-_RECENT_PROFILE_FILE_CONTEXT_MAX:]


def _should_append_recent_profile_file_context(text: str) -> bool:
    return bool(_RECENT_FILE_CONTEXT_TRIGGER_RE.search(str(text or "")))


def _workspace_alias_for_profile_file(path: Path, profile_home: Path) -> str:
    workspace_root = (profile_home / "workspace").resolve(strict=False)
    resolved = path.expanduser().resolve(strict=False)
    try:
        relative = resolved.relative_to(workspace_root)
    except ValueError:
        return str(resolved)
    return "/workspace/" + relative.as_posix()


def _recent_profile_files_from_history(prior_messages: list[dict], profile_home: Path) -> list[tuple[Path, Path]]:
    candidates: list[tuple[Path, Path]] = []
    seen: set[str] = set()
    for message in reversed(prior_messages[-_SESSION_HISTORY_MAX:]):
        if message.get("role") not in {"assistant", "tool"}:
            continue
        content = str(message.get("content") or "")
        for match in _PROFILE_FILE_PATH_RE.finditer(content):
            raw_path = match.group("path").strip().strip("`\"'")
            source = Path(raw_path).expanduser()
            if not source.is_absolute():
                source = profile_home / source
            source = source.resolve(strict=False)
            published = _publish_mentioned_profile_file(raw_path, profile_home)
            if published is None:
                continue
            resolved = published.resolve(strict=False)
            if str(resolved) in seen:
                continue
            seen.add(str(resolved))
            candidates.append((source, resolved))
            if len(candidates) >= _RECENT_PROFILE_FILE_CONTEXT_MAX:
                return candidates
    return candidates


def _append_recent_profile_file_context(
    text: str,
    *,
    profile_name: str,
    chat_id: str,
    profile_home: Path,
    prior_messages: list[dict],
) -> str:
    raw = str(text or "")
    if not _should_append_recent_profile_file_context(raw):
        return raw

    candidates: list[tuple[Path, Path]] = []
    seen: set[str] = set()
    for stored in _recent_profile_files_by_chat.get((profile_name, chat_id), []):
        path = Path(stored).expanduser().resolve(strict=False)
        if _is_deliverable_profile_file(path, profile_home) and str(path) not in seen:
            candidates.append((path, path))
            seen.add(str(path))

    for source, path in _recent_profile_files_from_history(prior_messages, profile_home):
        if str(path) not in seen:
            candidates.append((source, path))
            seen.add(str(path))
        if len(candidates) >= _RECENT_PROFILE_FILE_CONTEXT_MAX:
            break

    if not candidates:
        return raw

    lines = [
        "",
        "",
        "[Hermes context: 最近 Hermes 已投递给当前会话的文件]",
    ]
    for source, path in candidates[:_RECENT_PROFILE_FILE_CONTEXT_MAX]:
        lines.extend(
            [
                f"- file_name: {path.name}",
                f"  workspace_path: {_workspace_alias_for_profile_file(path, profile_home)}",
                f"  profile_path: {path}",
            ]
        )
        if source != path:
            lines.append(f"  source_path: {source}")
    lines.append("[/Hermes context]")
    return raw + "\n".join(lines)


def _resolve_profile_media_artifact(raw_path: str, profile_home: Path) -> Optional[Path]:
    """Map tool-reported temp media paths to same-name artifacts in the workspace."""
    name = Path(raw_path).name
    if not name:
        return None
    search_dirs = (
        profile_home / "home",
        profile_home / "home" / "Downloads",
        profile_home / "cache" / "images",
        profile_home / "tmp",
        profile_home / "data",
    )
    for directory in search_dirs:
        candidate = (directory / name).resolve(strict=False)
        if candidate.exists() and candidate.is_file() and profile_home in candidate.parents:
            return _publish_profile_media_artifact(candidate, profile_home)
    return None


def _publish_profile_media_artifact(source: Path, profile_home: Path) -> Optional[Path]:
    """Copy a profile-local generated artifact into the WebUI-visible workspace."""
    source = source.resolve(strict=False)
    root = profile_home.resolve(strict=False)
    if not (source.exists() and source.is_file() and root in source.parents):
        return None
    workspace_root = (root / "workspace").resolve(strict=False)
    if source == workspace_root or workspace_root in source.parents:
        return source
    artifact_dirs = (
        root / "home" / "Downloads",
        root / "cache" / "images",
        root / "tmp",
        root / "data",
    )
    direct_home_file = source.parent == (root / "home").resolve(strict=False)
    source_in_artifacts = any(
        source == directory.resolve(strict=False)
        or directory.resolve(strict=False) in source.parents
        for directory in artifact_dirs
    ) or direct_home_file
    if not source_in_artifacts:
        return None
    target_dir = (workspace_root / "Downloads").resolve(strict=False)
    target = (target_dir / source.name).resolve(strict=False)
    if not (target == root or root in target.parents):
        return None
    try:
        target_dir.mkdir(parents=True, exist_ok=True)
        if source != target:
            shutil.copy2(source, target)
        return target
    except Exception as exc:
        logger.warning(
            "multitenancy: failed to publish media artifact to workspace source=%s target=%s error=%s",
            source,
            target,
            exc,
        )
        return None


_IMAGE_FILE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".heic"}


def _event_has_image_media(event: Any) -> bool:
    media_urls = getattr(event, "media_urls", None) or []
    media_types = getattr(event, "media_types", None) or []
    if not media_urls:
        return False
    message_type_obj = getattr(event, "message_type", None)
    message_type_parts = [
        str(getattr(message_type_obj, "value", "") or ""),
        str(getattr(message_type_obj, "name", "") or ""),
        str(message_type_obj or ""),
    ]
    message_type = " ".join(message_type_parts).lower()
    if "photo" in message_type or "image" in message_type:
        return True
    for raw_path, raw_mtype in zip_longest(media_urls, media_types, fillvalue=""):
        raw = str(raw_path or "")
        suffix = Path(raw).suffix.lower()
        media_type = str(raw_mtype or "").lower()
        normalized = raw.replace("\\", "/").lower()
        if media_type.startswith("image") or suffix in _IMAGE_FILE_EXTENSIONS or "/cache/images/" in normalized:
            return True
    return False


def _image_vision_unavailable_response(event: Any, enriched_text: Optional[str]) -> Optional[str]:
    """Return a direct blocked reply when upstream image analysis is unavailable."""
    if not _event_has_image_media(event):
        return None
    text = str(enriched_text or "")
    lowered = text.lower()
    timeout_markers = (
        "vision auto-analysis timed out",
        "image preprocessing timed out",
    )
    failure_markers = (
        "something went wrong when i tried to look at it",
        "couldn't quite see it",
        "vision auto-analysis error",
    )
    is_timeout = any(marker in lowered for marker in timeout_markers)
    if not (is_timeout or any(marker in lowered for marker in failure_markers)):
        return None

    names: list[str] = []
    for raw_path in getattr(event, "media_urls", None) or []:
        name = Path(str(raw_path or "")).name
        if name:
            names.append(name)
    suffix = f"\n已收到图片附件：{', '.join(names[:3])}。" if names else ""
    message_id = _event_message_id(event)
    message_note = f"\nFeishu message_id: {message_id}" if message_id else ""
    if is_timeout:
        return (
            "无法读取图片内容：已收到图片附件，但自动视觉分析超时。"
            f"{suffix}{message_note}\n可以稍后重试，或提高图片预分析超时时间后再试。"
        )
    return (
        "无法读取图片内容：当前图片视觉分析不可用，vision_analyze provider rejected the request。"
        f"{suffix}{message_note}\n请修复 profile 的 vision provider/key 后重试。"
    )


def _image_prep_unavailable_note(event: Any, *, reason: str = "provider") -> str:
    paths = ", ".join(str(path) for path in (getattr(event, "media_urls", None) or [])[:3])
    suffix = f" using image_url: {paths}" if paths else ""
    message_id = _event_message_id(event)
    message_note = f" Feishu message_id: {message_id}." if message_id else ""
    if reason == "timeout":
        return (
            "[The user sent an image but vision auto-analysis timed out before vision_analyze completed. "
            "You can try again later or examine it yourself with vision_analyze"
            f"{suffix}.{message_note}]"
        )
    return (
        "[The user sent an image but something went wrong when I tried to look at it~ "
        "You can try examining it yourself with vision_analyze"
        f"{suffix}.{message_note}]"
    )


def _materialize_inbound_media_for_profile(event: Any, profile_home: Path) -> None:
    """Copy gateway-cached inbound media into the routed profile boundary."""
    media_urls = getattr(event, "media_urls", None) or []
    if not media_urls:
        return
    media_types = getattr(event, "media_types", None) or []
    root = profile_home.resolve(strict=False)
    rewritten: list[str] = []
    replacements: dict[str, str] = {}
    changed = False

    for raw_path, raw_mtype in zip_longest(media_urls, media_types, fillvalue=""):
        raw = str(raw_path or "")
        if not raw or re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*://", raw):
            rewritten.append(raw)
            continue
        source = Path(raw).expanduser().resolve(strict=False)
        if not (source.exists() and source.is_file()):
            rewritten.append(raw)
            continue
        if source == root or root in source.parents:
            rewritten.append(str(source))
            continue

        suffix = source.suffix.lower()
        media_type = str(raw_mtype or "").lower()
        target_dir = root / "cache" / "images" if (
            media_type.startswith("image") or suffix in _IMAGE_FILE_EXTENSIONS
        ) else root / "uploads"
        target = (target_dir / source.name).resolve(strict=False)
        if target.exists() and target.resolve(strict=False) != source:
            digest = hashlib.sha1(str(source).encode("utf-8")).hexdigest()[:10]
            target = (target_dir / f"{source.stem}-{digest}{source.suffix}").resolve(strict=False)
        if not (target == root or root in target.parents):
            rewritten.append(raw)
            continue
        try:
            target_dir.mkdir(parents=True, exist_ok=True)
            if target.resolve(strict=False) != source:
                shutil.copy2(source, target)
            rewritten_path = str(target)
            rewritten.append(rewritten_path)
            replacements[raw] = rewritten_path
            changed = True
            logger.warning(
                "multitenancy: materialized inbound media for profile source=%s target=%s",
                source,
                target,
            )
        except Exception as exc:
            logger.warning(
                "multitenancy: failed to materialize inbound media source=%s profile_home=%s error=%s",
                source,
                root,
                exc,
            )
            rewritten.append(raw)

    if changed:
        setattr(event, "media_urls", rewritten)
        text = getattr(event, "text", None)
        if isinstance(text, str):
            for old, new in replacements.items():
                text = text.replace(old, new)
            setattr(event, "text", text)


async def _enrich_via_hermes_pipeline(event: Any, gateway: Any) -> Optional[str]:
    """Delegate inbound preprocessing to hermes' ``_prepare_inbound_message_text``.

    This is the single call that mainstream uses to:
      - run vision_analyze_tool on attached images
      - run transcribe_audio on voice messages
      - inject text-file content (.txt / .md / .csv / etc.)
      - prepend reply-quoted context
      - attribute multi-user shared sessions

    By calling the same gateway method, the plugin behaves *identically* to
    mainstream for every multimodal input — no re-implementation, no drift.

    Caveat: this depends on a private GatewayRunner method. If hermes-agent
    refactors ``_prepare_inbound_message_text``, swap to local fallbacks
    (``_local_enrich_with_vision_only`` below as a minimal safety net).

    Returns:
        Enriched text string, or None on failure (caller falls back to event.text).
    """
    native_text: Optional[str] = None
    if _event_has_image_media(event):
        strategy = os.getenv("HERMES_MULTITENANCY_IMAGE_PREP_STRATEGY", "gateway").strip().lower()
        if strategy in {"blocked", "block", "skip", "disabled", "off"}:
            logger.warning(
                "multitenancy: image preprocessing blocked by strategy message_id=%s",
                _event_message_id(event) or "",
            )
            return _image_prep_unavailable_note(event)

    prep = getattr(gateway, "_prepare_inbound_message_text", None)
    if gateway is None or prep is None or not callable(prep):
        logger.debug("multitenancy: gateway._prepare_inbound_message_text unavailable")
    else:
        source = getattr(event, "source", None)
        if source is not None:
            try:
                prep_call = prep(event=event, source=source, history=[])
                if _event_has_image_media(event):
                    timeout_s = float(os.getenv("HERMES_MULTITENANCY_IMAGE_PREP_TIMEOUT_S", "30"))
                    native_text = await asyncio.wait_for(prep_call, timeout=max(0.1, timeout_s))
                else:
                    native_text = await prep_call
            except asyncio.TimeoutError:
                logger.warning("multitenancy: image preprocessing timed out")
                native_text = _image_prep_unavailable_note(event, reason="timeout")
            except Exception as exc:
                logger.debug("multitenancy: gateway._prepare_inbound_message_text failed (%s)", exc)

    local_file_text = _local_enrich_with_file_content(event, existing_text=native_text or "")
    if local_file_text:
        return _append_enrichment(native_text or getattr(event, "text", "") or "", local_file_text)

    if native_text:
        return native_text
    return await _local_enrich_with_vision_only(event)


def _append_enrichment(base: str, enrichment: str) -> str:
    base = str(base or "").strip()
    enrichment = str(enrichment or "").strip()
    if not enrichment:
        return base
    if not base:
        return enrichment
    return f"{enrichment}\n{base}"


_TEXT_FILE_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".yaml", ".yml",
    ".log", ".xml", ".html", ".htm",
}
_MAX_LOCAL_ENRICH_FILE_BYTES = 10 * 1024 * 1024
_MAX_LOCAL_TEXT_PREVIEW_BYTES = 100_000
_MAX_XLSX_XML_BYTES = 2 * 1024 * 1024


def _local_enrich_with_file_content(event: Any, *, existing_text: str = "") -> Optional[str]:
    """Plugin-owned fallback for Feishu document attachments Hermes does not inline.

    Hermes core remains the first choice for multimodal preprocessing. This
    fallback only covers plain/tabular files already cached as local event paths,
    keeping compatibility in multitenancy instead of patching Hermes-agent.
    """
    media_urls = getattr(event, "media_urls", None) or []
    media_types = getattr(event, "media_types", None) or []
    if not media_urls:
        return None
    existing = str(existing_text or "")
    parts: list[str] = []
    for raw_path, raw_mtype in zip_longest(media_urls, media_types, fillvalue=""):
        path = Path(str(raw_path))
        if not path.is_file():
            continue
        name = path.name
        suffix = path.suffix.lower()
        if suffix not in {".xlsx", ".docx", ".pdf"} and f"[Content of {name}]" in existing:
            continue
        try:
            if path.stat().st_size > _MAX_LOCAL_ENRICH_FILE_BYTES:
                logger.debug("multitenancy: local file enrichment skipped oversized file %s", path)
                continue
        except OSError:
            continue
        media_type = str(raw_mtype or "").lower()
        try:
            if suffix == ".xlsx":
                content = _extract_xlsx_text(path)
            elif suffix == ".docx":
                content = _extract_docx_text(path)
            elif suffix == ".pdf":
                content = _extract_pdf_text(path)
            elif media_type.startswith("text/") or suffix in _TEXT_FILE_EXTENSIONS:
                with path.open("rb") as handle:
                    content = handle.read(_MAX_LOCAL_TEXT_PREVIEW_BYTES).decode("utf-8", errors="replace")
            else:
                continue
        except Exception as exc:
            logger.debug("multitenancy: local file enrichment failed for %s: %s", path, exc)
            continue
        content = content.strip()
        if not content or content in existing:
            continue
        header = f"[Content of {name}]"
        if suffix in {".xlsx", ".docx", ".pdf"} and header in existing:
            header = f"[Content of {name} - multitenancy {suffix.lstrip('.')} fallback]"
        parts.append(f"{header}:\n{content}")
    if not parts:
        return None
    return "\n\n".join(parts)


def _extract_xlsx_text(path: Path, *, max_sheets: int = 3, max_rows: int = 50, max_cells: int = 20) -> str:
    """Extract a small text preview from an XLSX file using only stdlib."""
    ns = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    rel_ns = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
    with zipfile.ZipFile(path) as zf:
        shared_strings: list[str] = []
        if "xl/sharedStrings.xml" in zf.namelist():
            root = ET.fromstring(_read_zip_member_limited(zf, "xl/sharedStrings.xml"))
            for si in root.findall("main:si", ns):
                texts = [node.text or "" for node in si.findall(".//main:t", ns)]
                shared_strings.append("".join(texts))

        sheet_names: dict[str, str] = {}
        if "xl/workbook.xml" in zf.namelist():
            workbook = ET.fromstring(_read_zip_member_limited(zf, "xl/workbook.xml"))
            for idx, sheet in enumerate(workbook.findall(".//main:sheet", ns), start=1):
                rel_id = sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                sheet_names[rel_id or f"rId{idx}"] = sheet.attrib.get("name") or f"sheet{idx}"

        rel_targets: list[tuple[str, str]] = []
        if "xl/_rels/workbook.xml.rels" in zf.namelist():
            rels = ET.fromstring(_read_zip_member_limited(zf, "xl/_rels/workbook.xml.rels"))
            for rel in rels.findall("rel:Relationship", rel_ns):
                target = rel.attrib.get("Target") or ""
                if target.startswith("worksheets/"):
                    rel_targets.append((rel.attrib.get("Id") or "", "xl/" + target))
        if not rel_targets:
            rel_targets = [
                (f"rId{idx}", name)
                for idx, name in enumerate(sorted(
                    n for n in zf.namelist() if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")
                ), start=1)
            ]

        sections: list[str] = []
        for rel_id, sheet_path in rel_targets[:max_sheets]:
            if sheet_path not in zf.namelist():
                continue
            root = ET.fromstring(_read_zip_member_limited(zf, sheet_path))
            rows: list[str] = []
            for row in root.findall(".//main:sheetData/main:row", ns)[:max_rows]:
                cells: list[str] = []
                for cell in row.findall("main:c", ns)[:max_cells]:
                    value = cell.find("main:v", ns)
                    raw = value.text if value is not None else ""
                    cell_type = cell.attrib.get("t")
                    if cell_type == "s" and raw.isdigit():
                        idx = int(raw)
                        raw = shared_strings[idx] if idx < len(shared_strings) else raw
                    elif cell_type == "inlineStr":
                        texts = [node.text or "" for node in cell.findall(".//main:t", ns)]
                        raw = "".join(texts)
                    cells.append(raw or "")
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                sections.append(f"[{sheet_names.get(rel_id, rel_id or sheet_path)}]\n" + "\n".join(rows))
        return "\n\n".join(sections)


def _read_zip_member_limited(zf: zipfile.ZipFile, name: str) -> bytes:
    info = zf.getinfo(name)
    if info.file_size > _MAX_XLSX_XML_BYTES:
        raise ValueError(f"xlsx member too large: {name}")
    return zf.read(name)


def _extract_docx_text(path: Path, *, max_paragraphs: int = 80) -> str:
    """Extract a small text preview from a DOCX file using only stdlib."""
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as zf:
        if "word/document.xml" not in zf.namelist():
            return ""
        root = ET.fromstring(_read_zip_member_limited(zf, "word/document.xml"))
        paragraphs: list[str] = []
        for para in root.findall(".//w:p", ns)[:max_paragraphs]:
            texts = [node.text or "" for node in para.findall(".//w:t", ns)]
            text = "".join(texts).strip()
            if text:
                paragraphs.append(text)
        return "\n".join(paragraphs)


_PDF_STREAM_RE = re.compile(rb"stream\r?\n(?P<body>.*?)\r?\n?endstream", re.DOTALL)
_PDF_TEXT_STRING_RE = re.compile(rb"\((?:\\.|[^\\()])*\)")


def _extract_pdf_text(path: Path, *, max_bytes: int = _MAX_LOCAL_TEXT_PREVIEW_BYTES) -> str:
    """Best-effort PDF text preview for small generated/text PDFs.

    This intentionally stays lightweight. It covers the real UAT fixtures
    generated by ReportLab (ASCII85Decode + FlateDecode text streams) and
    simple uncompressed text streams; scanned PDFs still require OCR/vision.
    """
    raw = path.read_bytes()[: _MAX_LOCAL_ENRICH_FILE_BYTES]
    chunks: list[str] = []
    for match in _PDF_STREAM_RE.finditer(raw):
        stream = match.group("body").strip()
        decoded = _decode_pdf_stream(stream)
        if not decoded:
            continue
        chunks.extend(_extract_pdf_literal_strings(decoded))
        if sum(len(item) for item in chunks) >= max_bytes:
            break
    if not chunks:
        chunks.extend(_extract_pdf_literal_strings(raw))
    text = "\n".join(item for item in chunks if item.strip())
    return text[:max_bytes]


def _decode_pdf_stream(stream: bytes) -> bytes:
    candidates = [stream]
    try:
        candidates.append(base64.a85decode(stream, adobe=True))
    except Exception:
        pass
    for candidate in list(candidates):
        try:
            decoded = zlib.decompress(candidate)
        except Exception:
            continue
        candidates.append(decoded)
    return candidates[-1] if candidates else b""


def _extract_pdf_literal_strings(raw: bytes) -> list[str]:
    values: list[str] = []
    for match in _PDF_TEXT_STRING_RE.finditer(raw):
        token = match.group(0)[1:-1]
        token = (
            token.replace(rb"\(", b"(")
            .replace(rb"\)", b")")
            .replace(rb"\\", b"\\")
            .replace(rb"\n", b"\n")
            .replace(rb"\r", b"\r")
            .replace(rb"\t", b"\t")
        )
        text = token.decode("utf-8", errors="replace").strip()
        if text:
            values.append(text)
    return values


async def _local_enrich_with_vision_only(event: Any) -> Optional[str]:
    """Local fallback if hermes' ``_prepare_inbound_message_text`` is unavailable.

    Only handles images (the most common multimodal input). Audio / files /
    reply context degrade gracefully — the model will see ``event.text`` only.
    """
    media_urls = getattr(event, "media_urls", None) or []
    media_types = getattr(event, "media_types", None) or []
    if not media_urls:
        return None
    try:
        from tools.vision_tools import vision_analyze_tool  # type: ignore
    except ImportError:
        return None
    import json as _json
    descriptions: list[str] = []
    for path, mtype in zip(media_urls, media_types or [""] * len(media_urls)):
        if mtype and not mtype.startswith("image"):
            continue
        try:
            result_json = await vision_analyze_tool(
                image_url=path,
                user_prompt="Describe this image in thorough detail.",
            )
            result = _json.loads(result_json) if isinstance(result_json, str) else result_json
            if isinstance(result, dict) and result.get("success"):
                descriptions.append(f"[Image: {result.get('analysis', '')}]")
            else:
                error = ""
                if isinstance(result, dict):
                    error = str(result.get("error") or result.get("analysis") or "").strip()
                descriptions.append(_image_analysis_unavailable_note(path, error))
        except Exception as exc:
            logger.debug("multitenancy: local vision fallback error on %s: %s", path, exc)
            descriptions.append(_image_analysis_unavailable_note(path, str(exc)))
    if not descriptions:
        return None
    base = getattr(event, "text", "") or ""
    return "\n".join(descriptions) + ("\n" + base if base else "")


def _image_analysis_unavailable_note(path: Any, error: str = "") -> str:
    """Return recoverable context when an image is present but vision fails."""
    reason = re.sub(r"\s+", " ", str(error or "")).strip()
    if len(reason) > 300:
        reason = reason[:297] + "..."
    suffix = f" Reason: {reason}" if reason else ""
    return (
        "[Image analysis unavailable: the image is attached at "
        f"{path}, but automatic vision analysis failed.{suffix} "
        "If the user asks about the screenshot, explain that a vision-capable "
        "model or permission is required, or ask the user to describe the image.]"
    )


# -- Hook entry point --------------------------------------------------------


def _register_session_guard_for_dispatch(event: Any, gateway: Any, task: Any) -> None:
    """Mark this dispatch's session as active so flush-batch re-dispatch is no-op.

    Feishu adapter coalesces multiple WS text chunks into a single message via
    _pending_text_batches, then flushes ~0.6s later by calling handle_message
    again through _handle_message_with_guards. Without this guard, the flush
    sees no active session in base.handle_message and spawns a second
    handle_async task (creating a duplicate streaming card).
    """
    try:
        from gateway.session import build_session_key  # type: ignore
    except Exception:
        return
    try:
        adapter = None
        adapters = getattr(gateway, "adapters", None)
        if isinstance(adapters, dict):
            adapter = adapters.get("feishu")
        if adapter is None:
            adapter = getattr(gateway, "feishu_adapter", None)
        if adapter is None:
            return
        active = getattr(adapter, "_active_sessions", None)
        if active is None:
            return
        source = getattr(event, "source", None)
        if source is None:
            return
        config = getattr(adapter, "config", None)
        extra = getattr(config, "extra", {}) if config is not None else {}
        if not isinstance(extra, dict):
            extra = {}
        session_key = build_session_key(
            source,
            group_sessions_per_user=bool(extra.get("group_sessions_per_user", True)),
            thread_sessions_per_user=bool(extra.get("thread_sessions_per_user", False)),
        )
        if not session_key:
            return
        import asyncio as _asyncio
        guard = _asyncio.Event()
        existing = active.get(session_key)
        if existing is not None and _synthetic_session_guards.get(session_key) is not existing:
            return
        active[session_key] = guard
        _synthetic_session_guards[session_key] = guard

        def _cleanup(_t: Any) -> None:
            try:
                if active.get(session_key) is guard and _synthetic_session_guards.get(session_key) is guard:
                    active.pop(session_key, None)
                    _synthetic_session_guards.pop(session_key, None)
            except Exception:
                pass

        task.add_done_callback(_cleanup)
    except Exception as exc:
        logger.debug("multitenancy: session guard registration failed: %s", exc)


def on_pre_gateway_dispatch(*, event: Any, gateway: Any, session_store: Any = None, **_kwargs) -> dict:
    """Sync hook callback (registered to ``pre_gateway_dispatch``).

    Schedules the async work as a background task and returns immediately
    with ``action: skip`` so the gateway main flow halts for this event.
    """
    try:
        if _should_defer_gateway_processing_complete(event):
            _defer_gateway_processing_complete(event, gateway)
        loop = asyncio.get_running_loop()
        task = loop.create_task(handle_async(event=event, gateway=gateway))
        task.add_done_callback(_log_task_failure)
        # Prevent Feishu's text-batch flush (feishu.py:_flush_text_batch_now,
        # ~0.6s after WS dispatch) from re-routing this message through
        # base.handle_message → pre_gateway_dispatch a SECOND time, which
        # would spawn a duplicate handle_async task and create a second
        # streaming card. Register a synthetic guard in the adapter's
        # _active_sessions so handle_message sees the session as busy.
        _register_session_guard_for_dispatch(event, gateway, task)
    except RuntimeError:
        # Test-only fallback: hook called from sync context (no running loop).
        if os.environ.get("PYTEST_CURRENT_TEST"):
            try:
                asyncio.run(handle_async(event=event, gateway=gateway))
            except Exception as exc:
                logger.warning("multitenancy: sync fallback dispatch failed: %s", exc)
        else:
            logger.error(
                "multitenancy: pre_gateway_dispatch invoked without a running loop "
                "and not in pytest — dropping event"
            )
    except Exception as exc:
        logger.warning("multitenancy: failed to schedule handle_async: %s", exc)
    return {"action": "skip", "reason": "multitenancy router took over"}


def _should_defer_gateway_processing_complete(event: Any) -> bool:
    """Return True when the async router owns visible processing lifecycle."""
    from .commands import parse_command

    source = getattr(event, "source", None)
    fallback_sender = getattr(source, "user_id", "unknown") if source else "unknown"
    sender = _resolve_sender_for_routing(event, fallback=fallback_sender)
    sender_alt = getattr(source, "user_id_alt", None) if source else None
    text = getattr(event, "text", "") or ""
    if parse_command(text) is not None:
        return False

    chat_type = _extract_chat_type(event)
    if _is_group_chat_type(chat_type):
        chat_id = _extract_chat_id(event)
        if not chat_id:
            return False
        table = _get_routing_table()
        if table is not None and table.lookup_by_chat_id(chat_id) is not None:
            return True
        return (
            _auto_provision_enabled()
            and table is not None
            and _has_cached_chat_inviter(chat_id)
        )

    _profile_name, profile_home = _resolve_route(sender, alt_id=sender_alt)
    if profile_home is not None:
        return True
    return (
        _auto_provision_enabled()
        and _get_routing_table() is not None
        and bool(sender)
        and sender != "unknown"
    )


def _defer_gateway_processing_complete(event: Any, gateway: Any) -> None:
    adapter = _get_feishu_adapter(gateway)
    defer = getattr(adapter, "defer_processing_complete", None)
    if not callable(defer):
        return
    try:
        defer(event)
    except Exception as exc:
        logger.debug("multitenancy: defer_processing_complete failed: %s", exc)


# -- Async dispatch ----------------------------------------------------------


async def handle_async(*, event: Any, gateway: Any) -> None:
    """Async dispatch — orchestrates routing + pool + adapter calls + commands."""
    from .commands import parse_command

    try:
        source = getattr(event, "source", None)
        chat_id = getattr(source, "chat_id", "unknown") if source else "unknown"
        fallback_sender = getattr(source, "user_id", "unknown") if source else "unknown"
        sender = _resolve_sender_for_routing(event, fallback=fallback_sender)
        if _is_feishu_open_id(sender):
            setattr(event, "sender_open_id", sender)
        text = getattr(event, "text", "") or ""

        if _is_reaction_synthetic_event(event, text):
            logger.info(
                "multitenancy: skipping Feishu reaction synthetic event "
                "text=%r message_id=%s chat_id=%s",
                text,
                _event_message_id(event) or "",
                chat_id,
            )
            return

        sender_alt = getattr(source, "user_id_alt", None) if source else None
        if getattr(event, "media_urls", None):
            logger.info(
                "multitenancy: handle_async media event message_id=%s text_len=%s media_urls=%s media_types=%s message_type=%s",
                _event_message_id(event) or "",
                len(str(text or "")),
                list(getattr(event, "media_urls", None) or []),
                list(getattr(event, "media_types", None) or []),
                str(getattr(event, "message_type", "")),
            )

        # Group-chat profile resolution — when the event is from a group/topic
        # chat, the route is keyed by chat_id (not the @-er's open_id). This
        # branch runs before slash-command short-circuit so /status & friends
        # see the group profile instead of the @-er's private profile.
        chat_type = _extract_chat_type(event)
        is_group_chat = _is_group_chat_type(chat_type)
        group_profile_name: Optional[str] = None
        group_profile_home: Optional[Path] = None
        if is_group_chat and chat_id and chat_id != "unknown":
            group_profile_name, group_profile_home = (
                await resolve_or_auto_provision_group_route(
                    chat_id=chat_id, gateway=gateway,
                )
            )

        # Slash command short-circuit (resolve route first so /status / /new
        # know which profile's history to inspect). When _resolve_route signals
        # a miss with profile_home=None, surface profile_name=None so command
        # handlers reply "未路由" instead of leaking the sender id.
        # Group messages start with leading @_all / @_user_N tokens that
        # Feishu prepends; strip them before delegating to parse_command so
        # ``@bot /feishu_auth`` is recognised as a slash command.
        command_source_text = _strip_leading_at_mentions(text) if is_group_chat else text
        cmd_pair = parse_command(command_source_text)
        if cmd_pair is not None:
            # Group profiles do not own any UAT — the whole auth command
            # family is hard-rejected so a curious member can't trigger an
            # OAuth dance that would store a per-user token under the
            # group's profile_home. Normalise first: Feishu lets users send
            # `/feishu_auth@bot` or sneak zero-width chars, and an exact
            # string-equality gate would let those through the day
            # feishu_auth becomes gateway-dispatchable.
            if is_group_chat and _is_blocked_group_command(cmd_pair[0]):
                adapter = _get_feishu_adapter(gateway)
                if adapter is not None:
                    await _safe_call(
                        adapter.send,
                        chat_id,
                        "群聊模式下不支持 /feishu_auth。"
                        "如需以你本人的身份调用飞书数据，请在与我私聊时执行。",
                    )
                return

            if is_group_chat:
                cmd_profile_name = group_profile_name
                cmd_profile_home = group_profile_home
            else:
                cmd_profile_name, cmd_profile_home = _resolve_route(sender, alt_id=sender_alt)
            cmd_profile = cmd_profile_name if cmd_profile_home is not None else None
            if _should_check_skill_slash_command(cmd_pair[0], gateway):
                async with _profile_gateway_context(
                    gateway,
                    event,
                    sender=sender,
                    sender_alt=sender_alt,
                    profile_name=cmd_profile,
                    profile_home=cmd_profile_home,
                    chat_id=chat_id,
                ):
                    skill_handled, skill_reply = _maybe_rewrite_skill_slash_command(
                        cmd_pair,
                        event,
                        gateway,
                        sender=sender,
                        sender_alt=sender_alt,
                        profile_name=cmd_profile,
                        profile_home=cmd_profile_home,
                        chat_id=chat_id,
                    )
            else:
                skill_handled, skill_reply = False, None
            if skill_handled:
                if skill_reply:
                    adapter = _get_feishu_adapter(gateway)
                    if adapter is not None:
                        await _safe_call(adapter.send, chat_id, skill_reply)
                    return
                text = getattr(event, "text", "") or ""
            else:
                await _handle_command(
                    cmd_pair,
                    sender,
                    sender_alt,
                    cmd_profile,
                    cmd_profile_home,
                    chat_id,
                    gateway,
                    event,
                )
                return

        # Routing: group already resolved above; sender-based path for p2p.
        if is_group_chat:
            profile_name, profile_home = group_profile_name, group_profile_home
            if profile_home is None:
                logger.info(
                    "multitenancy: no group route for chat_id=%s (inviter "
                    "not captured), ignoring",
                    chat_id,
                )
                adapter = _get_feishu_adapter(gateway)
                if adapter is not None:
                    await _safe_call(
                        adapter.send,
                        chat_id,
                        "👋 我还没有这个群的专属 Profile。"
                        "请移除我后再次拉我进群，让我捕获邀请人身份。",
                    )
                return
        else:
            profile_name, profile_home = _resolve_or_auto_provision_route(sender, alt_id=sender_alt)
            if profile_home is None:
                logger.info("multitenancy: no route for sender=%s, ignoring", sender)
                return

        adapter = _get_feishu_adapter(gateway)
        # Detect whether adapter supports the streaming/reaction APIs we use.
        # Real FeishuAdapter does; unit-test mocks typically don't.
        feishu_full = (
            adapter is not None
            and hasattr(adapter, "edit_message")
            and hasattr(adapter, "on_processing_start")
            and hasattr(adapter, "on_processing_complete")
        )

        # Multi-modal enrichment must happen before RunRequest admission because
        # file-only Feishu events have empty event.text.  The enriched content is
        # the real prompt and the dedupe/admission key should reflect it.
        _materialize_inbound_media_for_profile(event, profile_home)
        enriched_text = await _enrich_via_hermes_pipeline(event, gateway)
        vision_blocked = _image_vision_unavailable_response(event, enriched_text)
        if vision_blocked:
            logger.info(
                "multitenancy: sending image vision unavailable response profile=%s message_id=%s",
                profile_name,
                _event_message_id(event) or "",
            )
            hist_key = _history_key(profile_name, sender, sender_alt)
            user_msg = _build_user_message(event, text_override=enriched_text)
            _persist_turn(hist_key, user_msg, vision_blocked)
            if adapter is not None:
                await _safe_call(adapter.send, chat_id, vision_blocked)
            return
        run_content = enriched_text or text
        if not run_content and getattr(event, "media_urls", None):
            run_content = "[media attachment]"

        run_request = _run_request_for_routed_event(
            event=event,
            profile_name=profile_name,
            sender=sender,
            sender_alt=sender_alt,
            chat_id=chat_id,
            text=run_content,
        )
        from .run_broker import RunRejected

        try:
            run_admission = await _make_routed_run_broker().admit(run_request)
        except RunRejected as exc:
            logger.warning("multitenancy: routed run rejected profile=%s sender=%s: %s", profile_name, sender, exc)
            if feishu_full:
                try:
                    out = _processing_outcome(failed=True)
                    complete_deferred = getattr(adapter, "complete_deferred_processing", None)
                    if callable(complete_deferred):
                        await complete_deferred(event, out)
                    else:
                        await adapter.on_processing_complete(event, out)
                except Exception as complete_exc:
                    logger.debug("multitenancy: rejected processing_complete failed: %s", complete_exc)
            return

        if run_admission.duplicate:
            logger.info(
                "multitenancy: duplicate inbound event skipped profile=%s sender=%s message_id=%s",
                profile_name,
                sender,
                _event_message_id(event) or "",
            )
            if feishu_full:
                try:
                    out = _processing_outcome(failed=False)
                    complete_deferred = getattr(adapter, "complete_deferred_processing", None)
                    if callable(complete_deferred):
                        await complete_deferred(event, out)
                    else:
                        await adapter.on_processing_complete(event, out)
                except Exception as exc:
                    logger.debug("multitenancy: duplicate processing_complete failed: %s", exc)
            return

        # Register self in the context-scoped in-flight slot (replace previous)
        current = asyncio.current_task()
        inflight_key = _inflight_key(profile_name, sender, sender_alt, chat_id)
        prev = _user_inflight_tasks.get(inflight_key)
        if prev is not None and not prev.done() and prev is not current:
            prev_hist_key = _user_inflight_history_keys.get(inflight_key)
            if prev_hist_key is not None:
                _persist_interruption_marker(prev_hist_key)
            _suppress_interruption_marker_tasks.add(prev)
            prev.cancel()
        if current is not None:
            _user_inflight_tasks[inflight_key] = current

        outcome_failed = False
        if feishu_full:
            try:
                await adapter.on_processing_start(event)
            except Exception as exc:
                logger.debug("multitenancy: on_processing_start failed: %s", exc)

        # Build the conversation: prior history + current user message (with
        # reply context spliced in). The runner prepends the profile's SOUL.
        # First lookup for a (profile, user) pair hydrates from SessionStore.
        hist_key = _history_key(profile_name, sender, sender_alt)
        prior = _load_history(hist_key)
        contextual_text = _append_recent_profile_file_context(
            enriched_text or text,
            profile_name=profile_name,
            chat_id=chat_id,
            profile_home=profile_home,
            prior_messages=prior,
        )
        user_msg = _build_user_message(event, text_override=contextual_text)
        conversation = prior + [user_msg]
        _persist_user_message(hist_key, user_msg)
        if current is not None and _user_inflight_tasks.get(inflight_key) is current:
            _user_inflight_history_keys[inflight_key] = hist_key
        agent_event = _event_with_text(event, user_msg["content"])

        try:
            if feishu_full:
                # Streaming path — card stream when available; text edit fallback.
                async def _dispatch_streaming(_request):
                    stream_kwargs = {"messages": conversation}
                    try:
                        if "gateway" in inspect.signature(_stream_into_feishu).parameters:
                            stream_kwargs["gateway"] = gateway
                    except (TypeError, ValueError):
                        stream_kwargs["gateway"] = gateway
                    stream_response = await _stream_into_feishu(
                        adapter, chat_id, profile_name, profile_home, agent_event,
                        **stream_kwargs,
                    )
                    if stream_response:
                        await _deliver_media_from_stream_response(
                            gateway, stream_response, agent_event, adapter, profile_home
                        )
                    return stream_response

                run_result = await _make_routed_run_broker(
                    dispatch_agent=_dispatch_streaming,
                ).run(run_request, admitted=True)
                response_text = run_result.content
            else:
                # Mock / minimal adapter — old non-stream path (send_typing + pool.dispatch + send)
                if adapter is not None:
                    await _safe_call(adapter.send_typing, chat_id)
                async def _dispatch_nonstream(_request):
                    return await _get_pool().dispatch(profile_name, profile_home, agent_event)

                run_result = await _make_routed_run_broker(
                    dispatch_agent=_dispatch_nonstream,
                ).run(run_request, admitted=True)
                response_text = run_result.content
                if adapter is not None:
                    await _safe_call(adapter.send, chat_id, response_text)

            # Record turn into history + persist to SessionStore.
            if response_text and isinstance(response_text, str):
                _persist_assistant_message(hist_key, response_text)

            _touch_route(sender, sender_alt)
        except asyncio.CancelledError:
            if current not in _suppress_interruption_marker_tasks:
                _persist_interruption_marker(hist_key)
            raise
        except Exception:
            outcome_failed = True
            _persist_failure_marker(hist_key)
            raise
        finally:
            if feishu_full:
                try:
                    out = _processing_outcome(failed=outcome_failed)
                    complete_deferred = getattr(adapter, "complete_deferred_processing", None)
                    if callable(complete_deferred):
                        await complete_deferred(event, out)
                    else:
                        await adapter.on_processing_complete(event, out)
                except Exception as exc:
                    logger.debug("multitenancy: on_processing_complete failed: %s", exc)
            if _user_inflight_tasks.get(inflight_key) is current:
                _user_inflight_tasks.pop(inflight_key, None)
                _user_inflight_history_keys.pop(inflight_key, None)
            if current is not None:
                _suppress_interruption_marker_tasks.discard(current)
    except asyncio.CancelledError:
        raise
    except Exception as exc:
        logger.exception("multitenancy: handle_async failed: %s", exc)


def _processing_outcome(*, failed: bool) -> Any:
    """Return Hermes' ProcessingOutcome enum, or a string-compatible fallback."""
    try:
        from gateway.platforms.base import ProcessingOutcome  # type: ignore

        return ProcessingOutcome.FAILURE if failed else ProcessingOutcome.SUCCESS
    except Exception:
        class _FallbackOutcome:
            def __str__(self) -> str:
                status = "FAILURE" if failed else "SUCCESS"
                return f"ProcessingOutcome.{status}"

        return _FallbackOutcome()


# -- Command dispatch --------------------------------------------------------


def _maybe_rewrite_skill_slash_command(
    pair: tuple[str, str],
    event: Any,
    gateway: Any,
    *,
    sender: str,
    sender_alt: Optional[str],
    profile_name: Optional[str],
    profile_home: Optional[Path],
    chat_id: str,
) -> tuple[bool, Optional[str]]:
    """Rewrite Hermes skill slash commands into the native skill invocation text.

    Native gateway treats ``/skill-name args`` as an agent turn after loading
    the skill instructions. The multitenancy router sees the slash first, so it
    must preserve that behavior instead of replying "unknown command".
    """
    cmd, args = pair
    try:
        from .commands import is_known_command

        if is_known_command(cmd) or _gateway_handler_for_command(gateway, cmd) is not None:
            return False, None
        if _get_quick_command(gateway, cmd) is not None:
            return False, None
        if _get_plugin_command_handler(cmd) is not None:
            return False, None

        from agent.skill_commands import (  # type: ignore
            build_skill_invocation_message,
            get_skill_commands,
            resolve_skill_command_key,
        )

        skill_cmds = get_skill_commands()
        cmd_key = resolve_skill_command_key(cmd)
        if cmd_key is None:
            alias = _SKILL_SLASH_ALIASES.get(cmd.replace("_", "-"))
            if alias:
                cmd_key = resolve_skill_command_key(alias)
        if cmd_key is None:
            return False, None

        skill_info = skill_cmds.get(cmd_key) or {}
        skill_name = skill_info.get("name", "")
        platform = _event_platform_value(event)
        if platform and skill_name:
            try:
                from agent.skill_utils import get_disabled_skill_names  # type: ignore

                if skill_name in get_disabled_skill_names(platform=platform):
                    return (
                        True,
                        f"The **{skill_name}** skill is disabled for {platform}.\n"
                        "Enable it with: `hermes skills config`",
                    )
            except Exception as exc:
                logger.debug("multitenancy: skill disabled check failed (%s)", exc)

        old_skill_dir = skill_info.get("skill_dir")
        relative_skill_dir = _profile_relative_skill_dir(skill_info, profile_home)
        if relative_skill_dir:
            skill_info["skill_dir"] = relative_skill_dir
        try:
            msg = build_skill_invocation_message(
                cmd_key,
                args.strip(),
                task_id=_multitenant_gateway_session_key(
                    event,
                    profile_name=profile_name,
                    sender=sender,
                    sender_alt=sender_alt,
                    chat_id=chat_id,
                ),
            )
        finally:
            if relative_skill_dir:
                skill_info["skill_dir"] = old_skill_dir
        if not msg:
            return False, None
        logger.info("Hermes skill slash invocation: %s profile=%s", cmd_key, profile_name or "")
        setattr(event, "text", msg)
        return True, None
    except Exception as exc:
        logger.debug("multitenancy: skill command passthrough failed (%s)", exc)
        return False, None


def _should_check_skill_slash_command(cmd: str, gateway: Any) -> bool:
    """Only unknown slash commands need profile-scoped skill alias lookup.

    Known Hermes commands such as ``/stop`` must not wait on the profile env
    lock; that lock may be held by the very in-flight run the command is trying
    to cancel.
    """
    try:
        from .commands import is_known_command

        if is_known_command(cmd):
            return False
    except Exception:
        pass
    if _gateway_handler_for_command(gateway, cmd) is not None:
        return False
    if _get_quick_command(gateway, cmd) is not None:
        return False
    if _get_plugin_command_handler(cmd) is not None:
        return False
    return True


async def _handle_command(
    pair: tuple[str, str],
    sender: str,
    sender_alt: Optional[str],
    profile_name: Optional[str],
    profile_home: Optional[Path],
    chat_id: str,
    gateway: Any,
    event: Any,
) -> None:
    """Execute a parsed slash command and reply via the shared adapter."""
    cmd, _args = pair
    adapter = _get_feishu_adapter(gateway)

    approval_reply = _handle_pending_approval_command(
        cmd,
        _args,
        event,
        profile_name=profile_name,
        sender=sender,
        sender_alt=sender_alt,
        chat_id=chat_id,
    )
    if approval_reply is not None:
        reply = approval_reply
    elif cmd == "stop":
        task = _cancel_inflight_task(
            _inflight_key(profile_name, sender, sender_alt, chat_id),
            preserve_resume_marker=True,
        )
        if task is not None and not task.done():
            reply = "已停止当前任务"
        else:
            reply = "没有进行中的任务"
    elif cmd == "status":
        task = _user_inflight_tasks.get(_inflight_key(profile_name, sender, sender_alt, chat_id))
        running = task is not None and not task.done()
        # Surface session memory size + profile so the user knows their context.
        if profile_name:
            hist = _session_history.get(_history_key(profile_name, sender, sender_alt), [])
            hist_len = len(hist)
        else:
            hist_len = 0
        reply = (
            f"状态: {'运行中' if running else '空闲'}\n"
            f"profile: {profile_name or '(未路由)'}\n"
            f"会话历史: {hist_len} 条消息"
        )
    elif cmd in ("new", "reset"):
        _cancel_inflight_task(
            _inflight_key(profile_name, sender, sender_alt, chat_id),
            preserve_resume_marker=False,
        )
        # Clear this user's per-profile history (cache + persistent SessionStore).
        if profile_name:
            key = _history_key(profile_name, sender, sender_alt)
            _clear_history(key)
            reply = "会话已重置 ✅"
        else:
            reply = "(未路由的用户) 没有历史可重置"
    elif cmd == "feishu-auth":
        await _handle_feishu_auth_command(
            args=_args,
            sender=sender,
            sender_alt=sender_alt,
            profile_name=profile_name,
            profile_home=profile_home,
            chat_id=chat_id,
            gateway=gateway,
            event=event,
        )
        return
    elif cmd == "auth":
        await _handle_auth_command(
            args=_args,
            sender=sender,
            sender_alt=sender_alt,
            profile_name=profile_name,
            profile_home=profile_home,
            chat_id=chat_id,
            gateway=gateway,
            event=event,
        )
        return
    elif cmd == "card":
        # Synthetic command from a card button click (core routes card actions as
        # "/card <tag> <value-json>"). Handle our credential-hub auth actions.
        handled = await _handle_hub_card_action(
            args=_args,
            sender=sender,
            sender_alt=sender_alt,
            profile_name=profile_name,
            profile_home=profile_home,
            chat_id=chat_id,
            gateway=gateway,
        )
        if handled:
            return
        # Not ours → ignore silently (other card actions are core's concern).
        return
    elif cmd == "help":
        reply = _gateway_help_text()
    else:
        dispatched = await _dispatch_gateway_command(
            cmd,
            event,
            gateway,
            sender=sender,
            sender_alt=sender_alt,
            profile_name=profile_name,
            profile_home=profile_home,
            chat_id=chat_id,
        )
        if dispatched is not None:
            reply = dispatched
        else:
            from .commands import is_known_command, unknown_command_message

            if is_known_command(cmd):
                reply = (
                    f"Command `/{cmd}` is recognized by Hermes, but this gateway does not "
                    "expose a reusable command dispatcher yet."
                )
            else:
                reply = unknown_command_message(cmd)
                logger.info("%s", reply)

    if adapter is not None:
        await _safe_call(adapter.send, chat_id, reply)


async def _handle_feishu_auth_command(
    *,
    args: str,
    sender: str,
    sender_alt: Optional[str],
    profile_name: Optional[str],
    profile_home: Optional[Path],
    chat_id: str,
    gateway: Any,
    event: Any,
) -> None:
    """Start a multitenancy-owned Feishu UAT device-flow auth session."""
    del profile_home
    from . import feishu_uat_auth
    from .feishu_auth_cards import auth_text_fallback, build_auth_card, send_auth_card

    adapter = _get_feishu_adapter(gateway)
    open_id = (
        _normalize_feishu_open_id(sender)
        or _normalize_feishu_open_id(sender_alt)
        or _profile_open_id_for_auth(profile_name)
        or ""
    )
    if not _is_feishu_open_id(open_id):
        if adapter is not None:
            await _safe_call(adapter.send, chat_id, "无法启动飞书授权：当前消息没有可用的 sender open_id。")
        return
    if not profile_name:
        if adapter is not None:
            await _safe_call(adapter.send, chat_id, "无法启动飞书授权：当前飞书用户还没有绑定 Hermes profile。")
        return

    scope = args.strip() or None
    try:
        session = feishu_uat_auth.start_session(
            profile_name=profile_name,
            open_id=open_id,
            scope=scope,
        )
    except feishu_uat_auth.FeishuUatAuthError as exc:
        if adapter is not None:
            await _safe_call(adapter.send, chat_id, f"无法启动飞书授权：{exc.message}")
        return

    session_id = str(session.get("session_id") or "")
    verification_uri = str(session.get("verification_uri") or "")
    user_code = str(session.get("user_code") or "")
    expires_at = float(session.get("expires_at") or 0)
    expires_min = 10
    if expires_at:
        expires_min = max(1, int((expires_at - time.time() + 59) // 60))
    auth_card = None
    if adapter is not None:
        card = build_auth_card(
            verification_uri=verification_uri,
            user_code=user_code,
            expires_min=expires_min,
            scope=scope,
        )
        auth_card = await send_auth_card(adapter=adapter, chat_id=chat_id, card=card)
        if auth_card is None:
            await _safe_call(
                adapter.send,
                chat_id,
                auth_text_fallback(verification_uri=verification_uri, user_code=user_code),
            )
    _start_feishu_auth_poll_task(
        session_id=session_id,
        profile_name=profile_name,
        open_id=open_id,
        chat_id=chat_id,
        gateway=gateway,
        event=event,
        interval=int(session.get("interval") or 3),
        auth_card=auth_card,
    )


def _profile_open_id_for_auth(profile_name: Optional[str]) -> Optional[str]:
    if not profile_name:
        return None
    table = _get_routing_table()
    if table is None:
        return None
    try:
        db_path = table.db_path
    except AttributeError:
        return None
    try:
        import sqlite3

        with sqlite3.connect(f"file:{db_path}?mode=ro", uri=True, timeout=2) as conn:
            row = conn.execute(
                "SELECT open_id FROM multitenancy_routing "
                "WHERE profile_name = ? AND active = 1 AND kind = 'user' "
                "AND open_id LIKE 'ou_%' ORDER BY updated_at DESC LIMIT 1",
                (profile_name,),
            ).fetchone()
    except Exception as exc:
        logger.debug("multitenancy: profile auth open_id lookup failed (%s)", exc)
        return None
    return str(row[0]) if row else None


def _start_feishu_auth_poll_task(
    *,
    session_id: str,
    profile_name: str,
    open_id: str,
    chat_id: str,
    gateway: Any,
    event: Any,
    interval: int,
    auth_card: Optional[dict[str, Any]] = None,
) -> None:
    if not session_id:
        return
    task = asyncio.create_task(
        _poll_feishu_auth_session_until_done(
            session_id=session_id,
            profile_name=profile_name,
            open_id=open_id,
            chat_id=chat_id,
            gateway=gateway,
            event=event,
            interval=interval,
            auth_card=auth_card,
        ),
        name=f"feishu-auth:{profile_name}:{open_id}:{session_id}",
    )
    task.add_done_callback(lambda t: logger.debug("Feishu auth poll task ended: %s", t.get_name()))


async def _poll_feishu_auth_session_until_done(
    *,
    session_id: str,
    profile_name: str,
    open_id: str,
    chat_id: str,
    gateway: Any,
    event: Any,
    interval: int,
    auth_card: Optional[dict[str, Any]] = None,
) -> None:
    from . import feishu_uat_auth
    from .feishu_auth_cards import (
        build_auth_failed_card,
        build_auth_identity_mismatch_card,
        build_auth_success_card,
        update_auth_card,
    )

    adapter = _get_feishu_adapter(gateway)
    current_interval = max(int(interval or 3), 2)
    while True:
        await asyncio.sleep(current_interval)
        try:
            session = await asyncio.to_thread(
                feishu_uat_auth.poll_session,
                session_id=session_id,
                profile_name=profile_name,
                open_id=open_id,
            )
        except feishu_uat_auth.FeishuUatAuthError as exc:
            if adapter is not None:
                message = str(exc.message or "")
                card = build_auth_identity_mismatch_card() if "does not match" in message else build_auth_failed_card(message)
                updated = await update_auth_card(adapter=adapter, auth_card=auth_card, card=card)
                if not updated:
                    await _safe_call(adapter.send, chat_id, f"飞书 UAT 授权失败：{exc.message}")
            return
        status = str(session.get("status") or "")
        if status == "pending":
            current_interval = max(int(session.get("interval") or current_interval), 2)
            continue
        if adapter is None:
            return
        if status == "success":
            updated = await update_auth_card(adapter=adapter, auth_card=auth_card, card=build_auth_success_card())
            if not updated:
                await _safe_call(adapter.send, chat_id, "✅ 飞书 UAT 授权完成，后续 lark_cli 将优先使用你的 user 身份。")
        elif status == "expired":
            updated = await update_auth_card(adapter=adapter, auth_card=auth_card, card=build_auth_failed_card("expired"))
            if not updated:
                await _safe_call(adapter.send, chat_id, "飞书 UAT 授权已过期，请重新发送 /feishu_auth。")
        else:
            error = str(session.get("error") or status or "unknown error")
            card = build_auth_identity_mismatch_card() if "does not match" in error else build_auth_failed_card(error)
            updated = await update_auth_card(adapter=adapter, auth_card=auth_card, card=card)
            if not updated:
                await _safe_call(adapter.send, chat_id, f"飞书 UAT 授权失败：{error}")
        return


async def _handle_auth_command(
    *,
    args: str,
    sender: str,
    sender_alt: Optional[str],
    profile_name: Optional[str],
    profile_home: Optional[Path],
    chat_id: str,
    gateway: Any,
    event: Any,
) -> None:
    """Render the ``/auth`` credential hub — a collection card of all credentials.

    ``/feishu_auth`` is the lark-cli/feishu row of this hub. lark-cli is wired
    end-to-end here (reuses the device-flow session to mint its authorize URL +
    background poll). keep-record / kep-cli report live status; their in-Feishu
    auth-start is a follow-up slice.
    """
    del event
    from . import credential_hub, feishu_uat_auth
    from .feishu_auth_cards import send_auth_card
    from .feishu_credential_hub_cards import build_hub_card

    adapter = _get_feishu_adapter(gateway)
    open_id = (
        _normalize_feishu_open_id(sender)
        or _normalize_feishu_open_id(sender_alt)
        or _profile_open_id_for_auth(profile_name)
        or ""
    )
    if not _is_feishu_open_id(open_id):
        if adapter is not None:
            await _safe_call(adapter.send, chat_id, "无法打开凭证中心：当前消息没有可用的 sender open_id。")
        return
    if not profile_name:
        if adapter is not None:
            await _safe_call(adapter.send, chat_id, "无法打开凭证中心：当前飞书用户还没有绑定 Hermes profile。")
        return

    # Use the router's authoritative profile home when available so the hub
    # reads the same dotfile root the agent runtime writes to.
    home_dir = (Path(profile_home) / "home") if profile_home else None
    rows = await asyncio.to_thread(
        credential_hub.collect_credential_statuses,
        profile_name=profile_name,
        open_id=open_id,
        home_dir=home_dir,
    )

    if adapter is None:
        return

    # /auth just renders status + callback buttons. No flow is started here — a
    # click delivers a synthetic /card COMMAND that _handle_hub_card_action picks
    # up and runs that credential's flow. (Click → flow → feedback, in-process.)
    card = build_hub_card(rows=rows)
    sent = await send_auth_card(adapter=adapter, chat_id=chat_id, card=card)
    if sent is None:
        lines = ["凭证中心："] + [
            f"- {row.title}: {'✅ 已认证' if row.authenticated else '⚠️ 未认证'}" for row in rows
        ]
        await _safe_call(adapter.send, chat_id, "\n".join(lines))


async def _handle_hub_card_action(
    *,
    args: str,
    sender: str,
    sender_alt: Optional[str],
    profile_name: Optional[str],
    profile_home: Optional[Path],
    chat_id: str,
    gateway: Any,
) -> bool:
    """Handle a credential-hub button click delivered as a synthetic /card command
    (text: ``card <tag> {value-json}``). Starts the credential's flow and sends
    the QR/URL card, then polls and pushes a 认证成功 card. Returns True if the
    action was ours (hub_action == "auth")."""
    import json as _json
    from . import credential_hub as ch, credential_hub_auth as cha, feishu_uat_auth
    from .feishu_auth_cards import send_auth_card
    from .feishu_credential_hub_cards import build_qr_card, build_url_card

    value: dict[str, Any] = {}
    brace = args.find("{")
    if brace >= 0:
        try:
            parsed = _json.loads(args[brace:])
            if isinstance(parsed, dict):
                value = parsed
        except Exception:
            value = {}
    logger.info("multitenancy: /card action args=%r parsed=%r", args, value)
    if value.get("hub_action") != "auth":
        return False
    cred = str(value.get("cred") or "")
    logger.info("multitenancy: /card hub auth cred=%r", cred)

    adapter = _get_feishu_adapter(gateway)
    open_id = (
        _normalize_feishu_open_id(sender)
        or _normalize_feishu_open_id(sender_alt)
        or _profile_open_id_for_auth(profile_name)
        or ""
    )
    if adapter is None or not _is_feishu_open_id(open_id) or not profile_name:
        return True
    shared = feishu_uat_auth.resolve_shared_home()
    pdir = Path(profile_home) if profile_home else (shared / "profiles" / profile_name)

    try:
        if cred == ch.LARK_CLI:
            session = await asyncio.to_thread(feishu_uat_auth.find_active_session,
                                              profile_name=profile_name, open_id=open_id)
            if not session:
                session = await asyncio.to_thread(feishu_uat_auth.start_session,
                                                  profile_name=profile_name, open_id=open_id)
            await send_auth_card(adapter=adapter, chat_id=chat_id,
                                 card=build_url_card("Lark-cli", str(session.get("verification_uri") or ""),
                                                     label_zh="前往授权"))
            _start_single_flow_poll(profile_name=profile_name, open_id=open_id, profile_dir=pdir,
                                    shared_home=shared, chat_id=chat_id, gateway=gateway, cred=cred,
                                    flow={"kind": "lark", "session_id": str(session.get("session_id") or "")})
        elif cred == ch.KEEP_RECORD:
            qr = await asyncio.to_thread(cha.start_keep_record_qr, pdir)
            image_key = await asyncio.to_thread(cha.fetch_qr_image_key, shared, qr["qrcode_url"])
            await send_auth_card(adapter=adapter, chat_id=chat_id,
                                 card=build_qr_card("Keep-record", image_key))
            _start_single_flow_poll(profile_name=profile_name, open_id=open_id, profile_dir=pdir,
                                    shared_home=shared, chat_id=chat_id, gateway=gateway, cred=cred,
                                    flow={"kind": "keep", "qrcode_id": qr["qrcode_id"]})
        elif cred == ch.KEP_CLI:
            origin = os.environ.get("HERMES_PUBLIC_CALLBACK_ORIGIN", "").strip() or None
            login = await asyncio.to_thread(cha.start_kep_cli_login, pdir, profile_name, shared,
                                            public_origin=origin)
            proc = login.get("_proc")
            _track_kep_login_proc(proc)
            await send_auth_card(adapter=adapter, chat_id=chat_id,
                                 card=build_url_card("kep-cli", login["verification_uri"], label_zh="前往登录"))
            _start_single_flow_poll(profile_name=profile_name, open_id=open_id, profile_dir=pdir,
                                    shared_home=shared, chat_id=chat_id, gateway=gateway, cred=cred,
                                    flow={"kind": "kep", "proc": proc})
        else:
            return True
    except cha.HubAuthError as exc:
        await _safe_call(adapter.send, chat_id, f"认证暂不可用：{exc.message}")
    except Exception as exc:  # pragma: no cover - defensive
        logger.debug("multitenancy: /card hub action %s failed (%s)", cred, exc)
        await _safe_call(adapter.send, chat_id, "认证暂时不可用，请稍后重试。")
    return True


def _start_single_flow_poll(*, profile_name, open_id, profile_dir, shared_home, chat_id, gateway, cred, flow) -> None:
    task = asyncio.create_task(
        _poll_single_flow(profile_name=profile_name, open_id=open_id, profile_dir=profile_dir,
                          shared_home=shared_home, chat_id=chat_id, gateway=gateway, cred=cred, flow=flow),
        name=f"auth-card:{cred}:{profile_name}:{open_id}",
    )
    task.add_done_callback(lambda t: logger.debug("auth card poll ended: %s", t.get_name()))


async def _poll_single_flow(*, profile_name, open_id, profile_dir, shared_home, chat_id, gateway, cred, flow) -> None:
    """Poll one credential's auth attempt; push a 认证成功 card on success. Silent
    on timeout/failure (button stays on the hub card for retry)."""
    from . import credential_hub as ch, credential_hub_auth as cha, feishu_uat_auth
    from .feishu_auth_cards import send_auth_card
    from .feishu_credential_hub_cards import build_success_card

    adapter = _get_feishu_adapter(gateway)
    if adapter is None:
        return
    titles = {ch.LARK_CLI: "Lark-cli", ch.KEEP_RECORD: "Keep-record", ch.KEP_CLI: "kep-cli"}
    ok = False
    for _ in range(40):
        try:
            if flow["kind"] == "lark":
                s = await asyncio.to_thread(feishu_uat_auth.poll_session, session_id=flow["session_id"],
                                            profile_name=profile_name, open_id=open_id)
                st = str(s.get("status") or "")
                if st == "success":
                    ok = True; break
                if st != "pending":
                    break
            elif flow["kind"] == "keep":
                r = await asyncio.to_thread(cha.poll_keep_record_once, profile_dir, flow["qrcode_id"])
                if r.get("status") == "authorized":
                    ok = True; break
            elif flow["kind"] == "kep":
                proc = flow.get("proc")
                rc = proc.poll() if proc is not None else 0
                if rc is not None:
                    ok = bool(rc == 0 and await asyncio.to_thread(
                        cha.kep_cli_logged_in, profile_dir, profile_name, shared_home))
                    break
        except Exception as exc:
            logger.debug("multitenancy: /card poll %s error (%s)", cred, exc)
            break
        await asyncio.sleep(3)
    proc = flow.get("proc")
    if proc is not None and proc.poll() is None:
        try:
            proc.kill()
        except Exception:
            pass
    if ok:
        expiry = ""
        try:
            rows = await asyncio.to_thread(ch.collect_credential_statuses, profile_name=profile_name,
                                           open_id=open_id, home_dir=profile_dir / "home")
            row = next((r for r in rows if r.id == cred), None)
            expiry = ch.human_expiry(row.expires_at) if row else ""
        except Exception:
            pass
        await send_auth_card(adapter=adapter, chat_id=chat_id,
                             card=build_success_card(titles.get(cred, cred), expiry_zh=expiry))


# Keep started kep-auth login subprocesses alive until their OAuth callback
# lands (a GC'd Popen would kill the login mid-flow).
_KEP_LOGIN_PROCS: set[Any] = set()


def _track_kep_login_proc(proc: Any) -> None:
    """Hold a reference to a live kep-auth login proc; prune finished ones so the
    set doesn't grow unbounded across /auth invocations."""
    if proc is None:
        return
    for old in [p for p in _KEP_LOGIN_PROCS if getattr(p, "poll", lambda: 0)() is not None]:
        _KEP_LOGIN_PROCS.discard(old)
    _KEP_LOGIN_PROCS.add(proc)


def _start_hub_flow_poll(
    *,
    profile_name: str,
    open_id: str,
    profile_dir: Path,
    shared_home: Path,
    chat_id: str,
    gateway: Any,
    hub_card: dict[str, Any],
    flows: dict[str, dict[str, Any]],
    auth_urls: dict[str, str],
    qr_image_keys: dict[str, str],
) -> None:
    if not flows:
        return
    task = asyncio.create_task(
        _poll_hub_flows(
            profile_name=profile_name,
            open_id=open_id,
            profile_dir=profile_dir,
            shared_home=shared_home,
            chat_id=chat_id,
            gateway=gateway,
            hub_card=hub_card,
            flows=flows,
            auth_urls=auth_urls,
            qr_image_keys=qr_image_keys,
        ),
        name=f"auth-hub:{profile_name}:{open_id}:{','.join(sorted(flows))}",
    )
    task.add_done_callback(lambda t: logger.debug("auth hub poll task ended: %s", t.get_name()))


async def _poll_hub_flows(
    *,
    profile_name: str,
    open_id: str,
    profile_dir: Path,
    shared_home: Path,
    chat_id: str,
    gateway: Any,
    hub_card: dict[str, Any],
    flows: dict[str, dict[str, Any]],
    auth_urls: dict[str, str],
    qr_image_keys: dict[str, str],
) -> None:
    """Poll the started auth flows. ONLY a credential the user actually completes
    produces feedback: its row flips to ✅已认证 via an in-place card update, and
    its button is dropped — every OTHER credential's button is preserved. Flows
    the user never acts on stay silent (no "未完成" noise); a failed attempt keeps
    its button so the user can retry."""
    from . import credential_hub, credential_hub_auth as cha, feishu_uat_auth
    from .feishu_auth_cards import send_auth_card, update_auth_card
    from .feishu_credential_hub_cards import build_hub_card, build_success_card

    adapter = _get_feishu_adapter(gateway)
    if adapter is None:
        return

    titles = {credential_hub.LARK_CLI: "Lark-cli", credential_hub.FEISHU_PROJECT: "飞书项目",
              credential_hub.KEEP_RECORD: "Keep-record", credential_hub.KEP_CLI: "kep-cli"}
    # Entries still offered on re-render. A credential drops out only once it
    # SUCCEEDS — so re-rendering after one completion keeps the others' buttons/QRs.
    remaining_urls = dict(auth_urls or {})
    remaining_qr = dict(qr_image_keys or {})

    async def _fresh_rows() -> list:
        try:
            return await asyncio.to_thread(
                credential_hub.collect_credential_statuses,
                profile_name=profile_name, open_id=open_id, home_dir=profile_dir / "home",
            )
        except Exception as exc:  # pragma: no cover - defensive
            logger.debug("multitenancy: /auth hub refresh failed (%s)", exc)
            return []

    async def _rerender(rows: list) -> None:
        await update_auth_card(adapter=adapter, auth_card=hub_card,
                               card=build_hub_card(rows=rows, auth_urls=remaining_urls,
                                                   pending_note={}, qr_image_keys=remaining_qr))

    pending = dict(flows)
    # Each flow keys off THIS attempt so re-auth of an already-authed credential
    # reports real completion, not a stale "already logged in". ~40 iterations;
    # keep's login-wait blocks ~15s/iter (QR window is minutes).
    for _ in range(40):
        if not pending:
            break
        succeeded: list[str] = []
        for cid, desc in list(pending.items()):
            try:
                if desc["kind"] == "lark":
                    s = await asyncio.to_thread(feishu_uat_auth.poll_session,
                                                session_id=desc["session_id"],
                                                profile_name=profile_name, open_id=open_id)
                    st = str(s.get("status") or "")
                    if st == "success":
                        succeeded.append(cid); pending.pop(cid, None)
                    elif st != "pending":  # expired/error: stop polling, keep button for retry
                        pending.pop(cid, None)
                elif desc["kind"] == "keep":
                    r = await asyncio.to_thread(cha.poll_keep_record_once, profile_dir, desc["qrcode_id"])
                    if r.get("status") == "authorized":
                        succeeded.append(cid); pending.pop(cid, None)
                    # not-yet-scanned → still pending; keep polling, no message
                elif desc["kind"] == "kep":
                    proc = desc.get("proc")
                    rc = proc.poll() if proc is not None else 0
                    if rc is not None:  # login proc exited
                        ok = await asyncio.to_thread(cha.kep_cli_logged_in, profile_dir, profile_name, shared_home)
                        if rc == 0 and ok:
                            succeeded.append(cid)
                        pending.pop(cid, None)  # proc finished either way; keep button if it failed
            except Exception as exc:  # stop polling a broken flow; keep its button, stay silent
                logger.debug("multitenancy: /auth flow %s poll error (%s)", cid, exc)
                pending.pop(cid, None)
        if succeeded:
            for cid in succeeded:
                remaining_urls.pop(cid, None)  # drop only the completed credential's entry
                remaining_qr.pop(cid, None)
            rows = await _fresh_rows()
            # Push a green 认证成功 card per completed credential (card-style feedback)…
            for cid in succeeded:
                row = next((r for r in rows if r.id == cid), None)
                expiry = credential_hub.human_expiry(row.expires_at) if row else ""
                await send_auth_card(adapter=adapter, chat_id=chat_id,
                                     card=build_success_card(titles.get(cid, cid), expiry_zh=expiry))
            # …and update the hub in place (completed row → ✅已认证, other buttons kept).
            await _rerender(rows)
        if pending:
            await asyncio.sleep(3)
    # Un-acted flows: kill any lingering kep login proc so it doesn't run forever.
    for desc in pending.values():
        proc = desc.get("proc")
        if proc is not None and proc.poll() is None:
            try:
                proc.kill()
            except Exception:
                pass


def _handle_pending_approval_command(
    cmd: str,
    args: str,
    event: Any,
    *,
    profile_name: Optional[str],
    sender: str,
    sender_alt: Optional[str],
    chat_id: str,
) -> Optional[str]:
    """Resolve a child AIAgent approval bridge before falling back to gateway commands."""
    if cmd not in {"approve", "deny"}:
        return None
    session_key = _multitenant_gateway_session_key(
        event,
        profile_name=profile_name,
        sender=sender,
        sender_alt=sender_alt,
        chat_id=chat_id,
    )
    if not session_key or not _pending_approval_requests.get(session_key):
        return None

    if cmd == "deny":
        resolve_all = "all" in str(args or "").lower().split()
        count = _resolve_pending_approval_requests(session_key, "deny", resolve_all=resolve_all)
        count_msg = f" ({count} commands)" if count > 1 else ""
        return f"❌ Command{'s' if count > 1 else ''} denied{count_msg}."

    parts = str(args or "").strip().lower().split()
    resolve_all = "all" in parts
    remaining = [part for part in parts if part != "all"]
    if any(part in {"always", "permanent", "permanently"} for part in remaining):
        choice = "always"
        scope_msg = " (pattern approved permanently)"
    elif any(part in {"session", "ses"} for part in remaining):
        choice = "session"
        scope_msg = " (pattern approved for this session)"
    else:
        choice = "once"
        scope_msg = ""
    count = _resolve_pending_approval_requests(session_key, choice, resolve_all=resolve_all)
    count_msg = f" ({count} commands)" if count > 1 else ""
    return f"✅ Command{'s' if count > 1 else ''} approved{scope_msg}{count_msg}. The agent is resuming..."


def _resolve_pending_approval_requests(
    session_key: str,
    choice: str,
    *,
    resolve_all: bool = False,
) -> int:
    queue = _pending_approval_requests.get(session_key) or []
    if not queue:
        return 0
    if resolve_all:
        targets = list(queue)
        queue.clear()
    else:
        targets = [queue.pop(0)]
    if queue:
        _pending_approval_requests[session_key] = queue
    else:
        _pending_approval_requests.pop(session_key, None)
    for entry in targets:
        raw_decision_path = str(entry.get("decision_path") or "").strip()
        if not raw_decision_path:
            continue
        decision_path = Path(raw_decision_path)
        try:
            decision_path.parent.mkdir(parents=True, exist_ok=True)
            decision_path.write_text(json.dumps({"choice": choice}), encoding="utf-8")
        except Exception as exc:
            logger.warning(
                "multitenancy: failed to write approval decision for %s: %s",
                entry.get("approval_id") or "?",
                exc,
            )
    return len(targets)


def _record_pending_approval(payload: dict) -> None:
    session_key = str(payload.get("session_key") or "").strip()
    decision_path = str(payload.get("decision_path") or "").strip()
    if not session_key or not decision_path:
        return
    approval_id = str(payload.get("approval_id") or decision_path)
    queue = _pending_approval_requests.setdefault(session_key, [])
    queue[:] = [
        item for item in queue
        if str(item.get("approval_id") or item.get("decision_path")) != approval_id
    ]
    queue.append(dict(payload))


def _clear_pending_approval(payload: dict) -> None:
    session_key = str(payload.get("session_key") or "").strip()
    approval_id = str(payload.get("approval_id") or "").strip()
    if not session_key or not approval_id:
        return
    queue = _pending_approval_requests.get(session_key) or []
    queue[:] = [item for item in queue if str(item.get("approval_id") or "") != approval_id]
    if queue:
        _pending_approval_requests[session_key] = queue
    else:
        _pending_approval_requests.pop(session_key, None)


async def _handle_child_approval_required(adapter: Any, chat_id: str, payload: Any) -> None:
    data = payload if isinstance(payload, dict) else {}
    _record_pending_approval(data)
    command = str(data.get("command") or "")
    description = str(data.get("description") or "dangerous command")
    logger.info(
        "multitenancy child approval_required session=%s approval_id=%s command=%s",
        str(data.get("session_key") or ""),
        str(data.get("approval_id") or ""),
        command[:120],
    )
    preview = command[:200] + "..." if len(command) > 200 else command
    message = (
        "⚠️ Dangerous command requires approval:\n"
        f"```\n{preview}\n```\n"
        f"Reason: {description}\n\n"
        "Reply `/approve` to execute, `/approve session` to approve this pattern "
        "for the session, `/approve always` to approve permanently, or `/deny` to cancel."
    )
    if adapter is not None:
        await _safe_call(adapter.send, chat_id, message)


async def _dispatch_gateway_command(
    cmd: str,
    event: Any,
    gateway: Any,
    *,
    sender: str,
    sender_alt: Optional[str],
    profile_name: Optional[str],
    profile_home: Optional[Path],
    chat_id: str,
) -> Optional[str]:
    """Delegate a Hermes-known slash command to the gateway when possible."""
    _ensure_command_event_methods(event, cmd)

    dispatcher = getattr(gateway, "_dispatch_slash_command", None)
    if callable(dispatcher):
        async with _profile_gateway_context(
            gateway,
            event,
            sender=sender,
            sender_alt=sender_alt,
            profile_name=profile_name,
            profile_home=profile_home,
            chat_id=chat_id,
        ):
            try:
                result = dispatcher(event, multitenancy_context={
                    "profile_name": profile_name,
                    "profile_home": str(profile_home) if profile_home else "",
                    "sender_open_id": sender,
                    "session_key_override": _multitenant_gateway_session_key(
                        event,
                        profile_name=profile_name,
                        sender=sender,
                        sender_alt=sender_alt,
                        chat_id=chat_id,
                    ),
                })
            except TypeError:
                result = dispatcher(event)
            if asyncio.iscoroutine(result):
                result = await result
            logger.info("Hermes gateway command handled: %s", cmd)
            return str(result) if result is not None else None

    handler = _gateway_handler_for_command(gateway, cmd)
    if handler is None:
        quick_result = await _dispatch_quick_command(
            cmd,
            event,
            gateway,
            sender=sender,
            sender_alt=sender_alt,
            profile_name=profile_name,
            profile_home=profile_home,
            chat_id=chat_id,
        )
        if quick_result is not None:
            return quick_result
        return await _dispatch_plugin_command(
            cmd,
            event,
            gateway,
            sender=sender,
            sender_alt=sender_alt,
            profile_name=profile_name,
            profile_home=profile_home,
            chat_id=chat_id,
        )
    async with _profile_gateway_context(
        gateway,
        event,
        sender=sender,
        sender_alt=sender_alt,
        profile_name=profile_name,
        profile_home=profile_home,
        chat_id=chat_id,
    ):
        result = handler(event)
        if asyncio.iscoroutine(result):
            result = await result
    logger.info("Hermes gateway command handled: %s", cmd)
    return str(result) if result is not None else None


async def _dispatch_quick_command(
    cmd: str,
    event: Any,
    gateway: Any,
    *,
    sender: str,
    sender_alt: Optional[str],
    profile_name: Optional[str],
    profile_home: Optional[Path],
    chat_id: str,
) -> Optional[str]:
    """Handle Hermes config.quick_commands without copying the command list."""
    qcmd = _get_quick_command(gateway, cmd)
    if not isinstance(qcmd, dict):
        return None

    kind = qcmd.get("type")
    if kind == "exec":
        exec_cmd = qcmd.get("command", "")
        if not exec_cmd:
            return f"Quick command '/{cmd}' has no command defined."
        if not _quick_exec_allowed(gateway, qcmd):
            return (
                f"Quick command '/{cmd}' exec is disabled for Feishu multitenancy. "
                "Enable only after profile sandboxing is in place."
            )
        logger.info("Hermes quick command exec: %s", cmd)
        try:
            env = os.environ.copy()
            if profile_home is not None:
                env["HERMES_HOME"] = str(profile_home)
            proc = await asyncio.create_subprocess_shell(
                exec_cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=env,
            )
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=30)
            output = (stdout or stderr).decode().strip()
            return output if output else "Command returned no output."
        except asyncio.TimeoutError:
            return "Quick command timed out (30s)."
        except Exception as exc:
            return f"Quick command error: {exc}"

    if kind == "alias":
        target = str(qcmd.get("target", "") or "").strip()
        if not target:
            return f"Quick command '/{cmd}' has no target defined."
        target = target if target.startswith("/") else f"/{target}"
        target_command = target.lstrip("/")
        new_cmd = target_command.split()[0] if target_command else ""
        if not new_cmd or new_cmd == cmd:
            return f"Quick command '/{cmd}' has invalid target."
        user_args = _command_args_from_event(event).strip()
        setattr(event, "text", f"{target} {user_args}".strip())
        _set_command_event_methods(event, new_cmd)
        logger.info("Hermes quick command alias: %s -> %s", cmd, target)
        return await _dispatch_gateway_command(
            new_cmd,
            event,
            gateway,
            sender=sender,
            sender_alt=sender_alt,
            profile_name=profile_name,
            profile_home=profile_home,
            chat_id=chat_id,
        )

    return f"Quick command '/{cmd}' has unsupported type (supported: 'exec', 'alias')."


def _get_quick_command(gateway: Any, cmd: str) -> Any:
    config = getattr(gateway, "config", None)
    if isinstance(config, dict):
        quick_commands = config.get("quick_commands", {}) or {}
    else:
        quick_commands = getattr(config, "quick_commands", {}) or {}
    if not isinstance(quick_commands, dict):
        return None
    return quick_commands.get(cmd)


def _quick_exec_allowed(gateway: Any, qcmd: dict[str, Any]) -> bool:
    """Return True only when multitenant Feishu quick exec is explicitly enabled."""
    qcmd_flag = qcmd.get("multitenancy_allow_exec")
    if qcmd_flag is None:
        qcmd_cfg = qcmd.get("multitenancy")
        if isinstance(qcmd_cfg, dict):
            qcmd_flag = qcmd_cfg.get("allow_exec")
    if qcmd_flag is not None:
        return _truthy(qcmd_flag)

    config = getattr(gateway, "config", None)
    plugin_cfg = None
    if isinstance(config, dict):
        plugin_cfg = config.get("multitenancy")
    else:
        plugin_cfg = getattr(config, "multitenancy", None)
    if isinstance(plugin_cfg, dict) and "allow_quick_exec" in plugin_cfg:
        return _truthy(plugin_cfg.get("allow_quick_exec"))

    env_value = os.getenv("HERMES_MULTITENANCY_ALLOW_QUICK_EXEC")
    return _truthy(env_value) if env_value is not None else False


def _truthy(value: Any) -> bool:
    return str(value).strip().lower() in {"1", "true", "yes", "on", "allow", "enabled"}


async def _dispatch_plugin_command(
    cmd: str,
    event: Any,
    gateway: Any,
    *,
    sender: str,
    sender_alt: Optional[str],
    profile_name: Optional[str],
    profile_home: Optional[Path],
    chat_id: str,
) -> Optional[str]:
    """Delegate plugin-registered slash commands to Hermes' plugin manager."""
    handler = _get_plugin_command_handler(cmd)
    if handler is None:
        return None
    logger.info("Hermes plugin slash handler: %s", cmd.replace("_", "-"))
    user_args = ""
    get_args = getattr(event, "get_command_args", None)
    if callable(get_args):
        user_args = (get_args() or "").strip()
    async with _profile_gateway_context(
        gateway,
        event,
        sender=sender,
        sender_alt=sender_alt,
        profile_name=profile_name,
        profile_home=profile_home,
        chat_id=chat_id,
    ):
        result = handler(user_args)
        if asyncio.iscoroutine(result):
            result = await result
    return str(result) if result else None


def _get_plugin_command_handler(cmd: str) -> Any:
    """Return Hermes' plugin command handler for ``cmd`` when available."""
    try:
        from hermes_cli.plugins import get_plugin_command_handler  # type: ignore

        return get_plugin_command_handler(cmd.replace("_", "-"))
    except Exception as exc:
        logger.debug("multitenancy: plugin command lookup failed (%s)", exc)
        return None


def _gateway_handler_for_command(gateway: Any, cmd: str) -> Any:
    """Return Hermes' handler method using naming conventions, not a command table."""
    normalized = cmd.replace("-", "_")
    candidates = [f"_handle_{normalized}_command"]
    if normalized == "sethome":
        candidates.append("_handle_set_home_command")
    for name in candidates:
        handler = getattr(gateway, name, None)
        if callable(handler):
            return handler
    return None


def _ensure_command_event_methods(event: Any, cmd: str) -> None:
    """Add minimal MessageEvent command helpers for tests/fallback objects."""
    args = _command_args_from_event(event)
    if not callable(getattr(event, "get_command", None)):
        setattr(event, "get_command", lambda: cmd)
    if not callable(getattr(event, "get_command_args", None)):
        setattr(event, "get_command_args", lambda: args)


def _set_command_event_methods(event: Any, cmd: str) -> None:
    args = _command_args_from_text(getattr(event, "text", "") or "")
    setattr(event, "get_command", lambda: cmd)
    setattr(event, "get_command_args", lambda: args)


def _command_args_from_event(event: Any) -> str:
    get_args = getattr(event, "get_command_args", None)
    if callable(get_args):
        return str(get_args() or "")
    return _command_args_from_text(getattr(event, "text", "") or "")


def _command_args_from_text(text: str) -> str:
    parts = text.split(maxsplit=1)
    return parts[1] if len(parts) > 1 else ""


def _event_platform_value(event: Any) -> Optional[str]:
    source = getattr(event, "source", None)
    platform = getattr(source, "platform", None) if source is not None else None
    value = getattr(platform, "value", platform)
    return str(value) if value else None


def _profile_relative_skill_dir(skill_info: dict[str, Any], profile_home: Optional[Path]) -> Optional[str]:
    if profile_home is None:
        return None
    raw_skill_dir = skill_info.get("skill_dir")
    if not isinstance(raw_skill_dir, str) or not raw_skill_dir:
        return None
    try:
        skill_dir = Path(raw_skill_dir)
        if not skill_dir.is_absolute():
            return None
        return str(skill_dir.relative_to(profile_home / "skills"))
    except Exception:
        return None


def _scope_profile_skill_loader(profile_home: Optional[Path]) -> list[tuple[Any, str, Any, bool]]:
    """Point upstream skill loader module caches at the routed profile."""
    if profile_home is None:
        return []
    states: list[tuple[Any, str, Any, bool]] = []

    def remember(module: Any, attr: str, value: Any) -> None:
        states.append((module, attr, getattr(module, attr, None), hasattr(module, attr)))
        setattr(module, attr, value)

    try:
        from tools import skills_tool  # type: ignore

        remember(skills_tool, "HERMES_HOME", profile_home)
        remember(skills_tool, "SKILLS_DIR", profile_home / "skills")
    except Exception as exc:
        logger.debug("multitenancy: profile skill loader scope skipped (%s)", exc)

    try:
        from agent import skill_commands  # type: ignore

        # ``agent.skill_commands`` caches slash skills per process, while this
        # router process serves many profiles.  Clear the cache inside the
        # scoped context so `/skill` resolves against this routed profile.
        remember(skill_commands, "_skill_commands", {})
        remember(skill_commands, "_skill_commands_platform", None)
    except Exception as exc:
        logger.debug("multitenancy: skill command cache scope skipped (%s)", exc)

    return states


def _restore_profile_skill_loader(states: list[tuple[Any, str, Any, bool]]) -> None:
    for module, attr, old_value, had_attr in reversed(states):
        try:
            if had_attr:
                setattr(module, attr, old_value)
            else:
                delattr(module, attr)
        except Exception:
            pass


@asynccontextmanager
async def _profile_gateway_context(
    gateway: Any,
    event: Any,
    *,
    sender: str,
    sender_alt: Optional[str],
    profile_name: Optional[str],
    profile_home: Optional[Path],
    chat_id: str,
):
    """Temporarily scope Hermes gateway helpers to the routed profile."""
    from .runtime import _get_env_lock

    async with _get_env_lock():
        old_home = os.environ.get("HERMES_HOME")
        had_home = "HERMES_HOME" in os.environ
        old_session_key_for_source = getattr(gateway, "_session_key_for_source", None)
        if profile_home is not None:
            os.environ["HERMES_HOME"] = str(profile_home)
        skill_loader_states = _scope_profile_skill_loader(profile_home)
        session_key = _multitenant_gateway_session_key(
            event,
            profile_name=profile_name,
            sender=sender,
            sender_alt=sender_alt,
            chat_id=chat_id,
        )
        if session_key:
            def _scoped_session_key_for_source(source):
                return session_key

            setattr(gateway, "_session_key_for_source", _scoped_session_key_for_source)
        try:
            yield
        finally:
            if old_session_key_for_source is not None:
                setattr(gateway, "_session_key_for_source", old_session_key_for_source)
            elif hasattr(gateway, "_session_key_for_source"):
                try:
                    delattr(gateway, "_session_key_for_source")
                except Exception:
                    pass
            _restore_profile_skill_loader(skill_loader_states)
            if had_home:
                os.environ["HERMES_HOME"] = old_home or ""
            else:
                os.environ.pop("HERMES_HOME", None)


def _multitenant_gateway_session_key(
    event: Any,
    *,
    profile_name: Optional[str],
    sender: str,
    sender_alt: Optional[str],
    chat_id: str,
) -> Optional[str]:
    if not profile_name:
        return None
    source = getattr(event, "source", None)
    platform = getattr(getattr(source, "platform", None), "value", None) or "feishu"
    user_key = _tenant_user_key(sender, sender_alt)
    return f"multitenancy:{platform}:{profile_name}:{chat_id}:{user_key}"


def _gateway_help_text() -> str:
    """Render help from Hermes' central command registry when available."""
    try:
        from hermes_cli.commands import gateway_help_lines  # type: ignore

        lines = gateway_help_lines()
        if lines:
            help_lines = [line for line in lines if "/help" in line]
            lines = (help_lines[:1] or ["`/help` -- 显示这条帮助"]) + [
                line for line in lines if "/help" not in line
            ]
            return "📖 可用命令\n" + "\n".join(lines[:30])
    except Exception:
        pass
    return (
        "📖 可用命令\n"
        "/help    — 显示这条帮助\n"
        "/status  — 查看当前 profile + 历史长度\n"
        "/new     — 重置会话历史 (开始新对话)\n"
        "/reset   — /new 的别名\n"
        "/stop    — 取消正在运行的任务\n"
        "/model   — 切换当前会话模型\n"
        "/reasoning — 管理推理强度\n"
        "/voice   — 切换语音回复模式\n"
    )


# -- Routing resolution ------------------------------------------------------


def _resolve_route(sender: str, *, alt_id: Optional[str] = None) -> tuple[str, Optional[Path]]:
    """Resolve sender → (profile_name, profile_home).

    The routing table's ``open_id`` column is overloaded as "any stable user
    identifier" — it can hold a real Feishu open_id (``ou_xxx``), a union_id
    (``on_xxx``), or any other tenant-stable token chosen by feishu-sync.

    Lookup order:
      1. SQLite RoutingTable WHERE open_id = sender (typical: real Feishu ou_* open_id)
      2. SQLite RoutingTable WHERE open_id = alt_id (legacy: source.user_id_alt = union_id)
      3. In-memory ``_SPIKE_ROUTING`` dict (Phase 1 compat / unit tests)

    Returns (sender, None) when no route hits.
    """
    from .runtime import resolve_profile_home as _spike_resolve

    table = _get_routing_table()
    candidates: list[str] = [sender]
    if alt_id and alt_id != sender:
        candidates.append(alt_id)

    if table is not None:
        for candidate in candidates:
            try:
                row = table.lookup_by_open_id(candidate)
            except Exception as exc:
                logger.debug("multitenancy: routing lookup failed (%s)", exc)
                continue
            if row is not None:
                return (row.profile_name, _profile_name_to_home(row.profile_name))
        # alt_id is the dedicated union_id channel — when sync chose to store a
        # tenant user_id placeholder in the open_id column, the real on_* still
        # lives in the union_id column. Query it directly so router doesn't end
        # up provisioning a duplicate route for the same physical user.
        if alt_id and isinstance(alt_id, str) and alt_id.startswith("on_"):
            try:
                row = table.lookup_by_union_id(alt_id)
            except Exception as exc:
                logger.debug("multitenancy: routing union_id lookup failed (%s)", exc)
            else:
                if row is not None:
                    return (row.profile_name, _profile_name_to_home(row.profile_name))

    # Spike routing dict (Phase 1 compat / unit tests)
    for candidate in candidates:
        spike_home = _spike_resolve(candidate)
        if spike_home is not None:
            return (spike_home.name, spike_home)
    return (sender, None)


def _resolve_or_auto_provision_route(
    sender: str,
    *,
    alt_id: Optional[str] = None,
) -> tuple[str, Optional[Path]]:
    """Resolve an existing route, or create a dedicated profile for a new sender."""
    alt_lookup = None if _is_feishu_open_id(sender) else alt_id
    profile_name, profile_home = _resolve_route(sender, alt_id=alt_lookup)
    if profile_home is not None:
        _repair_auto_profile(profile_name, profile_home, route_key=sender, sender=sender)
        return profile_name, profile_home
    if _auto_provision_enabled():
        provisioned = _auto_provision_route(sender, alt_id=alt_id)
        if provisioned is not None:
            return provisioned
        return profile_name, None
    return _resolve_route(sender, alt_id=alt_id)


def _repair_auto_profile(
    profile_name: str,
    profile_home: Path,
    *,
    route_key: str,
    sender: str,
) -> None:
    if not profile_name.startswith("feishu_"):
        return
    try:
        _ensure_auto_profile(profile_name, profile_home, route_key=route_key, sender=sender)
    except Exception as exc:
        logger.debug("multitenancy: auto profile repair failed for %s: %s", profile_name, exc)


def _auto_provision_enabled() -> bool:
    value = os.environ.get("HERMES_MULTITENANCY_AUTO_PROVISION", "1").strip().lower()
    return value not in {"0", "false", "no", "off"}


def _auto_provision_route(
    sender: str,
    *,
    alt_id: Optional[str] = None,
) -> Optional[tuple[str, Path]]:
    """Create a route/profile for an unseen Feishu user, then return it."""
    if not _auto_provision_enabled():
        return None
    table = _get_routing_table()
    if table is None:
        return None

    route_key = sender
    if not route_key or route_key == "unknown":
        return None

    profile_name = _auto_profile_name(route_key)
    profile_home = _profile_name_to_home(profile_name)
    try:
        _ensure_auto_profile(profile_name, profile_home, route_key=route_key, sender=sender)
        table.upsert(
            user_id=route_key,
            profile_name=profile_name,
            open_id=route_key,
            union_id=alt_id if alt_id and alt_id != sender else None,
        )
    except Exception as exc:
        logger.warning(
            "multitenancy: auto-provision failed sender=%s alt=%s: %s",
            sender,
            alt_id,
            exc,
        )
        return None

    logger.info(
        "multitenancy: auto-provisioned sender=%s alt=%s profile=%s",
        sender,
        alt_id,
        profile_name,
    )
    return profile_name, profile_home


def _auto_profile_name(route_key: str) -> str:
    """Return a deterministic, filesystem-safe profile name for a tenant key."""
    clean = re.sub(r"[^A-Za-z0-9_-]+", "_", route_key).strip("_")
    if not clean:
        clean = "unknown"
    return f"feishu_{clean}"


# -- Group-chat profile routing (Layer 2) -----------------------------------


def _extract_chat_type(event: Any) -> str:
    """Best-effort chat_type extraction.

    Feishu v1 message events expose ``source.chat_type`` (``"p2p"`` |
    ``"group"`` | ``"topic"``). Some test fixtures and webhook variants only
    have ``event.chat_type`` or ``event.message.chat_type``. Returns ``""``
    if nothing is available.
    """
    source = getattr(event, "source", None)
    candidates = (
        getattr(source, "chat_type", None) if source is not None else None,
        getattr(event, "chat_type", None),
        getattr(getattr(event, "message", None), "chat_type", None),
    )
    for candidate in candidates:
        if candidate:
            return str(candidate).strip().lower()
    return ""


def _extract_chat_id(event: Any) -> str:
    """Best-effort chat_id extraction with the same fallback chain as feishu.py."""
    source = getattr(event, "source", None)
    candidates = (
        getattr(source, "chat_id", None) if source is not None else None,
        getattr(source, "parent_chat_id", None) if source is not None else None,
        getattr(source, "chat_id_alt", None) if source is not None else None,
        getattr(event, "chat_id", None),
        getattr(getattr(event, "message", None), "chat_id", None),
    )
    for candidate in candidates:
        if candidate:
            return str(candidate)
    return ""


def _is_group_chat_type(chat_type: str) -> bool:
    return chat_type.lower() in _GROUP_CHAT_TYPES


# Commands that bind/replace a per-user Feishu identity. None of them make
# sense in a group profile (it owns no UAT), so all are hard-rejected.
_GROUP_BLOCKED_COMMANDS: frozenset[str] = frozenset(
    {"feishu_auth", "feishu-auth", "feishu_logout", "feishu-logout", "feishu_reauth", "feishu-reauth"}
)
# Zero-width / bidi chars Feishu clients occasionally inject; stripped before
# the command-name membership check so they can't smuggle a blocked command
# past an exact-match gate.
_INVISIBLE_CHARS = str.maketrans(
    "", "", "​‌‍‎‏﻿"
)


def _is_blocked_group_command(command: str) -> bool:
    """True when ``command`` is an auth-family command barred inside groups.

    Normalises before comparing: lowercase, strip a ``@bot`` suffix Feishu
    appends to at-mentioned commands, and drop zero-width characters. An
    exact ``== "feishu_auth"`` check let ``/feishu_auth@bot`` and
    zero-width-padded variants slip through.
    """
    if not command:
        return False
    normalized = command.translate(_INVISIBLE_CHARS).strip().lower()
    normalized = normalized.split("@", 1)[0]
    return normalized in _GROUP_BLOCKED_COMMANDS


_LEADING_MENTIONED_PREFIX_RE = re.compile(
    r"^\s*\[Mentioned:[^\]]*\]\s*",
)
_LEADING_AT_MENTION_RE = re.compile(
    r"^(?:\s*@\S+\s+)+",
)


def _strip_leading_at_mentions(text: str) -> str:
    """Strip leading Feishu @-mentions so slash commands are still parseable.

    The hermes-agent Feishu adapter normalises group messages into a form
    like ``[Mentioned: @all]\\n\\n@all /feishu_auth``. ``parse_command``
    requires the slash to be the first character, so for group messages
    we peel both the ``[Mentioned: …]`` prefix and the @-token tail off
    the front and hand the result to the parser.
    """
    if not text:
        return text
    text = _LEADING_MENTIONED_PREFIX_RE.sub("", text)
    text = _LEADING_AT_MENTION_RE.sub("", text)
    return text


def _short_chat_id(chat_id: str) -> str:
    """Return a deterministic short tag for a chat_id, safe in profile names.

    The 12-char prefix is cosmetic only — Feishu ``oc_`` ids are not
    uniform in their leading bytes, so similar-vintage ids share prefixes.
    Uniqueness rests entirely on a 16-hex (64-bit) sha1 suffix: collision
    probability stays below ~N²/2^65, negligible even for millions of
    chats. (The previous 4-hex/16-bit suffix collided at ~1/65536 per
    shared prefix, which would silently route two groups to one profile —
    the exact isolation this feature promises.)
    """
    raw = re.sub(r"[^A-Za-z0-9]+", "", str(chat_id or ""))
    if raw.lower().startswith("oc"):
        raw = raw[2:]
    prefix = raw[:12] or "noid"
    suffix = hashlib.sha1(str(chat_id).encode("utf-8")).hexdigest()[:16]
    return f"{prefix}_{suffix}"


def _make_group_profile_name(chat_id: str) -> str:
    return f"{_GROUP_PROFILE_PREFIX}{_short_chat_id(chat_id)}"


def is_group_profile_name(name: Optional[str]) -> bool:
    return bool(name) and str(name).startswith(_GROUP_PROFILE_PREFIX)


def register_chat_inviter(
    chat_id: str,
    inviter_open_id: str,
    *,
    chat_name: Optional[str] = None,
    inviter_display: Optional[str] = None,
) -> None:
    """Layer 4 hook entry — cache who pulled the bot into a chat.

    The cache survives only until the first @mention in the chat (or until
    the process restarts); the durable owner record is the routing row's
    ``owner_open_id`` column written by ``_provision_group_route``.
    """
    if not chat_id or not _is_feishu_open_id(inviter_open_id):
        return
    now = time.time()
    with _chat_inviter_cache_lock:
        _chat_inviter_cache[chat_id] = {
            "inviter_open_id": inviter_open_id,
            "chat_name": chat_name,
            "inviter_display": inviter_display,
            "_ts": now,
        }
        _chat_inviter_cache.move_to_end(chat_id)
        # Drop expired entries, then trim oldest if still over the cap.
        for cid in [
            c
            for c, e in _chat_inviter_cache.items()
            if now - e.get("_ts", 0) > _CHAT_INVITER_CACHE_TTL_S
        ]:
            _chat_inviter_cache.pop(cid, None)
        while len(_chat_inviter_cache) > _CHAT_INVITER_CACHE_MAX:
            _chat_inviter_cache.popitem(last=False)
    try:
        table = _get_routing_table()
        if table is None:
            logger.debug(
                "multitenancy: pending inviter store unavailable for chat_id=%s",
                chat_id,
            )
            return
        table.put_pending_inviter(chat_id, inviter_open_id)
        table.prune_pending_inviters(int(now), _CHAT_INVITER_CACHE_TTL_S)
    except Exception as exc:
        logger.debug(
            "multitenancy: failed to persist pending inviter chat_id=%s: %s",
            chat_id,
            exc,
        )


def _resolve_group_inviter_from_cache(chat_id: str) -> Optional[dict[str, Any]]:
    with _chat_inviter_cache_lock:
        entry = _chat_inviter_cache.get(chat_id)
        if entry is None:
            return None
        if time.time() - entry.get("_ts", 0) > _CHAT_INVITER_CACHE_TTL_S:
            _chat_inviter_cache.pop(chat_id, None)
            return None
        return dict(entry)


def _pop_chat_inviter(chat_id: str) -> None:
    with _chat_inviter_cache_lock:
        _chat_inviter_cache.pop(chat_id, None)


def _has_cached_chat_inviter(chat_id: str) -> bool:
    return _resolve_group_inviter_from_cache(chat_id) is not None


async def resolve_or_auto_provision_group_route(
    *,
    chat_id: str,
    gateway: Any,
) -> tuple[Optional[str], Optional[Path]]:
    """Resolve a group route by chat_id, auto-provisioning on first use.

    Group provisioning is intentionally independent from the unknown-user
    auto-provision toggle: a group route can only be created after a trusted
    ``bot_added`` event captures the inviter, so it does not open the same
    attack surface as creating user profiles for arbitrary senders. Returns
    ``(profile_name, profile_home)`` on success, ``(None, None)`` if no route
    exists and no inviter was captured for this chat. Inviter identity MUST
    come from the
    ``im.chat.member.bot.added_v1.operator_id`` cache entry (see
    ``register_chat_inviter``); we never fall back to the group's current
    owner_id because that is not the same person who added the bot.
    """
    if not chat_id:
        return None, None
    table = _get_routing_table()
    if table is None:
        return None, None
    row = table.lookup_by_chat_id(chat_id)
    if row is not None:
        profile_home = _profile_name_to_home(row.profile_name)
        try:
            _ensure_group_profile(
                profile_name=row.profile_name,
                profile_home=profile_home,
                chat_id=chat_id,
                owner_open_id=row.owner_open_id or "",
                display_label=row.display_label or row.profile_name,
            )
        except Exception as exc:
            logger.debug(
                "multitenancy: failed to normalize existing group profile chat_id=%s profile=%s: %s",
                chat_id,
                row.profile_name,
                exc,
            )
        return row.profile_name, profile_home
    return await _provision_group_route(chat_id=chat_id, gateway=gateway)


async def _provision_group_route(
    *,
    chat_id: str,
    gateway: Any,
) -> tuple[Optional[str], Optional[Path]]:
    """Create a new group routing row + profile skeleton.

    Returns ``(None, None)`` if no inviter is known for the chat — that is
    the explicit "refuse silently" path. The caller is responsible for
    surfacing a user-facing message.
    """
    table = _get_routing_table()
    if table is None:
        return None, None
    inviter_open_id: Optional[str] = None
    chat_name: Optional[str] = None
    inviter_display: Optional[str] = None
    try:
        table.prune_pending_inviters(
            int(time.time()),
            _CHAT_INVITER_CACHE_TTL_S,
        )
        pending_inviter = table.get_pending_inviter(chat_id)
    except Exception as exc:
        logger.debug(
            "multitenancy: pending inviter lookup failed chat_id=%s: %s",
            chat_id,
            exc,
        )
        pending_inviter = None
    if _is_feishu_open_id(pending_inviter):
        inviter_open_id = str(pending_inviter)
    else:
        cached = _resolve_group_inviter_from_cache(chat_id)
        if cached and _is_feishu_open_id(cached.get("inviter_open_id")):
            inviter_open_id = cached["inviter_open_id"]
            chat_name = cached.get("chat_name")
            inviter_display = cached.get("inviter_display")
    if not inviter_open_id:
        logger.warning(
            "multitenancy: cannot auto-provision group route chat_id=%s "
            "(bot_added inviter not captured — re-add the bot to retry)",
            chat_id,
        )
        return None, None

    if not chat_name:
        chat_name = await _fetch_chat_name(chat_id, gateway) or chat_id
    if not inviter_display:
        inviter_display = (
            await _fetch_user_display(inviter_open_id, gateway) or inviter_open_id
        )

    display_label = f"{inviter_display}-{chat_name}".strip("-") or chat_id
    profile_name = _make_group_profile_name(chat_id)
    profile_home = _profile_name_to_home(profile_name)
    try:
        _ensure_group_profile(
            profile_name=profile_name,
            profile_home=profile_home,
            chat_id=chat_id,
            owner_open_id=inviter_open_id,
            display_label=display_label,
        )
        table.upsert_group(
            chat_id=chat_id,
            profile_name=profile_name,
            owner_open_id=inviter_open_id,
            display_label=display_label,
        )
    except Exception as exc:
        logger.warning(
            "multitenancy: failed to provision group route chat_id=%s: %s",
            chat_id,
            exc,
        )
        return None, None

    logger.info(
        "multitenancy: auto-provisioned group profile chat_id=%s profile=%s owner=%s",
        chat_id,
        profile_name,
        inviter_open_id,
    )
    # Eagerly evict the cache entry so a later bot-removal + re-add cycle
    # repopulates from the fresh event rather than reusing stale chat_name.
    _pop_chat_inviter(chat_id)
    try:
        table.clear_pending_inviter(chat_id)
    except Exception as exc:
        logger.debug(
            "multitenancy: failed to clear pending inviter chat_id=%s: %s",
            chat_id,
            exc,
        )
    return profile_name, profile_home


def _ensure_group_profile(
    *,
    profile_name: str,
    profile_home: Path,
    chat_id: str,
    owner_open_id: str,
    display_label: str,
) -> None:
    """Create the on-disk profile skeleton for a group profile.

    Mirrors ``_ensure_auto_profile`` for shared config + SOUL.md, but adds a
    ``group_profile.json`` marker so downstream tools can detect the group
    context and refuse UAT-dependent operations.
    """
    profile_home.mkdir(parents=True, exist_ok=True)
    shared_home = _shared_home_for_profile(profile_home)

    config_path = profile_home / "config.yaml"
    if config_path.exists():
        _normalize_profile_config_file(config_path, shared_home=shared_home)
    else:
        config_path.write_text(
            _dump_profile_config(_profile_config_from_shared_home(shared_home)),
            encoding="utf-8",
        )

    soul_path = profile_home / "SOUL.md"
    if not soul_path.exists():
        soul_path.write_text(
            "\n".join(
                [
                    f"# Hermes Group Profile {profile_name}",
                    "",
                    f"You are a Feishu group-chat agent for chat `{chat_id}`.",
                    f"This profile is owned by Feishu user `{owner_open_id}` "
                    f"(display label: `{display_label}`).",
                    "Identity rules (strict):",
                    "- 你不能以群成员任何一个个人的身份操作飞书数据。",
                    "- /feishu_auth 在群聊模式下被禁用，任何用户的 UAT 不会被加载。",
                    "- 该群 profile 使用 `lark_cli` 时默认走 bot identity。",
                    "- 仅响应被 @ 的消息；群里的旁白不要回复。",
                    "- 该群的对话和记忆与其它群、与个人私聊完全隔离。",
                    "",
                    _LARK_CLI_SOUL_GUIDANCE,
                    "",
                    _GROUP_EXTERNAL_TOOL_SOUL_GUIDANCE,
                    "",
                ]
            ),
            encoding="utf-8",
        )
    else:
        _ensure_soul_guidance(soul_path, _LARK_CLI_SOUL_GUIDANCE)
        _ensure_soul_guidance(soul_path, _GROUP_EXTERNAL_TOOL_SOUL_GUIDANCE)

    # Group profiles get an empty feishu_uat/ directory but no per-user JSON;
    # the marker file below tells UAT helpers to refuse to load user tokens.
    (profile_home / "feishu_uat").mkdir(parents=True, exist_ok=True, mode=0o700)
    upstream_profile_home = _upstream_profile_home_for_owner(owner_open_id, shared_home)
    _sync_default_skills_for_profile(
        profile_home,
        shared_home,
        include_default_skills=True,
        upstream_profile_home=upstream_profile_home,
    )

    marker_path = profile_home / "group_profile.json"
    if not marker_path.exists():
        marker_path.write_text(
            json.dumps(
                {
                    "kind": "group",
                    "chat_id": chat_id,
                    "owner_open_id": owner_open_id,
                    "display_label": display_label,
                    "feishu_auth_disabled": True,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

    # The group profile's home channel IS the group itself. Pre-write
    # FEISHU_HOME_CHANNEL into .env so the gateway's "no home channel set"
    # onboarding prompt never fires for a group profile.
    _write_group_profile_env(profile_home, shared_home, chat_id)

    # .env is owned by us above; auth.json still follows the shared symlink.
    _ensure_shared_profile_file(profile_home, shared_home, "auth.json")


def _ensure_webui_agent_profile(
    *,
    profile_name: str,
    profile_home: Path,
    owner_open_id: str,
    display_label: str,
    agent_id: str,
    upstream_profile: str | None = None,
) -> None:
    """Create the on-disk skeleton for a WebUI group-chat agent profile.

    WebUI group-chat agents are not real Feishu ``oc_*`` groups, but they must
    run with the same tokenless/bot-only boundary as Feishu group profiles.
    """
    profile_home.mkdir(parents=True, exist_ok=True)
    shared_home = _shared_home_for_profile(profile_home)

    config_path = profile_home / "config.yaml"
    if config_path.exists():
        _normalize_profile_config_file(config_path, shared_home=shared_home)
        _disable_webui_agent_feishu_platform_file(config_path)
    else:
        config = _profile_config_from_shared_home(shared_home)
        _disable_webui_agent_feishu_platform(config)
        config_path.write_text(_dump_profile_config(config), encoding="utf-8")

    soul_path = profile_home / "SOUL.md"
    if not soul_path.exists():
        soul_path.write_text(
            "\n".join(
                [
                    f"# Hermes WebUI Group Agent Profile {profile_name}",
                    "",
                    f"You are a WebUI group-chat agent named `{display_label}`.",
                    f"This profile is owned by Feishu user `{owner_open_id}`.",
                    "Identity rules (strict):",
                    "- 你不能以 WebUI 群成员任何一个个人的身份操作飞书数据。",
                    "- /feishu_auth 在 WebUI 群聊 agent 模式下被禁用，任何用户的 UAT 不会被加载。",
                    "- 该 profile 使用 `lark_cli` 时默认走 bot identity。",
                    "- 该 profile 的对话和记忆与其它 profile 完全隔离。",
                    "",
                    _LARK_CLI_SOUL_GUIDANCE,
                    "",
                    _GROUP_EXTERNAL_TOOL_SOUL_GUIDANCE,
                    "",
                ]
            ),
            encoding="utf-8",
        )
    else:
        _ensure_soul_guidance(soul_path, _LARK_CLI_SOUL_GUIDANCE)
        _ensure_soul_guidance(soul_path, _GROUP_EXTERNAL_TOOL_SOUL_GUIDANCE)

    (profile_home / "feishu_uat").mkdir(parents=True, exist_ok=True, mode=0o700)

    upstream_profile_home: Path | None = None
    if upstream_profile:
        candidate = shared_home / "profiles" / upstream_profile
        upstream_profile_home = candidate if candidate.exists() else None
    if upstream_profile_home is None:
        upstream_profile_home = _upstream_profile_home_for_owner(owner_open_id, shared_home)
    _sync_default_skills_for_profile(
        profile_home,
        shared_home,
        include_default_skills=True,
        upstream_profile_home=upstream_profile_home,
    )

    marker_path = profile_home / "group_profile.json"
    marker_payload = {
        "kind": "group",
        "chat_id": agent_id,
        "owner_open_id": owner_open_id,
        "display_label": display_label,
        "feishu_auth_disabled": True,
        "source": "webui-agent",
    }
    if not marker_path.exists():
        marker_path.write_text(
            json.dumps(marker_payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    _write_group_profile_env(profile_home, shared_home, None)
    _ensure_shared_profile_file(profile_home, shared_home, "auth.json")


def _write_group_profile_env(profile_home: Path, shared_home: Path, chat_id: str | None) -> None:
    """Materialize <profile>/.env with ONLY model keys and optional FEISHU_HOME_CHANNEL.

    The subprocess env is whitelisted (``_SUBPROCESS_ENV_ALLOWLIST``) and
    does NOT pass model provider API keys, so the AIAgent child reads them
    from ``<profile>/.env`` via its own dotenv path. But the previous
    "copy every shared line" approach also duplicated the credential-vault
    master key, gateway config, and unrelated tool tokens into every group
    profile dir — secret sprawl flagged in code review. We now copy only
    the model provider key/base-url names (``_MODEL_ENV_ALLOWLIST``, the
    same source of truth the model resolver uses) plus the per-group
    FEISHU_HOME_CHANNEL override. ``HERMES_MULTITENANCY_CREDENTIAL_KEY``
    already flows through the allowlisted subprocess env, so it never
    needs to land on disk in a tenant dir.

    Real Feishu group profiles also preempt the gateway's home-channel
    onboarding prompt by declaring the group chat itself as the home channel.
    WebUI group agents pass no chat_id because they are not real Feishu chats.
    Re-runs are idempotent.
    """
    from .agent_real import _MODEL_ENV_ALLOWLIST

    shared_env = shared_home / ".env"
    target = profile_home / ".env"
    base_lines: list[str] = []
    if shared_env.exists():
        try:
            shared_text = shared_env.read_text(encoding="utf-8")
        except (OSError, UnicodeDecodeError) as exc:
            logger.debug(
                "multitenancy: shared .env unreadable (%s); group .env will only set FEISHU_HOME_CHANNEL",
                exc,
            )
            shared_text = ""
        for line in shared_text.splitlines():
            stripped = line.lstrip()
            if not stripped or stripped.startswith("#") or "=" not in stripped:
                continue
            key = stripped.split("=", 1)[0].strip()
            if key in _MODEL_ENV_ALLOWLIST:
                base_lines.append(line)
    if chat_id:
        base_lines.append(f"FEISHU_HOME_CHANNEL={chat_id}")
    # Atomic write: a crash between unlink() and write_text() used to leave
    # the profile with NO .env (all inherited creds gone until reprovision).
    # Write a sibling temp file then os.replace() — the rename is atomic on
    # POSIX and also transparently replaces a stale symlink without ever
    # following it into the shared .env.
    if target.is_symlink():
        target.unlink()
    tmp = target.with_name(target.name + ".tmp")
    tmp.write_text("\n".join(base_lines) + "\n", encoding="utf-8")
    os.replace(tmp, target)


async def _fetch_chat_name(chat_id: str, gateway: Any) -> Optional[str]:
    """Pull chat_name from the live Feishu adapter via its async API."""
    adapter = _get_feishu_adapter(gateway)
    getter = getattr(adapter, "get_chat_info", None) if adapter else None
    if getter is None:
        return None
    try:
        info = await getter(chat_id)
    except Exception as exc:
        logger.debug("multitenancy: get_chat_info(%s) failed: %s", chat_id, exc)
        return None
    if isinstance(info, dict):
        for key in ("name", "chat_name", "title"):
            value = info.get(key)
            if value:
                return str(value)
    return None


async def _fetch_user_display(open_id: str, gateway: Any) -> Optional[str]:
    """Look up a Feishu user's display name through any adapter method
    that the running Feishu adapter exposes. Tested method names are kept
    intentionally small — these match the names the adapter and its sync
    helpers currently expose; new names would need a follow-up edit here."""
    adapter = _get_feishu_adapter(gateway)
    for attr in ("get_user_name", "get_user_display", "get_user_info"):
        fn = getattr(adapter, attr, None)
        if fn is None:
            continue
        try:
            result = await fn(open_id)
        except Exception as exc:
            logger.debug("multitenancy: %s(%s) failed: %s", attr, open_id, exc)
            continue
        if isinstance(result, str) and result:
            return result
        if isinstance(result, dict):
            for key in ("name", "display_name", "nickname"):
                value = result.get(key)
                if value:
                    return str(value)
    return None



def _ensure_auto_profile(
    profile_name: str,
    profile_home: Path,
    *,
    route_key: str,
    sender: str,
) -> None:
    """Create the on-disk profile skeleton required by AIAgent."""
    profile_home.mkdir(parents=True, exist_ok=True)
    shared_home = _shared_home_for_profile(profile_home)

    config_path = profile_home / "config.yaml"
    if config_path.exists():
        _normalize_profile_config_file(config_path, shared_home=shared_home)
    else:
        config_path.write_text(
            _dump_profile_config(_profile_config_from_shared_home(shared_home)),
            encoding="utf-8",
        )

    soul_path = profile_home / "SOUL.md"
    if not soul_path.exists():
        soul_path.write_text(
            "\n".join(
                [
                    f"# Hermes Profile {profile_name}",
                    "",
                    f"You are the dedicated Hermes tenant profile for Feishu route `{route_key}`.",
                    f"The current Feishu sender open_id is `{sender}`.",
                    "Keep tools, files, memory, and responses isolated to this profile.",
                    "Do not claim to be another Hermes profile.",
                    "",
                    _LARK_CLI_SOUL_GUIDANCE,
                    "",
                ]
            ),
            encoding="utf-8",
        )
    else:
        _ensure_soul_guidance(soul_path, _LARK_CLI_SOUL_GUIDANCE)

    _ensure_shared_profile_file(profile_home, shared_home, "auth.json")
    ensure_profile_local_env(profile_home, shared_home)
    _sync_default_skills_for_profile(profile_home, shared_home)


def _sync_default_skills_for_profile(
    profile_home: Path,
    shared_home: Path,
    *,
    include_default_skills: bool = False,
    upstream_profile_home: Path | None = None,
) -> None:
    # Use a relative import so the call works both when the package is loaded
    # as ``hermes_multitenancy`` (pytest/direct install) and when the Hermes
    # plugin loader exposes it as ``hermes_plugins.multitenancy.hermes_multitenancy``.
    from .sync.feishu_org import Employee, _sync_default_profile_skills

    employee = None
    if include_default_skills:
        employee = Employee(
            open_id="",
            user_id=profile_home.name,
            agent_id=profile_home.name,
            profile_name=profile_home.name,
            name=profile_home.name,
            dept_id="",
            dept_name="",
            leader_user_id=None,
        )
    _sync_default_profile_skills(
        profile_home,
        shared_home,
        employee=employee,
        upstream_profile_home=upstream_profile_home,
    )


def _upstream_profile_home_for_owner(owner_open_id: str, shared_home: Path) -> Path | None:
    table = _get_routing_table()
    if table is None or not owner_open_id:
        return None
    try:
        owner = table.resolve_owner_root(owner_open_id) or table.lookup_by_open_id(owner_open_id)
    except Exception as exc:
        logger.debug(
            "multitenancy: owner profile lookup failed owner=%s: %s",
            owner_open_id,
            exc,
        )
        return None
    if owner is None or not owner.profile_name:
        return None
    candidate = shared_home / "profiles" / owner.profile_name
    return candidate if candidate.exists() else None


def _shared_home_for_profile(profile_home: Path) -> Path:
    """Infer the shared Hermes root from a profile path."""
    if profile_home.parent.name == "profiles":
        return profile_home.parent.parent
    return Path.home() / ".hermes"


def _profile_config_from_shared_home(shared_home: Path) -> dict[str, Any]:
    """Build a minimal profile config from the shared Hermes config."""
    config: dict[str, Any] = {}
    shared_config = shared_home / "config.yaml"
    if shared_config.exists():
        try:
            import yaml

            loaded = yaml.safe_load(shared_config.read_text(encoding="utf-8")) or {}
            if isinstance(loaded, dict):
                for key in ("model", "fallback", "platform_toolsets"):
                    value = loaded.get(key)
                    if value:
                        config[key] = value
                feishu_platform = ((loaded.get("platforms") or {}).get("feishu") or None)
                if feishu_platform:
                    config["platforms"] = {"feishu": feishu_platform}
        except Exception as exc:
            logger.debug("multitenancy: failed to read shared config %s: %s", shared_config, exc)

    _apply_lark_cli_profile_defaults(config)
    return _normalize_profile_config(config)


def _normalize_profile_config(config: dict[str, Any]) -> dict[str, Any]:
    model = config.get("model")
    if isinstance(model, dict) and model.get("default"):
        default_model = str(model.get("default") or "").strip()
        provider = str(model.get("provider") or "").strip()
        if default_model and provider and "/" not in default_model:
            model["default"] = f"{provider}/{default_model}"
    return config


def _normalize_profile_config_file(config_path: Path, *, shared_home: Path) -> None:
    try:
        import yaml

        loaded = yaml.safe_load(config_path.read_text(encoding="utf-8")) or {}
    except Exception as exc:
        logger.debug("multitenancy: failed to normalize profile config %s: %s", config_path, exc)
        return
    if not isinstance(loaded, dict):
        return
    before = json.dumps(loaded, sort_keys=True, ensure_ascii=True)
    _merge_shared_feishu_platform(loaded, shared_home)
    _apply_lark_cli_profile_defaults(loaded)
    normalized = _normalize_profile_config(loaded)
    after = json.dumps(normalized, sort_keys=True, ensure_ascii=True)
    if after != before:
        config_path.write_text(_dump_profile_config(normalized), encoding="utf-8")


def _disable_webui_agent_feishu_platform_file(config_path: Path) -> None:
    try:
        import yaml

        loaded = yaml.safe_load(config_path.read_text(encoding="utf-8")) or {}
    except Exception as exc:
        logger.debug("multitenancy: failed to read WebUI agent config %s: %s", config_path, exc)
        return
    if not isinstance(loaded, dict):
        return
    before = json.dumps(loaded, sort_keys=True, ensure_ascii=True)
    _disable_webui_agent_feishu_platform(loaded)
    normalized = _normalize_profile_config(loaded)
    after = json.dumps(normalized, sort_keys=True, ensure_ascii=True)
    if after != before:
        config_path.write_text(_dump_profile_config(normalized), encoding="utf-8")


def _disable_webui_agent_feishu_platform(config: dict[str, Any]) -> None:
    """WebUI group agents use lark-cli bot auth but must not connect Feishu."""
    platforms = config.setdefault("platforms", {})
    if isinstance(platforms, dict):
        platforms["feishu"] = {"enabled": False}


def _apply_lark_cli_profile_defaults(config: dict[str, Any]) -> None:
    """Keep generated Feishu profiles thin: identity/display in Hermes, OAPI in lark-cli."""
    platform_toolsets = config.setdefault("platform_toolsets", {})
    if isinstance(platform_toolsets, dict):
        for platform_key in ("feishu", "api_server", "webui"):
            platform_toolsets[platform_key] = list(_LARK_CLI_PROFILE_TOOLSETS)
    multitenancy = config.setdefault("multitenancy", {})
    if isinstance(multitenancy, dict):
        multitenancy.setdefault("toolsets_mode", "explicit")
        platform_modes = multitenancy.setdefault("platform_toolsets_mode", {})
        if isinstance(platform_modes, dict):
            # feishu must merge (not replace) the platform default toolset.
            # `explicit` returned only ["lark-cli"] and silently dropped the
            # `skills` toolset, so lark-* skill bodies symlinked into the
            # profile never surfaced in the agent manifest (user report:
            # 技能里看不到 lark skills). Keep aligned with api_server/webui.
            for platform_key in ("feishu", "api_server", "webui"):
                explicit_toolsets = _normalize_string_list(platform_toolsets.get(platform_key))
                current_mode = str(platform_modes.get(platform_key) or "").strip().lower()
                if explicit_toolsets == _LARK_CLI_PROFILE_TOOLSETS and current_mode in {
                    "",
                    "explicit",
                    "strict",
                    "replace",
                }:
                    platform_modes[platform_key] = "merge_default"
                else:
                    platform_modes.setdefault(platform_key, "merge_default")


def _normalize_string_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def _ensure_soul_guidance(soul_path: Path, guidance: str) -> None:
    try:
        text = soul_path.read_text(encoding="utf-8")
    except OSError:
        return
    guidance_lines = guidance.splitlines()
    missing_lines = [line for line in guidance_lines if line and line not in text]
    if not missing_lines:
        return
    if guidance_lines and guidance_lines[0] in text:
        suffix = "" if text.endswith("\n") else "\n"
        soul_path.write_text(f"{text}{suffix}" + "\n".join(missing_lines) + "\n", encoding="utf-8")
        return
    suffix = "" if text.endswith("\n") else "\n"
    soul_path.write_text(f"{text}{suffix}\n{guidance}\n", encoding="utf-8")


def _merge_shared_feishu_platform(config: dict[str, Any], shared_home: Path) -> None:
    shared_config = shared_home / "config.yaml"
    if not shared_config.exists():
        return
    try:
        import yaml

        loaded = yaml.safe_load(shared_config.read_text(encoding="utf-8")) or {}
    except Exception as exc:
        logger.debug("multitenancy: failed to read shared Feishu platform config %s: %s", shared_config, exc)
        return
    if not isinstance(loaded, dict):
        return
    feishu_platform = ((loaded.get("platforms") or {}).get("feishu") or None)
    if not feishu_platform:
        return
    platforms = config.setdefault("platforms", {})
    if isinstance(platforms, dict):
        platforms["feishu"] = feishu_platform


def _dump_profile_config(config: dict[str, Any]) -> str:
    try:
        import yaml

        return yaml.safe_dump(config, sort_keys=False, allow_unicode=False)
    except Exception:
        return json.dumps(config, indent=2, ensure_ascii=True) + "\n"


def _ensure_shared_profile_file(profile_home: Path, shared_home: Path, name: str) -> None:
    source = shared_home / name
    target = profile_home / name
    if target.exists() or target.is_symlink() or not source.exists():
        return
    try:
        target.symlink_to(source)
    except OSError:
        shutil.copy2(source, target)


def ensure_profile_local_env(profile_home: Path, shared_home: Path) -> bool:
    """Ensure a tenant profile has a local .env, not a symlink to shared secrets.

    User-scoped UAT and provider credentials are brokered separately; a tenant
    ``.env`` exists for compatibility only and must not point at the service
    root's secret-bearing ``.env``.
    """
    target = profile_home / ".env"
    shared_env = shared_home / ".env"
    if target.is_symlink() and not _symlink_points_to(target, shared_env):
        return False
    if target.exists() and not target.is_symlink():
        _chmod_private_file(target)
        return False
    _write_empty_profile_env(target)
    return True


def repair_profile_local_envs(
    *,
    shared_home: Path | None = None,
    profiles_root: Path | None = None,
    dry_run: bool = False,
) -> dict[str, int]:
    """Backfill profile .env symlinks that point at the shared Hermes .env."""
    root = (shared_home or Path(os.environ.get("HERMES_HOME", "~/.hermes"))).expanduser()
    profiles = profiles_root or (root / "profiles")
    stats = {
        "scanned": 0,
        "repaired": 0,
        "created": 0,
        "planned_repaired": 0,
        "planned_created": 0,
        "kept": 0,
        "skipped_service": 0,
        "skipped_non_shared_symlink": 0,
        "errors": 0,
    }
    if not profiles.exists():
        return stats
    for profile_home in sorted(path for path in profiles.iterdir() if path.is_dir()):
        stats["scanned"] += 1
        if _is_service_profile(profile_home.name):
            stats["skipped_service"] += 1
            continue
        target = profile_home / ".env"
        if target.is_symlink() and not _symlink_points_to(target, root / ".env"):
            stats["skipped_non_shared_symlink"] += 1
            continue
        if target.exists() and not target.is_symlink():
            stats["kept"] += 1
            continue
        was_missing = not target.exists() and not target.is_symlink()
        if dry_run:
            if was_missing:
                stats["planned_created"] += 1
            else:
                stats["planned_repaired"] += 1
            continue
        try:
            ensure_profile_local_env(profile_home, root)
        except OSError:
            logger.exception("multitenancy: failed to repair profile .env profile=%s", profile_home.name)
            stats["errors"] += 1
            continue
        if was_missing:
            stats["created"] += 1
        else:
            stats["repaired"] += 1
    return stats


def _is_service_profile(profile_name: str) -> bool:
    return profile_name == os.environ.get("HERMES_MULTITENANCY_ROUTER_PROFILE", "multitenancy_router")


def _symlink_points_to(path: Path, expected: Path) -> bool:
    try:
        return path.resolve(strict=False) == expected.resolve(strict=False)
    except OSError:
        return False


def _write_empty_profile_env(target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    tmp = target.with_name(f".{target.name}.tmp.{os.getpid()}")
    tmp.write_text("", encoding="utf-8")
    _chmod_private_file(tmp)
    tmp.replace(target)


def _chmod_private_file(path: Path) -> None:
    try:
        path.chmod(0o600)
    except OSError:
        logger.debug("multitenancy: failed to chmod private file %s", path, exc_info=True)


def _profile_name_to_home(profile_name: str) -> Path:
    """Map profile_name to its on-disk profile home directory.

    Mirrors ``hermes_cli/profiles.py`` convention: ``~/.hermes/profiles/<name>``.
    """
    return Path.home() / ".hermes" / "profiles" / profile_name


def _touch_route(sender: str, sender_alt: Optional[str] = None) -> None:
    """Best-effort last_active_at update; no-op if no SQLite table or row.

    Touch both identifiers because auto-provisioned rows are keyed by
    app-scoped ``sender`` while older synced rows may still be keyed by
    ``sender_alt`` / union_id.
    """
    table = _get_routing_table()
    if table is None:
        return
    keys = [sender]
    if sender_alt and sender_alt != sender:
        keys.append(sender_alt)
    for key in keys:
        try:
            table.touch_active(key)
        except Exception as exc:
            logger.debug("multitenancy: touch_active failed: %s", exc)


# -- Lazy singletons (RoutingTable + RuntimePool) ----------------------------


_routing_table: Any = None
_routing_db_path: Optional[str] = None
_pool: Any = None
_session_store: Any = None
_session_db_path: Optional[str] = None  # None → DEFAULT_DB_PATH inside SessionStore


def _get_session_store():
    """Lazy-init module-level SessionStore. Returns None if init fails (in-memory only)."""
    global _session_store
    if _session_store is None:
        try:
            from .sessions import SessionStore
            _session_store = SessionStore(_session_db_path)
        except Exception as exc:
            logger.debug("multitenancy: SessionStore init failed (%s)", exc)
            return None
    return _session_store


def override_session_store(store_or_path) -> None:
    """Test helper — inject a SessionStore (or db path string, or None to disable)."""
    global _session_store, _session_db_path, _session_loaded
    if _session_store is not None and _session_store is not store_or_path:
        try:
            _session_store.close()
        except Exception:
            pass
    _session_loaded.clear()
    if store_or_path is None or isinstance(store_or_path, (str, Path)):
        _session_store = None
        _session_db_path = str(store_or_path) if store_or_path is not None else None
    else:
        _session_store = store_or_path
        _session_db_path = None


def _get_routing_table():
    """Lazy-init module-level RoutingTable. Returns None if init fails."""
    global _routing_table
    if _routing_table is None:
        try:
            from .routing import RoutingTable
            _routing_table = RoutingTable(_routing_db_path)
        except Exception as exc:
            logger.debug("multitenancy: RoutingTable init failed (%s)", exc)
            return None
    return _routing_table


def _get_pool():
    """Lazy-init module-level RuntimePool."""
    global _pool
    if _pool is None:
        from .pool import RuntimePool
        _pool = RuntimePool()
    return _pool


def override_routing_table(db_path: Optional[str | Path]) -> None:
    """Test helper — reset the routing-table singleton, optionally pointing it at a different db."""
    global _routing_table, _routing_db_path
    if _routing_table is not None:
        try:
            _routing_table.close()
        except Exception:
            pass
    _routing_table = None
    _routing_db_path = str(db_path) if db_path is not None else None


def override_pool(pool) -> None:
    """Test helper — inject a custom RuntimePool (or None to reset)."""
    global _pool
    _pool = pool


# -- Adapter resolution ------------------------------------------------------


def _get_feishu_adapter(gateway: Any) -> Any:
    """Pull the FeishuAdapter from the gateway, returning None if unavailable."""
    if gateway is None:
        return None
    adapters = getattr(gateway, "adapters", None)
    if adapters is None:
        return None
    try:
        from gateway.platforms.base import Platform  # type: ignore
        result = adapters.get(Platform.FEISHU)
        if result is not None:
            return result
    except Exception as exc:  # pragma: no cover — only triggers when import fails
        logger.debug("multitenancy: Platform import unavailable (%s)", exc)
    if isinstance(adapters, dict):
        return adapters.get("feishu")
    return None


async def _safe_call(fn, *args, **kwargs):
    """Await fn(*args, **kwargs) whether it is sync or async."""
    result = fn(*args, **kwargs)
    if asyncio.iscoroutine(result):
        return await result
    return result


# Detect Feishu rate-limit errors. Strings vary across SDK versions, so we
# match on broad hints rather than exact codes.
_RATE_LIMIT_HINTS = ("429", "rate limit", "too many requests", "ratelimit")


def _is_rate_limit_error(exc: BaseException) -> bool:
    msg = (str(exc) or "").lower()
    return any(h in msg for h in _RATE_LIMIT_HINTS)


async def _edit_with_retry(adapter, chat_id, message_id, content, *, finalize=False):
    """Wrap adapter.edit_message with exponential backoff on 429.

    On 429: backoff 0.5 → 1.0 → 2.0s, max 3 retries (4 total attempts), then
    log a warning and return None so the caller can continue streaming.
    On non-429: 1 retry after 0.2s, then propagate.
    """
    backoffs = (0.5, 1.0, 2.0)
    for attempt in range(4):
        try:
            if finalize:
                try:
                    return await adapter.edit_message(chat_id, message_id, content, finalize=True)
                except TypeError:
                    return await adapter.edit_message(chat_id, message_id, content)
            return await adapter.edit_message(chat_id, message_id, content)
        except Exception as exc:
            if _is_rate_limit_error(exc):
                if attempt < len(backoffs):
                    logger.debug(
                        "multitenancy: edit_message 429, backoff %ss (attempt %d)",
                        backoffs[attempt], attempt + 1,
                    )
                    await asyncio.sleep(backoffs[attempt])
                    continue
                logger.warning("multitenancy: edit_message rate-limited 4x, giving up")
                return None
            if attempt == 0:
                logger.debug("multitenancy: edit_message non-429 retry: %s", exc)
                await asyncio.sleep(0.2)
                continue
            raise
    return None


def _adapter_supports_streaming_card(adapter) -> bool:
    """Return True when the shared Feishu adapter can drive card streaming."""
    if adapter is None:
        return False
    try:
        from .feishu_cardkit_compat import ensure_feishu_cardkit_streaming

        ensure_feishu_cardkit_streaming(adapter)
    except Exception as exc:
        logger.debug("multitenancy: Feishu CardKit compat install skipped: %s", exc)
    supports = getattr(adapter, "supports_streaming_card", None)
    if callable(supports):
        try:
            return bool(supports())
        except Exception as exc:
            logger.debug("multitenancy: supports_streaming_card failed: %s", exc)
            return False
    return bool(getattr(adapter, "SUPPORTS_STREAMING_CARD", False))


async def _start_feishu_stream_target(
    adapter,
    chat_id,
    *,
    reply_to: Optional[str] = None,
    metadata: Optional[dict[str, Any]] = None,
) -> tuple[str, Optional[str]]:
    """Start a card stream when possible, otherwise create the text placeholder."""
    if _adapter_supports_streaming_card(adapter):
        starter = getattr(adapter, "start_streaming_card", None)
        updater = getattr(adapter, "update_streaming_card", None)
        if callable(starter) and callable(updater):
            try:
                result = await starter(chat_id=chat_id, reply_to=reply_to, metadata=metadata)
            except Exception as exc:
                logger.debug("multitenancy: start_streaming_card failed: %s", exc)
            else:
                message_id = getattr(result, "message_id", None)
                if getattr(result, "success", False) and message_id:
                    logger.info("multitenancy: streaming_card started message_id=%s", message_id)
                    return ("card", str(message_id))
                logger.debug(
                    "multitenancy: start_streaming_card unsuccessful: %s",
                    getattr(result, "error", None),
                )

    placeholder_send = await adapter.send(chat_id, _STREAM_INVISIBLE_PLACEHOLDER, reply_to=reply_to, metadata=metadata)
    message_id = (
        placeholder_send.message_id
        if getattr(placeholder_send, "success", False)
        else None
    )
    return ("edit", str(message_id) if message_id else None)


async def _update_feishu_stream_target(
    adapter, chat_id, message_id, content, *, mode: str, finalize: bool = False
):
    """Update the current Feishu streaming surface."""
    if mode == "card":
        result = await adapter.update_streaming_card(
            chat_id=chat_id,
            message_id=message_id,
            content=content,
            finalize=finalize,
        )
        if not getattr(result, "success", False):
            logger.debug(
                "multitenancy: update_streaming_card unsuccessful: %s",
                getattr(result, "error", None),
            )
        return result
    return await _edit_with_retry(
        adapter, chat_id, message_id, content, finalize=finalize
    )


async def _abort_feishu_stream_target(
    adapter, chat_id, message_id, content, *, mode: str
):
    """Force the current streaming surface into an aborted terminal state."""
    if mode == "card":
        aborter = getattr(adapter, "abort_streaming_card", None)
        if callable(aborter):
            result = await aborter(
                chat_id=chat_id,
                message_id=message_id,
                content=content,
            )
            if not getattr(result, "success", False):
                logger.debug(
                    "multitenancy: abort_streaming_card unsuccessful: %s",
                    getattr(result, "error", None),
                )
            return result
    return await _update_feishu_stream_target(
        adapter,
        chat_id,
        message_id,
        content or "Aborted.",
        mode=mode,
        finalize=True,
    )


async def _run_terminal_stream_update(update_coro, *, label: str):
    """Run terminal card/edit update even if the caller is being cancelled."""
    task = asyncio.create_task(update_coro)
    try:
        return await asyncio.shield(task)
    except asyncio.CancelledError:
        try:
            result = await task
            logger.info("multitenancy: %s completed while task was cancelling", label)
            return result
        except Exception as exc:
            logger.debug("multitenancy: %s failed while task was cancelling: %s", label, exc)
        raise


async def _update_feishu_stream_reasoning(
    adapter, chat_id, message_id, content, *, mode: str
):
    """Update reasoning/status text without polluting the final answer block."""
    if mode == "card":
        updater = getattr(adapter, "update_streaming_card_reasoning", None)
        if callable(updater):
            return await updater(
                chat_id=chat_id,
                message_id=message_id,
                content=content,
            )
    return await _update_feishu_stream_target(
        adapter, chat_id, message_id, content, mode=mode
    )


async def _update_feishu_stream_status(
    adapter, chat_id, message_id, content, *, mode: str
):
    """Update an ephemeral in-progress status without retaining it as reasoning."""
    if mode == "card":
        updater = getattr(adapter, "update_streaming_card_status", None)
        if callable(updater):
            return await updater(
                chat_id=chat_id,
                message_id=message_id,
                content=content,
            )
    return await _update_feishu_stream_reasoning(
        adapter, chat_id, message_id, content, mode=mode
    )


async def _update_feishu_stream_tool_event(
    adapter, chat_id, message_id, payload, *, mode: str, completed: bool, profile_home: Optional[Path] = None
):
    """Update active/completed tool state on the streaming surface."""
    payload = _sanitize_tool_event_payload(payload, profile_home)
    tool_name = str(payload.get("name") or payload.get("tool_name") or "tool")
    if mode == "card":
        method_name = (
            "update_streaming_card_tool_completed"
            if completed
            else "update_streaming_card_tool_started"
        )
        updater = getattr(adapter, method_name, None)
        if callable(updater):
            if completed:
                return await updater(
                    chat_id=chat_id,
                    message_id=message_id,
                    tool_name=tool_name,
                    duration=payload.get("duration"),
                    is_error=bool(payload.get("is_error")),
                )
            return await updater(
                chat_id=chat_id,
                message_id=message_id,
                tool_name=tool_name,
                preview=payload.get("preview"),
                args=payload.get("args"),
            )

    status = (
        f"✅ 工具完成: {tool_name}"
        if completed and not payload.get("is_error")
        else f"⚠️ 工具失败: {tool_name}"
        if completed
        else f"🔧 正在调用工具: {tool_name}"
    )
    return await _update_feishu_stream_target(
        adapter, chat_id, message_id, status, mode=mode
    )


# Streaming-card flush throttles match openclaw-lark:
# CARDKIT_MS=100 for cardElement.content and PATCH_MS=1500 for legacy edits.
# The character thresholds are Hermes' local stream-consumer coalescing floor;
# time-based throttle is the cross-project contract.
_STREAM_CONTENT_MIN_CHARS = 120
_STREAM_CONTENT_MIN_SECONDS = 1.5
_STREAM_CARDKIT_CONTENT_MIN_CHARS = 30
_STREAM_CARDKIT_CONTENT_MIN_SECONDS = 0.1
_STREAM_THINKING_MIN_SECONDS = 2.0
_STREAM_CARD_REASONING_MIN_CHARS = 100
_STREAM_CARD_REASONING_MIN_SECONDS = 2.0
_STREAM_CARD_PRIME_STATUS = ""
_STREAM_INVISIBLE_PLACEHOLDER = "\u200b"
_STREAM_CARD_IDLE_HEARTBEAT_SECONDS = 2.5
_STREAM_STATUS_ANIMATION_MARKERS = ("\u200b", "\u200c", "\u200d", "\ufeff")
_STREAM_ABORT_FALLBACK = "Aborted."
# Soft per-card segment target for the legacy CardKit compat path. This is not
# a Feishu/CardKit platform limit; shared GatewayStreamConsumer uses its own
# adapter-specific splitting and must receive the complete stream.
_STREAM_MAX_VISIBLE_CHARS = 3_000


def _is_aiagent_stream_idle_timeout(exc: BaseException) -> bool:
    return "AIAgent subprocess produced no stream events" in str(exc)


def _aiagent_stream_timeout_notice(exc: BaseException) -> str:
    return (
        "\n\n⚠️ 当前任务长时间没有新的运行事件，已停止本次流式执行。\n"
        f"{exc}"
    )


def _stream_card_idle_status(tick: int) -> str:
    """Return a changing pre-token status marker with no visible waiting text."""
    marker = _STREAM_STATUS_ANIMATION_MARKERS[(max(1, int(tick)) - 1) % len(_STREAM_STATUS_ANIMATION_MARKERS)]
    return marker


def _strip_stream_status_animation_markers(text: str) -> str:
    result = str(text or "")
    for marker in _STREAM_STATUS_ANIMATION_MARKERS:
        result = result.replace(marker, "")
    return result


async def _stream_into_feishu_shared_consumer(
    adapter,
    chat_id,
    profile_name,
    profile_home,
    event,
    *,
    gateway: Any = None,
    messages: Optional[list[dict]] = None,
) -> Optional[str]:
    """Stream Feishu card output through Hermes' shared GatewayStreamConsumer.

    Returns None when the shared card surface cannot be started, allowing the
    caller to fall back to the legacy text-edit transport.
    """
    if GatewayStreamConsumer is None or StreamConsumerConfig is None:
        return None
    required_methods = (
        "ensure_streaming_card_started",
        "run",
        "on_delta",
        "finish",
        "update_streaming_card_status",
        "update_streaming_card_reasoning",
        "update_streaming_card_tool_started",
        "update_streaming_card_tool_completed",
    )
    if not all(hasattr(GatewayStreamConsumer, method) for method in required_methods):
        logger.debug("multitenancy: shared GatewayStreamConsumer lacks card methods; using adapter surface")
        return None

    import time
    from .agent_real import stream_run_agent, real_run_agent
    from .runtime import _PROFILE_HOME_VAR

    stream_started_at = time.monotonic()
    metadata = _thread_metadata_for_media_delivery(gateway, event) if gateway is not None else None
    reply_to = _event_reply_to_message_id(event)
    consumer = GatewayStreamConsumer(
        adapter,
        chat_id,
        StreamConsumerConfig(
            edit_interval=_STREAM_CARDKIT_CONTENT_MIN_SECONDS,
            buffer_threshold=_STREAM_CARDKIT_CONTENT_MIN_CHARS,
            cursor=" ▉",
        ),
        metadata=metadata,
        initial_reply_to_id=reply_to,
    )
    required_consumer_methods = (
        "ensure_streaming_card_started",
        "update_streaming_card_status",
        "update_streaming_card_reasoning",
        "update_streaming_card_tool_started",
        "update_streaming_card_tool_completed",
    )
    if not all(hasattr(consumer, name) for name in required_consumer_methods):
        logger.debug(
            "multitenancy: GatewayStreamConsumer lacks shared-card methods; falling back"
        )
        return None
    consumer_task: Optional[asyncio.Task] = None
    idle_heartbeat_task: Optional[asyncio.Task] = None
    terminal_update_sent = False
    first_agent_event_seen = False
    content_delta_seen = False
    content = ""
    thinking = ""
    last_reasoning_edit = 0.0
    last_reasoning_len = 0

    async def _idle_card_heartbeat() -> None:
        tick = 2
        while True:
            await asyncio.sleep(_STREAM_CARD_IDLE_HEARTBEAT_SECONDS)
            if content_delta_seen:
                return
            try:
                await consumer.update_streaming_card_status(_stream_card_idle_status(tick))
            except asyncio.CancelledError:
                raise
            except Exception as exc:
                logger.debug("multitenancy: shared card idle heartbeat failed: %s", exc)
            tick += 1

    async def _stop_idle_card_heartbeat() -> None:
        if idle_heartbeat_task is None or idle_heartbeat_task.done():
            return
        idle_heartbeat_task.cancel()
        try:
            await idle_heartbeat_task
        except asyncio.CancelledError:
            pass

    def _abort_content() -> str:
        raw = content if content else (thinking if thinking else _STREAM_ABORT_FALLBACK)
        return _clean_stream_display_text(raw, profile_home)

    async def _finish_consumer() -> None:
        if consumer_task is None:
            return
        consumer.finish()
        await consumer_task

    try:
        start_task = asyncio.create_task(consumer.ensure_streaming_card_started())
        try:
            started = await asyncio.shield(start_task)
        except asyncio.CancelledError:
            try:
                started = await start_task
            except Exception as exc:
                logger.debug("multitenancy: shared card start failed while cancelling: %s", exc)
                started = False
            if started:
                aborter = getattr(consumer, "abort_streaming_card", None)
                if aborter is not None:
                    try:
                        await aborter(_STREAM_ABORT_FALLBACK)
                    except Exception as exc:
                        logger.debug("multitenancy: shared card abort-after-start failed: %s", exc)
            raise

        if not started:
            logger.debug("multitenancy: shared card start unavailable; falling back")
            return None

        logger.info(
            "multitenancy: shared stream card ready elapsed=%.3fs",
            time.monotonic() - stream_started_at,
        )
        consumer_task = asyncio.create_task(consumer.run())

        prime_task = asyncio.create_task(
            consumer.update_streaming_card_status(_stream_card_idle_status(1))
        )
        try:
            await asyncio.shield(prime_task)
        except asyncio.CancelledError:
            try:
                await prime_task
            except Exception as exc:
                logger.debug("multitenancy: shared card prime failed while cancelling: %s", exc)
            aborter = getattr(consumer, "abort_streaming_card", None)
            if aborter is not None:
                try:
                    await aborter(_STREAM_ABORT_FALLBACK)
                except Exception as exc:
                    logger.debug("multitenancy: shared card abort-after-prime failed: %s", exc)
            raise

        idle_heartbeat_task = asyncio.create_task(_idle_card_heartbeat())

        token = _PROFILE_HOME_VAR.set(profile_home)
        try:
            try:
                async for kind, delta in stream_run_agent(event, profile_home, messages=messages):
                    if not first_agent_event_seen:
                        first_agent_event_seen = True
                        logger.info(
                            "multitenancy: shared stream first agent event kind=%s total=%.3fs",
                            kind,
                            time.monotonic() - stream_started_at,
                        )

                    if kind == "thinking":
                        thinking += str(delta or "")
                        now = time.monotonic()
                        if (
                            not last_reasoning_len
                            or len(thinking) - last_reasoning_len >= _STREAM_CARD_REASONING_MIN_CHARS
                            or now - last_reasoning_edit >= _STREAM_CARD_REASONING_MIN_SECONDS
                        ):
                            await consumer.update_streaming_card_reasoning(thinking)
                            last_reasoning_len = len(thinking)
                            last_reasoning_edit = now
                        continue

                    if kind == "status":
                        text = str(delta or "").strip()
                        if text:
                            try:
                                await consumer.update_streaming_card_status(text)
                            except Exception as exc:
                                logger.debug("multitenancy: shared card status update failed: %s", exc)
                        continue

                    if kind == "tool_started":
                        payload = _sanitize_tool_event_payload(delta, profile_home)
                        await consumer.update_streaming_card_tool_started(
                            str(payload.get("name") or payload.get("tool_name") or "tool"),
                            preview=payload.get("preview"),
                            args=payload.get("args"),
                        )
                        continue

                    if kind == "tool_completed":
                        payload = _sanitize_tool_event_payload(delta, profile_home)
                        await consumer.update_streaming_card_tool_completed(
                            str(payload.get("name") or payload.get("tool_name") or "tool"),
                            duration=payload.get("duration"),
                            is_error=bool(payload.get("is_error")),
                        )
                        continue

                    if kind == "approval_required":
                        await _handle_child_approval_required(adapter, chat_id, delta)
                        try:
                            await consumer.update_streaming_card_status("等待用户审批: /approve 或 /deny")
                        except Exception as exc:
                            logger.debug("multitenancy: approval status update failed: %s", exc)
                        continue

                    if kind == "approval_resolved":
                        if isinstance(delta, dict):
                            _clear_pending_approval(delta)
                        continue

                    if kind == "done":
                        continue

                    piece = str(delta or "")
                    if not piece:
                        continue
                    if thinking and len(thinking) > last_reasoning_len:
                        await consumer.update_streaming_card_reasoning(thinking)
                        last_reasoning_len = len(thinking)
                        last_reasoning_edit = time.monotonic()
                    content += piece
                    consumer.on_delta(_clean_stream_delta_text(piece, profile_home))
                    content_delta_seen = True
            except Exception as exc:
                if _is_aiagent_stream_idle_timeout(exc):
                    logger.warning("multitenancy: shared streaming stopped on idle timeout: %s", exc)
                    timeout_notice = _aiagent_stream_timeout_notice(exc)
                    content += timeout_notice
                    consumer.on_delta(_clean_stream_display_text(timeout_notice, profile_home))
                    content_delta_seen = True
                else:
                    logger.info("multitenancy: shared streaming failed (%s) — falling back to non-stream", exc)
                    try:
                        content = await real_run_agent(event, profile_home, messages=messages)
                    except Exception as fallback_exc:
                        logger.warning("multitenancy: LLM fully unavailable: %s", fallback_exc)
                        content = (
                            "⚠️ 模型暂时不可用 (LLM provider rejected the request).\n"
                            "请检查 profile 的 config.yaml 模型/凭据, 或稍后再试。"
                        )
                    if not content_delta_seen:
                        consumer.on_delta(_clean_stream_display_text(content, profile_home))
                        content_delta_seen = True
        finally:
            _PROFILE_HOME_VAR.reset(token)
            await _stop_idle_card_heartbeat()

        full = content if content else (thinking if thinking else "(empty response)")
        if not content_delta_seen:
            consumer.on_delta(_clean_stream_display_text(full, profile_home))

        await _finish_consumer()
        terminal_update_sent = True
        return full
    except asyncio.CancelledError:
        await _stop_idle_card_heartbeat()
        if not terminal_update_sent:
            aborter = getattr(consumer, "abort_streaming_card", None)
            if aborter is not None:
                try:
                    await _run_terminal_stream_update(
                        aborter(_abort_content()),
                        label="shared stream abort update",
                    )
                except asyncio.CancelledError:
                    raise
                except Exception as abort_exc:
                    logger.debug("multitenancy: shared stream abort update failed: %s", abort_exc)
        if consumer_task is not None:
            consumer.finish()
            try:
                await consumer_task
            except Exception:
                pass
        raise


async def _stream_into_feishu(
    adapter,
    chat_id,
    profile_name,
    profile_home,
    event,
    *,
    gateway: Any = None,
    messages: Optional[list[dict]] = None,
) -> str:
    """Stream LLM tokens into Feishu.

    Uses OpenClaw-style interactive card streaming when the shared Feishu
    adapter supports it. Falls back to the legacy text placeholder +
    ``edit_message`` loop for older adapters.

    Falls back to a single non-streamed send() if streaming returns empty or
    the initial stream target fails. Returns the final concatenated text.

    ``messages`` (optional): full conversation including prior history. When
    omitted the runner builds a single-turn system+user prompt from the event.
    """
    import time
    from .agent_real import stream_run_agent, real_run_agent
    from .runtime import _PROFILE_HOME_VAR

    stream_started_at = time.monotonic()
    reply_to = _event_reply_to_message_id(event)
    metadata = _thread_metadata_for_media_delivery(gateway, event) if gateway is not None else None

    # Without an adapter we can still produce text (used in unit tests).
    if adapter is None:
        token = _PROFILE_HOME_VAR.set(profile_home)
        try:
            content_parts: list[str] = []
            async for kind, c in stream_run_agent(event, profile_home, messages=messages):
                if kind == "content":
                    content_parts.append(c)
            return (
                "".join(content_parts)
                or await real_run_agent(event, profile_home, messages=messages)
            )
        finally:
            _PROFILE_HOME_VAR.reset(token)

    if _adapter_supports_streaming_card(adapter):
        shared_response = await _stream_into_feishu_shared_consumer(
            adapter,
            chat_id,
            profile_name,
            profile_home,
            event,
            gateway=gateway,
            messages=messages,
        )
        if shared_response is not None:
            return shared_response

    stream_mode = "edit"
    placeholder_id: Optional[str] = None
    target_ready_at = stream_started_at
    thinking = ""
    content = ""
    full_content = ""
    last_edit_time = 0.0
    last_render_len = 0
    last_reasoning_render_len = 0
    content_started = False
    first_agent_event_seen = False
    terminal_update_sent = False
    card_reasoning_sent = False
    idle_heartbeat_task: Optional[asyncio.Task] = None

    async def _idle_card_heartbeat() -> None:
        tick = 2
        while True:
            await asyncio.sleep(_STREAM_CARD_IDLE_HEARTBEAT_SECONDS)
            if content_started or placeholder_id is None:
                return
            try:
                await _update_feishu_stream_status(
                    adapter,
                    chat_id,
                    placeholder_id,
                    _stream_card_idle_status(tick),
                    mode=stream_mode,
                )
            except asyncio.CancelledError:
                raise
            except Exception as exc:
                logger.debug("multitenancy: card idle heartbeat failed: %s", exc)
            tick += 1

    async def _stop_idle_card_heartbeat() -> None:
        if idle_heartbeat_task is None or idle_heartbeat_task.done():
            return
        idle_heartbeat_task.cancel()
        try:
            await idle_heartbeat_task
        except asyncio.CancelledError:
            pass

    def render() -> str:
        if content:
            return _clean_stream_display_text(content, profile_home)
        preview = thinking[-160:].strip() if thinking else ""
        return _clean_stream_display_text(preview, profile_home) if preview else _STREAM_INVISIBLE_PLACEHOLDER

    def abort_content() -> str:
        raw = content if content else (thinking if thinking else _STREAM_ABORT_FALLBACK)
        return _clean_stream_display_text(raw, profile_home)

    async def _flush_current_segment(*, finalize: bool) -> None:
        nonlocal last_edit_time, last_render_len, terminal_update_sent
        if placeholder_id is None:
            return
        rendered = render()
        try:
            await _run_terminal_stream_update(
                _update_feishu_stream_target(
                    adapter,
                    chat_id,
                    placeholder_id,
                    rendered,
                    mode=stream_mode,
                    finalize=finalize,
                ),
                label="stream segment update",
            )
            last_edit_time = time.monotonic()
            last_render_len = len(rendered)
            if finalize:
                terminal_update_sent = True
        except Exception as exc:
            logger.debug("multitenancy: stream segment update failed: %s", exc)

    async def _start_next_stream_segment() -> None:
        nonlocal stream_mode, placeholder_id, content, content_started
        nonlocal last_edit_time, last_render_len, terminal_update_sent
        if placeholder_id is not None:
            await _flush_current_segment(finalize=True)
        stream_mode, placeholder_id = await _start_feishu_stream_target(
            adapter,
            chat_id,
            reply_to=reply_to,
            metadata=metadata,
        )
        terminal_update_sent = False
        content = ""
        content_started = False
        last_edit_time = time.monotonic()
        last_render_len = 0
        if placeholder_id is None:
            return
        if stream_mode == "card":
            try:
                await _update_feishu_stream_status(
                    adapter,
                    chat_id,
                    placeholder_id,
                    _STREAM_INVISIBLE_PLACEHOLDER,
                    mode=stream_mode,
                )
            except Exception as exc:
                logger.debug("multitenancy: continuation card prime update failed: %s", exc)

    try:
        # Create/send can complete remotely after this task is cancelled. Shield
        # it so we can still obtain the message_id and close the card instead of
        # leaving a Generating card behind.
        start_task = asyncio.create_task(
            _start_feishu_stream_target(
                adapter,
                chat_id,
                reply_to=reply_to,
                metadata=metadata,
            )
        )
        try:
            stream_mode, placeholder_id = await asyncio.shield(start_task)
        except asyncio.CancelledError:
            try:
                stream_mode, placeholder_id = await start_task
                logger.info(
                    "multitenancy: stream target start completed while task was cancelling "
                    "mode=%s message_id=%s",
                    stream_mode,
                    placeholder_id,
                )
            except Exception as exc:
                logger.debug("multitenancy: stream target start failed while cancelling: %s", exc)
            raise

        target_ready_at = time.monotonic()
        logger.info(
            "multitenancy: stream target ready mode=%s message_id=%s elapsed=%.3fs",
            stream_mode,
            placeholder_id,
            target_ready_at - stream_started_at,
        )

        if placeholder_id is None:
            # Couldn't get a message to edit — degrade to one-shot non-stream.
            text = await real_run_agent(event, profile_home, messages=messages)
            await adapter.send(chat_id, text, reply_to=reply_to, metadata=metadata)
            return text

        if stream_mode == "card":
            try:
                await _update_feishu_stream_status(
                    adapter,
                    chat_id,
                    placeholder_id,
                    _stream_card_idle_status(1),
                    mode=stream_mode,
                )
                logger.info(
                    "multitenancy: stream card primed message_id=%s elapsed=%.3fs",
                    placeholder_id,
                    time.monotonic() - target_ready_at,
                )
                idle_heartbeat_task = asyncio.create_task(_idle_card_heartbeat())
            except asyncio.CancelledError:
                raise
            except Exception as exc:
                logger.debug("multitenancy: card prime update failed: %s", exc)

        token = _PROFILE_HOME_VAR.set(profile_home)
        try:
            try:
                async for kind, delta in stream_run_agent(event, profile_home, messages=messages):
                    if not first_agent_event_seen:
                        first_agent_event_seen = True
                        logger.info(
                            "multitenancy: stream first agent event kind=%s "
                            "since_target=%.3fs total=%.3fs",
                            kind,
                            time.monotonic() - target_ready_at,
                            time.monotonic() - stream_started_at,
                        )
                    if kind == "thinking":
                        thinking += str(delta or "")
                        if stream_mode == "card" and thinking:
                            now = time.monotonic()
                            should_update_card_reasoning = (
                                not card_reasoning_sent
                                or len(thinking) - last_reasoning_render_len >= _STREAM_CARD_REASONING_MIN_CHARS
                                or now - last_edit_time >= _STREAM_CARD_REASONING_MIN_SECONDS
                            )
                            if should_update_card_reasoning:
                                try:
                                    await _update_feishu_stream_reasoning(
                                        adapter,
                                        chat_id,
                                        placeholder_id,
                                        thinking,
                                        mode=stream_mode,
                                    )
                                    card_reasoning_sent = True
                                    last_reasoning_render_len = len(thinking)
                                except Exception as exc:
                                    logger.debug("multitenancy: card reasoning update failed: %s", exc)
                                last_edit_time = now
                                last_render_len = len(render())
                            continue
                    elif kind == "tool_started":
                        try:
                            await _update_feishu_stream_tool_event(
                                adapter,
                                chat_id,
                                placeholder_id,
                                delta,
                                mode=stream_mode,
                                completed=False,
                                profile_home=profile_home,
                            )
                        except Exception as exc:
                            logger.debug("multitenancy: card tool-start update failed: %s", exc)
                        last_edit_time = time.monotonic()
                        last_render_len = len(render())
                        continue
                    elif kind == "tool_completed":
                        try:
                            await _update_feishu_stream_tool_event(
                                adapter,
                                chat_id,
                                placeholder_id,
                                delta,
                                mode=stream_mode,
                                completed=True,
                                profile_home=profile_home,
                            )
                        except Exception as exc:
                            logger.debug("multitenancy: card tool-complete update failed: %s", exc)
                        last_edit_time = time.monotonic()
                        last_render_len = len(render())
                        continue
                    elif kind == "approval_required":
                        await _handle_child_approval_required(adapter, chat_id, delta)
                        try:
                            await _update_feishu_stream_status(
                                adapter,
                                chat_id,
                                placeholder_id,
                                "等待用户审批: /approve 或 /deny",
                                mode=stream_mode,
                            )
                        except Exception as exc:
                            logger.debug("multitenancy: approval status update failed: %s", exc)
                        last_edit_time = time.monotonic()
                        last_render_len = len(render())
                        continue
                    elif kind == "approval_resolved":
                        if isinstance(delta, dict):
                            _clear_pending_approval(delta)
                        continue
                    elif kind == "status":
                        status_text = str(delta or "").strip()
                        if status_text:
                            try:
                                await _update_feishu_stream_status(
                                    adapter,
                                    chat_id,
                                    placeholder_id,
                                    status_text,
                                    mode=stream_mode,
                                )
                            except Exception as exc:
                                logger.debug("multitenancy: stream status update failed: %s", exc)
                            last_edit_time = time.monotonic()
                            last_render_len = len(render())
                        continue
                    elif kind == "done":
                        continue
                    else:
                        piece = str(delta or "")
                        if not piece:
                            continue
                        full_content += piece
                        while piece:
                            remaining = _STREAM_MAX_VISIBLE_CHARS - len(content)
                            if remaining <= 0:
                                logger.info(
                                    "multitenancy: stream content segment finalized "
                                    "message_id=%s max_chars=%s",
                                    placeholder_id,
                                    _STREAM_MAX_VISIBLE_CHARS,
                                )
                                await _start_next_stream_segment()
                                if placeholder_id is None:
                                    break
                                remaining = _STREAM_MAX_VISIBLE_CHARS
                            segment_piece = piece[:remaining]
                            piece = piece[remaining:]
                            content += segment_piece
                            rendered = render()
                            now = time.monotonic()
                            if not content_started:
                                # Force an immediate edit on phase transition so the user
                                # sees the answer start the moment reasoning ends.
                                content_started = True
                                try:
                                    await _update_feishu_stream_target(
                                        adapter,
                                        chat_id,
                                        placeholder_id,
                                        rendered,
                                        mode=stream_mode,
                                    )
                                except Exception as exc:
                                    logger.debug(
                                        "multitenancy: phase-transition stream update failed: %s",
                                        exc,
                                    )
                                last_edit_time = time.monotonic()
                                last_render_len = len(rendered)
                            elif (
                                piece
                                or len(rendered) - last_render_len >= _STREAM_CONTENT_MIN_CHARS
                                or now - last_edit_time >= _STREAM_CONTENT_MIN_SECONDS
                            ):
                                try:
                                    await _update_feishu_stream_target(
                                        adapter,
                                        chat_id,
                                        placeholder_id,
                                        rendered,
                                        mode=stream_mode,
                                    )
                                except Exception as exc:
                                    logger.debug("multitenancy: stream update mid-stream failed: %s", exc)
                                last_edit_time = now
                                last_render_len = len(rendered)
                            if piece:
                                logger.info(
                                    "multitenancy: stream content segment split "
                                    "message_id=%s max_chars=%s",
                                    placeholder_id,
                                    _STREAM_MAX_VISIBLE_CHARS,
                                )
                                await _start_next_stream_segment()
                                if placeholder_id is None:
                                    break
                        continue

                    now = time.monotonic()
                    rendered = render()
                    if content_started:
                        should_edit = (
                            len(rendered) - last_render_len >= _STREAM_CONTENT_MIN_CHARS
                            or now - last_edit_time >= _STREAM_CONTENT_MIN_SECONDS
                        )
                    else:
                        # Reasoning phase — heartbeat-only edits, no char threshold.
                        should_edit = now - last_edit_time >= _STREAM_THINKING_MIN_SECONDS
                    if should_edit:
                        try:
                            await _update_feishu_stream_target(
                                adapter,
                                chat_id,
                                placeholder_id,
                                rendered,
                                mode=stream_mode,
                            )
                        except Exception as exc:
                            logger.debug("multitenancy: stream update mid-stream failed: %s", exc)
                        last_edit_time = now
                        last_render_len = len(rendered)
            except Exception as exc:
                if _is_aiagent_stream_idle_timeout(exc):
                    logger.warning("multitenancy: streaming stopped on idle timeout: %s", exc)
                    timeout_notice = _aiagent_stream_timeout_notice(exc)
                    content += timeout_notice
                    full_content += timeout_notice
                    try:
                        await _update_feishu_stream_status(
                            adapter,
                            chat_id,
                            placeholder_id,
                            "任务长时间没有新的运行事件，已停止。",
                            mode=stream_mode,
                        )
                    except Exception as status_exc:
                        logger.debug("multitenancy: idle-timeout status update failed: %s", status_exc)
                else:
                    logger.info("multitenancy: streaming failed (%s) — falling back to non-stream", exc)
                    try:
                        content = await real_run_agent(event, profile_home, messages=messages)
                        full_content = content
                    except Exception as fallback_exc:
                        # Both stream + non-stream LLM paths failed (e.g. region block,
                        # exhausted credentials). Surface a user-visible error instead
                        # of leaving the "..." placeholder hanging.
                        logger.warning("multitenancy: LLM fully unavailable: %s", fallback_exc)
                        content = (
                            "⚠️ 模型暂时不可用 (LLM provider rejected the request).\n"
                            "请检查 profile 的 config.yaml 模型/凭据, 或稍后再试。"
                        )
                        full_content = content
        finally:
            _PROFILE_HOME_VAR.reset(token)
            await _stop_idle_card_heartbeat()

        full = full_content or content or (thinking if thinking else "(empty response)")
        display_current = _clean_stream_display_text(content or full, profile_home)

        # 3. Final commit. finalize=True signals end of stream to Feishu.
        try:
            await _run_terminal_stream_update(
                _update_feishu_stream_target(
                    adapter,
                    chat_id,
                    placeholder_id,
                    display_current,
                    mode=stream_mode,
                    finalize=True,
                ),
                label="stream final update",
            )
            terminal_update_sent = True
        except Exception as exc:
            logger.debug("multitenancy: final stream update failed: %s", exc)

        return full
    except asyncio.CancelledError:
        if placeholder_id is not None and not terminal_update_sent:
            full = abort_content()
            logger.info(
                "multitenancy: stream cancelled; aborting target mode=%s message_id=%s content_len=%s",
                stream_mode,
                placeholder_id,
                len(full),
            )
            try:
                await _run_terminal_stream_update(
                    _abort_feishu_stream_target(
                        adapter,
                        chat_id,
                        placeholder_id,
                        full,
                        mode=stream_mode,
                    ),
                    label="stream abort update",
                )
            except asyncio.CancelledError:
                raise
            except Exception as abort_exc:
                logger.debug("multitenancy: stream abort update failed: %s", abort_exc)
        raise


def _log_task_failure(task: asyncio.Task) -> None:
    """Done-callback for fire-and-forget tasks — surfaces silent exceptions."""
    if task.cancelled():
        return
    exc = task.exception()
    if exc is not None:
        logger.error("multitenancy: background task crashed: %r", exc)
